"""stamp_signature_above_name — anchor the approval signature in an authored docx."""

import zipfile
from pathlib import Path
from typing import Any

from docx import Document

from app.core.docx_engine import stamp_signature_above_name

# ---------------------------------------------------------------------------
# Shared helpers
# ---------------------------------------------------------------------------

_ANCHOR_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


def _make_letter(tmp_path: Path, closing_name: str) -> Path:
    doc = Document()
    doc.add_paragraph("نص الكتاب التجريبي")
    doc.add_paragraph("")  # signature gap
    doc.add_paragraph(closing_name)
    doc.add_paragraph("مدير مشروع")
    p = tmp_path / "letter.docx"
    doc.save(str(p))
    return p


def _make_sig(tmp_path: Path) -> Path:
    from PIL import Image

    sig = tmp_path / "sig.png"
    Image.new("RGBA", (60, 30), (0, 0, 200, 255)).save(sig)
    return sig


def _document_xml(docx: Path) -> bytes:
    with zipfile.ZipFile(docx) as z:
        return z.read("word/document.xml")


def _para_has_anchor(para: Any) -> bool:
    return bool(para._p.findall(f".//{{{_ANCHOR_NS}}}anchor"))


def _make_report_sig_block(tmp_path: Path) -> Path:
    """Docx with the four-paragraph Report signature block:
    [0] الاسم : فلان
    [1] المسمى الوظيفي : كذا
    [2] التوقيـــــع:        (tatweel-stretched label)
    [3] ""                  (blank — expected anchor for the drawing)
    """
    doc = Document()
    doc.add_paragraph("الاسم : فلان")
    doc.add_paragraph("المسمى الوظيفي : كذا")
    doc.add_paragraph("التوقيـــــع:")  # tatweel in the label
    doc.add_paragraph("")  # blank — drawing lands here
    p = tmp_path / "report_sig.docx"
    doc.save(str(p))
    return p


# ---------------------------------------------------------------------------
# Existing General Book tests (regression guard — must stay green)
# ---------------------------------------------------------------------------


def test_stamps_on_exact_name(tmp_path: Path) -> None:
    docx = _make_letter(tmp_path, "سعيد راشد اليحيائي")
    ok = stamp_signature_above_name(
        docx, str(_make_sig(tmp_path)), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert ok
    # The float is an anchored drawing in the paragraph above the name.
    assert b"<wp:anchor" in _document_xml(docx)


def test_stamps_despite_tatweel_stretching(tmp_path: Path) -> None:
    # Hand-made templates stretch names with tatweel: سعيــــد راشــــد
    docx = _make_letter(tmp_path, "سعيــــــــــد راشــــــــــد اليحيائــــــــــي")
    ok = stamp_signature_above_name(
        docx, str(_make_sig(tmp_path)), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert ok
    assert b"<wp:anchor" in _document_xml(docx)


def test_falls_back_to_last_paragraph_when_name_missing(tmp_path: Path) -> None:
    docx = _make_letter(tmp_path, "اسم آخر تماماً")
    ok = stamp_signature_above_name(
        docx, str(_make_sig(tmp_path)), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert ok  # fallback anchor, still stamped
    assert b"<wp:anchor" in _document_xml(docx)


def test_noop_without_signature_file(tmp_path: Path) -> None:
    docx = _make_letter(tmp_path, "سعيد راشد اليحيائي")
    ok = stamp_signature_above_name(
        docx, str(tmp_path / "missing.png"), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert not ok


def test_cc_line_does_not_steal_the_anchor(tmp_path: Path) -> None:
    """A CC line AFTER the closing block can mention the manager's name — the
    exact-equality pass must win so the float lands above the NAME, not above
    the CC line (review M/S1)."""
    doc = Document()
    doc.add_paragraph("نص الكتاب")  # 0
    doc.add_paragraph("")  # 1 — signature gap (expected anchor)
    doc.add_paragraph("سعيد راشد اليحيائي")  # 2
    doc.add_paragraph("مدير مشروع")  # 3
    doc.add_paragraph("نسخة إلى: مكتب سعيد راشد اليحيائي")  # 4 — decoy
    p = tmp_path / "cc.docx"
    doc.save(str(p))
    ok = stamp_signature_above_name(
        p, str(_make_sig(tmp_path)), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert ok
    reloaded = Document(str(p))
    with_drawing = [
        i
        for i, para in enumerate(reloaded.paragraphs)
        if para._p.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor"
        )
    ]
    assert with_drawing == [1]  # the gap above the name — not the title (2/3) or CC (4)


def test_stamps_inside_tables(tmp_path: Path) -> None:
    """Word-paste-into-tables letters keep every paragraph inside table cells —
    doc.paragraphs alone finds nothing; the stamp must search cells too."""
    doc = Document()
    table = doc.add_table(rows=3, cols=1)
    table.rows[0].cells[0].paragraphs[0].add_run("نص داخل جدول")
    table.rows[1].cells[0].paragraphs[0].add_run("سعيد راشد اليحيائي")
    table.rows[2].cells[0].paragraphs[0].add_run("مدير مشروع")
    p = tmp_path / "table.docx"
    doc.save(str(p))
    ok = stamp_signature_above_name(
        p, str(_make_sig(tmp_path)), ["سعيد راشد اليحيائي"], size_mm=32.0, boldness=2
    )
    assert ok
    assert b"<wp:anchor" in _document_xml(p)


# ---------------------------------------------------------------------------
# New Report-paper signature-label tests (Task 6)
# ---------------------------------------------------------------------------


def test_sig_label_anchors_below(tmp_path: Path) -> None:
    """التوقيع label paragraph triggers anchor on para[3] (blank below), not
    above the name at para[0]."""
    docx = _make_report_sig_block(tmp_path)
    ok = stamp_signature_above_name(
        docx, str(_make_sig(tmp_path)), ["فلان"], size_mm=32.0, boldness=2
    )
    assert ok
    reloaded = Document(str(docx))
    paras = reloaded.paragraphs
    # Drawing must be in para[3] (blank below the label), NOT para[0] (above name)
    anchored = [i for i, p in enumerate(paras) if _para_has_anchor(p)]
    assert anchored == [3], f"expected drawing at para[3], got {anchored}"


def test_sig_label_beats_name_anchor(tmp_path: Path) -> None:
    """Even though 'فلان' matches para[0], the التوقيع rule wins."""
    docx = _make_report_sig_block(tmp_path)
    ok = stamp_signature_above_name(
        docx, str(_make_sig(tmp_path)), ["فلان"], size_mm=32.0, boldness=2
    )
    assert ok
    reloaded = Document(str(docx))
    paras = reloaded.paragraphs
    # Anchor must NOT be at para[0] or para[-1] (the name gap)
    anchored = [i for i, p in enumerate(paras) if _para_has_anchor(p)]
    assert 0 not in anchored, "name anchor must not win when label is present"
    assert anchored == [3]


def test_date_below_written(tmp_path: Path) -> None:
    """date_below text is written into the blank anchor paragraph, as an RTL run."""
    docx = _make_report_sig_block(tmp_path)
    ok = stamp_signature_above_name(
        docx,
        str(_make_sig(tmp_path)),
        ["فلان"],
        size_mm=32.0,
        boldness=2,
        date_below="27/07/2026",
    )
    assert ok
    reloaded = Document(str(docx))
    # Date text must appear somewhere in the document
    all_text = " ".join(p.text for p in reloaded.paragraphs)
    assert "27/07/2026" in all_text, "date_below text not written to document"
    # The run carrying it must be marked RTL (w:rtl element present in rPr)
    anchor_para = reloaded.paragraphs[3]
    assert anchor_para.text.strip() == "27/07/2026"
    from docx.oxml.ns import qn

    has_rtl = any(run._element.find(f".//{qn('w:rtl')}") is not None for run in anchor_para.runs)
    assert has_rtl, "date run must be marked RTL"


def test_date_below_none_writes_nothing(tmp_path: Path) -> None:
    """Default call (no date_below) writes no date text."""
    docx = _make_report_sig_block(tmp_path)
    stamp_signature_above_name(docx, str(_make_sig(tmp_path)), ["فلان"], size_mm=32.0, boldness=2)
    reloaded = Document(str(docx))
    # para[3] was blank; it must still have no visible text (only anchor XML)
    assert reloaded.paragraphs[3].text.strip() == ""


def test_date_never_clobbers_text(tmp_path: Path) -> None:
    """When the anchor paragraph already has text, date_below is skipped but
    the image is still placed."""
    doc = Document()
    doc.add_paragraph("الاسم : فلان")
    doc.add_paragraph("المسمى الوظيفي : كذا")
    doc.add_paragraph("التوقيـــــع:")
    doc.add_paragraph("نص موجود")  # pre-existing text — must not be overwritten
    p = tmp_path / "nonempty_anchor.docx"
    doc.save(str(p))
    ok = stamp_signature_above_name(
        p,
        str(_make_sig(tmp_path)),
        ["فلان"],
        size_mm=32.0,
        boldness=2,
        date_below="27/07/2026",
    )
    assert ok
    reloaded = Document(str(p))
    # Original text preserved — date NOT written over it
    assert reloaded.paragraphs[3].text.strip() == "نص موجود"
    # Image still placed
    assert _para_has_anchor(reloaded.paragraphs[3])
