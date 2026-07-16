# backend/tests/test_arabic_rtl_merged_cells.py
from docx import Document
from docx.oxml.ns import qn

from app.core.arabic_rtl import html_to_docx


def _table(html):
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(html, p)
    return doc.tables[0]


def _tcPr(table, r, c):
    return table.rows[r]._tr.findall(qn("w:tc"))[c].find(qn("w:tcPr"))


def test_colspan_produces_gridspan():
    t = _table('<table><tr><td colspan="2">H</td></tr><tr><td>A</td><td>B</td></tr></table>')
    gs = _tcPr(t, 0, 0).find(qn("w:gridSpan"))
    assert gs is not None and gs.get(qn("w:val")) == "2"
    assert t.rows[0].cells[0].text.strip() == "H"
    assert t.rows[1].cells[0].text.strip() == "A"
    assert t.rows[1].cells[1].text.strip() == "B"


def test_rowspan_produces_vmerge():
    t = _table('<table><tr><td rowspan="2">S</td><td>A</td></tr><tr><td>B</td></tr></table>')
    vm0 = _tcPr(t, 0, 0).find(qn("w:vMerge"))
    vm1 = _tcPr(t, 1, 0).find(qn("w:vMerge"))
    assert vm0 is not None and vm0.get(qn("w:val")) == "restart"
    assert vm1 is not None and vm1.get(qn("w:val")) in (None, "continue")
    # The cell displaced by the rowspan lands in grid column 1.
    assert t.rows[1].cells[1].text.strip() == "B"


def test_colspan_cell_width_spans_columns():
    t = _table(
        '<table><tr><td colspan="2">H</td><td>X</td></tr>'
        "<tr><td>A</td><td>B</td><td>C</td></tr></table>"
    )
    w_h = int(_tcPr(t, 0, 0).find(qn("w:tcW")).get(qn("w:w")))
    w_a = int(_tcPr(t, 1, 0).find(qn("w:tcW")).get(qn("w:w")))
    w_b = int(_tcPr(t, 1, 1).find(qn("w:tcW")).get(qn("w:w")))
    assert abs(w_h - (w_a + w_b)) <= 2  # rounding tolerance


def test_span_free_tables_unchanged():
    t = _table("<table><tr><td>A</td><td>B</td></tr></table>")
    assert len(t.columns) == 2
    assert _tcPr(t, 0, 0).find(qn("w:gridSpan")) is None


def test_rowspan_and_colspan_same_row():
    t = _table(
        '<table><tr><td colspan="2">H</td><td rowspan="2">R</td></tr>'
        "<tr><td>A</td><td>B</td></tr></table>"
    )
    # H spans grid cols 0-1; R occupies grid col 2 rows 0-1.
    gs = _tcPr(t, 0, 0).find(qn("w:gridSpan"))
    assert gs is not None and gs.get(qn("w:val")) == "2"
    # Row 0 raw tcs after the colspan merge: [H(merged), R] -> R at raw index 1.
    vm = _tcPr(t, 0, 1).find(qn("w:vMerge"))
    assert vm is not None and vm.get(qn("w:val")) == "restart"
    # Row 1: [A, B, continuation] -> continuation at raw index 2.
    vm1 = _tcPr(t, 1, 2).find(qn("w:vMerge"))
    assert vm1 is not None and vm1.get(qn("w:val")) in (None, "continue")
    assert t.rows[0].cells[0].text.strip() == "H"
    assert t.rows[1].cells[0].text.strip() == "A"
    assert t.rows[1].cells[1].text.strip() == "B"


def test_cell_with_both_rowspan_and_colspan():
    t = _table(
        '<table><tr><td colspan="2" rowspan="2">S</td><td>X</td></tr>'
        "<tr><td>Y</td></tr>"
        "<tr><td>A</td><td>B</td><td>C</td></tr></table>"
    )
    tcPr = _tcPr(t, 0, 0)
    gs = tcPr.find(qn("w:gridSpan"))
    vm = tcPr.find(qn("w:vMerge"))
    assert gs is not None and gs.get(qn("w:val")) == "2"
    assert vm is not None and vm.get(qn("w:val")) == "restart"
    assert t.rows[0].cells[0].text.strip() == "S"
    assert t.rows[2].cells[2].text.strip() == "C"
