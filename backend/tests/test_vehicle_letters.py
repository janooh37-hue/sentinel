"""Vehicle fines and accident letter generation contracts.

These tests deliberately exercise the real DOCX rendering and Records filing
pipeline. PDF conversion alone is stubbed so the suite does not depend on Word.
"""

from __future__ import annotations

import re
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

import pytest
from docx import Document as DocxDocument
from sqlalchemy.orm import Session

from app.config import Settings
from app.db.models import (
    Book,
    BookCategory,
    Document,
    Employee,
    User,
    Vehicle,
    VehicleAccident,
    VehicleFine,
    VehicleSite,
)
from app.services import document_service, vehicle_letter_service

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"
FINES_TEMPLATE = TEMPLATES_DIR / "GSSG-VF_300-001_Vehicle_Fines.docx"
ACCIDENT_TEMPLATE = TEMPLATES_DIR / "GSSG-VA_300-001_Vehicle_Accident_Report.docx"


@dataclass(frozen=True)
class _GenerationEnv:
    db: Session
    settings: Settings
    user: User
    vehicle: Vehicle
    employees: tuple[Employee, Employee]
    fines: tuple[VehicleFine, VehicleFine]
    accident: VehicleAccident


def _seed_vehicle_categories(db: Session) -> None:
    categories = (
        ("VF", "Vehicle Fines", "مخالفات المركبات"),
        ("VA", "Vehicle Accidents", "حوادث المركبات"),
    )
    for category_id, name_en, name_ar in categories:
        if db.get(BookCategory, category_id) is None:
            db.add(
                BookCategory(
                    id=category_id,
                    prefix=category_id,
                    name_en=name_en,
                    name_ar=name_ar,
                )
            )
    db.commit()


@pytest.fixture()
def gen_env(db_session, tmp_path, monkeypatch) -> _GenerationEnv:
    """Use temporary output and seed categories omitted by ``create_all``."""
    settings = Settings(data_dir=tmp_path / "data")
    monkeypatch.setattr(document_service, "get_settings", lambda: settings)
    monkeypatch.setattr(document_service, "convert_docx_to_pdf", lambda path: None)
    _seed_vehicle_categories(db_session)

    fine_employee_1 = Employee(
        id="G4101",
        name_en="Salim Test",
        name_ar="سالم الاختباري",
    )
    fine_employee_2 = Employee(
        id="G4102",
        name_en="Maryam Test",
        name_ar="مريم التجريبية",
    )
    submitter = Employee(
        id="G9900",
        name_en="Vehicle Letter Submitter",
        name_ar="مُصدر خطابات المركبات",
    )
    site = VehicleSite(name_ar="موقع الاختبار", name_en="Test Site")
    db_session.add_all([fine_employee_1, fine_employee_2, submitter, site])
    db_session.flush()

    user = User(
        email="vehicle.letters@example.test",
        password_hash="not-used",
        employee_id=submitter.id,
        display_name="Vehicle Letter Tester",
        role="admin",
        status="active",
    )
    vehicle = Vehicle(
        plate_code="14",
        plate_number="58216",
        traffic_code="1180021637",
        type_ar="تويوتا لاند كروزر",
        type_en="Toyota Land Cruiser",
        class_ar="مركبة خفيفة",
        class_en="Light vehicle",
        vin="JT111111111111111",
        site_id=site.id,
        license_start=date(2026, 1, 1),
        license_expiry=date(2026, 12, 31),
    )
    db_session.add_all([user, vehicle])
    db_session.flush()

    fine_1 = VehicleFine(
        vehicle_id=vehicle.id,
        employee_id=fine_employee_1.id,
        date=date(2026, 7, 12),
        time="08:15",
        amount=349,
        black_points=4,
        source="manual",
        location="أبوظبي",
        created_by_user_id=user.id,
    )
    fine_2 = VehicleFine(
        vehicle_id=vehicle.id,
        employee_id=fine_employee_2.id,
        date=date(2026, 7, 15),
        time="16:40",
        amount=725,
        black_points=6,
        source="manual",
        location="العين",
        created_by_user_id=user.id,
    )
    accident = VehicleAccident(
        vehicle_id=vehicle.id,
        employee_id=fine_employee_1.id,
        date=date(2026, 8, 17),
        time="14:35",
        location_ar="بوابة الموقع الشرقية",
        location_en="East site gate",
        description_ar="اصطدام بالحاجز أثناء الرجوع إلى الخلف",
        description_en="Collision with the barrier while reversing",
        police_ref="AUH-ACC-260817",
        damage_cost=2600,
        status="open",
        photo_file_ids=[],
    )
    db_session.add_all([fine_1, fine_2, accident])
    db_session.commit()

    return _GenerationEnv(
        db=db_session,
        settings=settings,
        user=user,
        vehicle=vehicle,
        employees=(fine_employee_1, fine_employee_2),
        fines=(fine_1, fine_2),
        accident=accident,
    )


def _container_text(container: Any) -> list[str]:
    chunks = [paragraph.text for paragraph in container.paragraphs]
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(_container_text(cell))
    return chunks


def _all_word_xml_text(path: Path) -> str:
    """Extract text nodes from every Word story, including text boxes."""
    chunks: list[str] = []
    with zipfile.ZipFile(path) as archive:
        for name in sorted(archive.namelist()):
            if not (name.startswith("word/") and name.endswith(".xml")):
                continue
            root = ElementTree.fromstring(archive.read(name))
            chunks.append(
                "".join(
                    node.text or "" for node in root.iter() if node.tag.rsplit("}", 1)[-1] == "t"
                )
            )
    return "\n".join(chunks)


def _all_docx_text(path: Path) -> str:
    """Extract body, table, header, and footer text from every section."""
    document = DocxDocument(str(path))
    chunks = _container_text(document)
    for section in document.sections:
        stories = (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for story in stories:
            chunks.extend(_container_text(story))
    chunks.append(_all_word_xml_text(path))
    return "\n".join(chunks)


def _rendered_docx(env: _GenerationEnv, document_id: int) -> tuple[Document, Path]:
    document = env.db.get(Document, document_id)
    assert document is not None
    assert document.docx_path is not None
    path = Path(document.docx_path)
    if not path.is_absolute():
        path = env.settings.data_dir / path
    assert path.is_file()
    return document, path


def _table_text(table: Any) -> str:
    return "\n".join(
        text for row in table.rows for cell in row.cells for text in _container_text(cell)
    )


def _row_text(row: Any) -> str:
    return "\n".join(text for cell in row.cells for text in _container_text(cell))


def test_two_fines_letter_renders_docx_and_files_vf_book(gen_env: _GenerationEnv) -> None:
    result = vehicle_letter_service.generate_fines_letter(
        gen_env.db,
        gen_env.vehicle.id,
        fine_ids=[fine.id for fine in reversed(gen_env.fines)],
        hide_names=False,
        user=gen_env.user,
    )

    assert re.fullmatch(r"VF-\d{4}", result.ref_number)
    assert result.pdf_available is False
    book = gen_env.db.get(Book, result.book_id)
    assert book is not None
    assert book.category_id == "VF"
    assert book.ref_number == result.ref_number

    document, path = _rendered_docx(gen_env, result.document_id)
    assert document.template_id == "Vehicle Fines"
    text = _all_docx_text(path)
    assert "14 / 58216" in text
    for employee in gen_env.employees:
        assert employee.name_ar is not None
        assert employee.name_ar in text
        assert employee.id in text
    assert "349 درهم" in text
    assert "725 درهم" in text
    assert result.ref_number in text


def test_hidden_names_fines_letter_redacts_names_and_g_numbers(
    gen_env: _GenerationEnv,
) -> None:
    result = vehicle_letter_service.generate_fines_letter(
        gen_env.db,
        gen_env.vehicle.id,
        fine_ids=[fine.id for fine in gen_env.fines],
        hide_names=True,
        user=gen_env.user,
    )

    book = gen_env.db.get(Book, result.book_id)
    assert book is not None and book.category_id == "VF"
    _, path = _rendered_docx(gen_env, result.document_id)
    text = _all_docx_text(path)
    for employee in gen_env.employees:
        assert employee.name_ar is not None
        assert employee.name_ar not in text
        assert employee.name_en not in text
        assert employee.id not in text
    assert "نسخة بدون أسماء" in text


def test_accident_letter_renders_docx_files_va_book_and_links_accident(
    gen_env: _GenerationEnv,
) -> None:
    result = vehicle_letter_service.generate_accident_letter(
        gen_env.db,
        gen_env.vehicle.id,
        gen_env.accident.id,
        user=gen_env.user,
    )

    assert re.fullmatch(r"VA-\d{4}", result.ref_number)
    book = gen_env.db.get(Book, result.book_id)
    assert book is not None
    assert book.category_id == "VA"
    assert book.ref_number == result.ref_number
    gen_env.db.refresh(gen_env.accident)
    assert gen_env.accident.letter_book_id == result.book_id

    document, path = _rendered_docx(gen_env, result.document_id)
    assert document.template_id == "Vehicle Accident Report"
    text = _all_docx_text(path)
    expected = (
        result.ref_number,
        "14 / 58216",
        gen_env.vehicle.type_ar,
        gen_env.vehicle.vin,
        gen_env.vehicle.site.name_ar,
        "17/08/2026 14:35",
        gen_env.employees[0].name_ar,
        gen_env.accident.location_ar,
        gen_env.accident.police_ref,
        "2600 درهم",
        "مفتوح",
        gen_env.accident.description_ar,
    )
    for value in expected:
        assert value is not None
        assert value in text


def test_generating_accident_letter_twice_reuses_linked_va_book(
    gen_env: _GenerationEnv,
) -> None:
    first = vehicle_letter_service.generate_accident_letter(
        gen_env.db,
        gen_env.vehicle.id,
        gen_env.accident.id,
        user=gen_env.user,
    )
    second = vehicle_letter_service.generate_accident_letter(
        gen_env.db,
        gen_env.vehicle.id,
        gen_env.accident.id,
        user=gen_env.user,
    )

    assert second.book_id == first.book_id
    assert second.ref_number == first.ref_number
    va_books = gen_env.db.query(Book).filter_by(category_id="VA").all()
    assert [book.id for book in va_books] == [first.book_id]
    gen_env.db.refresh(gen_env.accident)
    assert gen_env.accident.letter_book_id == first.book_id


@pytest.mark.parametrize("template_path", [FINES_TEMPLATE, ACCIDENT_TEMPLATE])
def test_vehicle_letter_templates_contain_no_literal_g_numbers(
    template_path: Path,
) -> None:
    text = _all_docx_text(template_path)
    assert re.search(r"G\d{4}", text) is None


def test_fines_template_has_only_one_templated_data_row() -> None:
    """Only header, loop controls, and one data row may survive the source."""
    document = DocxDocument(str(FINES_TEMPLATE))
    matching_tables = [
        table for table in document.tables if "{%tr for f in fines %}" in _table_text(table)
    ]
    assert len(matching_tables) == 1
    fines_table = matching_tables[0]
    assert len(fines_table.rows) == 4

    row_texts = [_row_text(row).strip() for row in fines_table.rows]
    assert row_texts[1] == "{%tr for f in fines %}"
    assert row_texts[3] == "{%tr endfor %}"

    data_rows = [row_text for row_text in row_texts if "{{ f." in row_text]
    assert data_rows == [row_texts[2]]
    data_tokens = {
        "{{ f.points }}",
        "{{ f.amount }}",
        "{{ f.date }}",
        "{{ f.g_number }}",
        "{{ f.employee_name }}",
        "{{ f.seq }}",
    }
    missing = {token for token in data_tokens if token not in row_texts[2]}
    assert not missing, f"missing fines data-row tokens: {sorted(missing)}"


def test_fines_template_contains_expected_jinja_tokens() -> None:
    text = _all_word_xml_text(FINES_TEMPLATE)
    expected_tokens = {
        "{{ ref }}",
        "{{ today }}",
        "{{ hijri_date }}",
        "{{ submitter_g }}",
        "{{ plate }}",
        "{% if hide_names %}",
        "{% endif %}",
        "{%tr for f in fines %}",
        "{%tr endfor %}",
        "{{ f.points }}",
        "{{ f.amount }}",
        "{{ f.date }}",
        "{{ f.g_number }}",
        "{{ f.employee_name }}",
        "{{ f.seq }}",
    }
    missing = {token for token in expected_tokens if token not in text}
    assert not missing, f"missing fines template tokens: {sorted(missing)}"


def test_accident_template_contains_expected_jinja_tokens() -> None:
    text = _all_word_xml_text(ACCIDENT_TEMPLATE)
    expected_tokens = {
        "{{ ref }}",
        "{{ today }}",
        "{{ hijri_date }}",
        "{{ submitter_g }}",
        "{{ plate }}",
        "{{ vehicle_type }}",
        "{{ vin }}",
        "{{ site }}",
        "{{ date_time }}",
        "{{ employee }}",
        "{{ location }}",
        "{{ police_ref }}",
        "{{ damage_cost }}",
        "{{ status }}",
        "{{ description }}",
    }
    missing = {token for token in expected_tokens if token not in text}
    assert not missing, f"missing accident template tokens: {sorted(missing)}"
