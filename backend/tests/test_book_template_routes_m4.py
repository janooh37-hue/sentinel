"""TDD M4-7: GET /books/word-templates/{name}/table + DELETE /books/word-templates/{name}.

RED → implement → GREEN.
"""

from __future__ import annotations

import secrets
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

from app.api.deps import get_current_user
from app.core.book_template_retokenize import retokenize_general_book
from app.db import session as session_mod
from app.db.models import Base, User
from app.db.session import attach_sqlite_pragmas, get_db
from app.main import create_app
from app.services import book_template_service, perm_service

# ---------------------------------------------------------------------------
# Fixtures (mirrors test_word_book_routes.py pattern)
# ---------------------------------------------------------------------------


@pytest.fixture()
def api_db(monkeypatch, tmp_path) -> Session:
    db_file = tmp_path / "routes_m4.db"
    eng = create_engine(
        f"sqlite:///{db_file}",
        future=True,
        connect_args={"check_same_thread": False},
    )
    attach_sqlite_pragmas(eng, wal=False)
    Base.metadata.create_all(eng)
    TestSession = sessionmaker(bind=eng, autoflush=False, expire_on_commit=False, future=True)
    monkeypatch.setattr(session_mod, "engine", eng)
    monkeypatch.setattr(session_mod, "SessionLocal", TestSession)
    db = TestSession()
    perm_service.seed_role_defaults(db)
    yield db
    db.close()


def _make_user(db: Session, role: str = "operator") -> User:
    u = User(
        email=f"{secrets.token_hex(4)}@test.ae",
        password_hash="x",
        role=role,
        status="active",
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


def _client(db: Session, user: User, monkeypatch, tmp_path: Path) -> TestClient:
    """Build a test client wired to db + user."""
    from app.config import Settings
    from app.services import word_book_service

    settings = Settings(
        data_dir=tmp_path / "data",
        templates_dir=tmp_path / "templates",
    )
    (tmp_path / "templates").mkdir(parents=True, exist_ok=True)

    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)

    app = create_app()
    app.dependency_overrides[get_db] = lambda: db
    app.dependency_overrides[get_current_user] = lambda: user
    return TestClient(app, raise_server_exceptions=True)


def _admin_client(api_db, monkeypatch, tmp_path):
    """books.manage-capable client (admin role)."""
    user = _make_user(api_db, role="admin")
    return _client(api_db, user, monkeypatch, tmp_path), user


def _plain_client(api_db, monkeypatch, tmp_path):
    """No books.manage (operator role)."""
    user = _make_user(api_db, role="operator")
    return _client(api_db, user, monkeypatch, tmp_path), user


# ---------------------------------------------------------------------------
# Template helpers
# ---------------------------------------------------------------------------


def _plain_docx(path: Path) -> None:
    """Retokenized General Book docx with no table."""
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: خطاب نصّي بدون جدول")
    doc.add_paragraph("نص")
    doc.save(str(path))
    retokenize_general_book(path)


def _table_docx(path: Path, headers: list[str]) -> None:
    """Retokenized General Book docx with a table."""
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: كتاب جدول بيانات للاختبار")
    t = doc.add_table(rows=2, cols=len(headers))
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for i in range(len(headers)):
        t.cell(1, i).text = f"بيانات {i}"
    doc.save(str(path))
    retokenize_general_book(path)


# ---------------------------------------------------------------------------
# GET /books/word-templates/{name}/table
# ---------------------------------------------------------------------------


def test_get_table_plain_template(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)
    _plain_docx(tpl_lib / "plain.docx")

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates/plain.docx/table")
    assert r.status_code == 200
    body = r.json()
    assert body["has_table"] is False
    assert body["columns"] == []


def test_get_table_table_template(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)
    headers = ["الاسم", "الرقم الوظيفي", "الجهة"]
    _table_docx(tpl_lib / "table.docx", headers)

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates/table.docx/table")
    assert r.status_code == 200
    body = r.json()
    assert body["has_table"] is True
    assert body["columns"] == headers


def test_get_table_missing_template_404(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates/ghost.docx/table")
    assert r.status_code == 404


def test_get_table_requires_books_manage(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _plain_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates/anything.docx/table")
    assert r.status_code == 403


# ---------------------------------------------------------------------------
# DELETE /books/word-templates/{name}
# ---------------------------------------------------------------------------


def test_delete_existing_template_204(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)
    _plain_docx(tpl_lib / "del.docx")
    assert (tpl_lib / "del.docx").exists()

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.delete("/api/v1/books/word-templates/del.docx")
    assert r.status_code == 204
    assert not (tpl_lib / "del.docx").exists()


def test_delete_missing_template_404(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.delete("/api/v1/books/word-templates/ghost.docx")
    assert r.status_code == 404


def test_delete_requires_books_manage(api_db, monkeypatch, tmp_path):
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _plain_client(api_db, monkeypatch, tmp_path)
    r = c.delete("/api/v1/books/word-templates/anything.docx")
    assert r.status_code == 403
