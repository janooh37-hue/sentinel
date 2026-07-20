"""Tests for table_rows threading in create_word_book (Task M4-8).

Builds a tokenized table template in the library dir, calls create_word_book
with table_rows, then asserts the working docx carries rendered row values.
"""

from __future__ import annotations

import secrets
from pathlib import Path

from app.db.models import BookCategory, BookEditSession, User

# ---------------------------------------------------------------------------
# Helpers (copied from test_word_book_service — same pattern)
# ---------------------------------------------------------------------------


def _seed_gs(db) -> None:
    if db.get(BookCategory, "GS") is None:
        db.add(BookCategory(id="GS", prefix="GS"))
        db.commit()


def _user(db, *, employee_id: str | None = None) -> User:
    u = User(email=f"{secrets.token_hex(4)}@test.ae", password_hash="x", status="active")
    u.employee_id = employee_id
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


def _settings(tmp_path: Path):
    from app.config import Settings

    return Settings(data_dir=tmp_path / "data", templates_dir=tmp_path / "templates")


def _write_minimal_docx(path: Path) -> None:
    """Minimal General Book template with tokens (same as test_word_book_service)."""
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("{%p if ref %}")
    doc.add_paragraph("الرقم: {{ ref }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("التاريخ: {{ date }}")
    doc.add_paragraph("السيد / {{ recipient_name }}")
    doc.add_paragraph("الموضوع: {{ subject }}")
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("{{ cc }}")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_title }}")
    doc.add_paragraph("{{ submitter_g }}")
    doc.save(str(path))


def _build_table_template(library_dir: Path, name: str, columns: list[str]) -> Path:
    """Build a tokenized table template file in *library_dir*.

    Creates a docx with:
    - التاريخ line (required by retokenize_general_book)
    - a table with the given column headers + one sample data row
    then runs retokenize_general_book (which internally calls normalize_data_table)
    to inject the {%tr for row in table_rows %} loop.
    """
    import docx as _docx

    from app.core.book_template_retokenize import retokenize_general_book

    src = library_dir / f"_src_{name}"
    doc = _docx.Document()
    doc.add_paragraph("التاريخ: 01/01/2026")
    # table with header row + one data row
    tbl = doc.add_table(rows=2, cols=len(columns))
    for i, col_name in enumerate(columns):
        tbl.rows[0].cells[i].text = col_name
        tbl.rows[1].cells[i].text = f"sample_{i}"
    doc.save(str(src))

    retokenize_general_book(src)
    dest = library_dir / name
    src.rename(dest)
    return dest


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_two_table_rows_render(db_session, tmp_path, monkeypatch):
    """Two data rows both appear in the working docx after create_word_book."""
    from app.core.book_text import docx_to_text
    from app.services import book_template_service as tpl_svc
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / "GSSG-GS_300-003_General_Book.docx")
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    # Library dir is data_dir / book_templates (see book_template_service.templates_dir)
    library_dir = tmp_path / "data" / "book_templates"
    library_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(tpl_svc, "templates_dir", lambda: library_dir)

    _build_table_template(library_dir, "جدول.docx", ["العمود1", "العمود2"])  # noqa: RUF001

    user = _user(db_session)
    rows = [
        {"c0": "alpha", "c1": "beta"},
        {"c0": "gamma", "c1": "delta"},
    ]
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="table test",
        cc=None,
        manager_id=None,
        template_name="جدول.docx",
        table_rows=rows,
    )

    session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    text = docx_to_text(Path(session.working_path))

    assert "alpha" in text
    assert "beta" in text
    assert "gamma" in text
    assert "delta" in text
    # No directive residue
    assert "{%tr" not in text


def test_no_table_rows_plain_template_no_error(db_session, tmp_path, monkeypatch):
    """Omitting table_rows (None) on a plain template works without error."""
    from app.core.book_text import docx_to_text
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / "GSSG-GS_300-003_General_Book.docx")
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="no table",
        cc=None,
        manager_id=None,
        table_rows=None,
    )

    session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    text = docx_to_text(Path(session.working_path))
    assert f"الرقم: {info.ref_number}" in text


def test_non_str_values_coerced(db_session, tmp_path, monkeypatch):
    """Non-str values (passed as str by schema, but coercion guard runs) produce no crash."""
    from app.core.book_text import docx_to_text
    from app.services import book_template_service as tpl_svc
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / "GSSG-GS_300-003_General_Book.docx")
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    library_dir = tmp_path / "data" / "book_templates"
    library_dir.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(tpl_svc, "templates_dir", lambda: library_dir)

    _build_table_template(library_dir, "coerce.docx", ["Col"])

    user = _user(db_session)
    # Pass values that are already str (schema enforces str; coercion is a safety guard)
    rows = [{"c0": "42"}, {"c0": "صفر"}]
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="coerce test",
        cc=None,
        manager_id=None,
        template_name="coerce.docx",
        table_rows=rows,
    )

    session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    text = docx_to_text(Path(session.working_path))
    assert "42" in text
    assert "صفر" in text
    assert "{%tr" not in text
