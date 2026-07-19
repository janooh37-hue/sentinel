"""Signing a Word-authored General Book must keep the authored body."""

import zipfile
from pathlib import Path

import pytest
from docx import Document as DocxFile
from sqlalchemy.orm import Session

from app.core.book_text import docx_to_text
from app.db.models import Book, BookCategory, BookVersion, Document
from app.services import document_service

BODY_LINE = "نرجو الموافقة على أعمال الصيانة العاجلة"


@pytest.fixture
def word_version(db_session: Session, tmp_path: Path) -> BookVersion:
    """A finished word-authored book: Document docx on disk, version.fields == {}."""
    docx_path = tmp_path / "1-11-GSSG-9.docx"
    d = DocxFile()
    d.add_paragraph(BODY_LINE)
    d.add_paragraph("")
    d.add_paragraph("سعيد راشد اليحيائي")
    d.save(str(docx_path))

    if db_session.get(BookCategory, "GS") is None:
        db_session.add(BookCategory(id="GS", prefix="GS"))
        db_session.flush()
    book = Book(category_id="GS", ref_number="1/11/GSSG/9", subject="اختبار التوقيع")
    db_session.add(book)
    db_session.flush()
    doc = Document(
        template_id="General Book",
        ref_number=book.ref_number,
        docx_path=str(docx_path),
        submission_id="t-sign",
        role="primary",
    )
    db_session.add(doc)
    db_session.flush()
    version = BookVersion(
        book_id=book.id,
        version_no=1,
        trigger="initial",
        status="none",
        template_id="General Book",
        fields={},
        document_id=doc.id,
    )
    db_session.add(version)
    db_session.commit()
    return version


def _sig(tmp_path: Path) -> str:
    from PIL import Image

    p = tmp_path / "sig.png"
    Image.new("RGBA", (60, 30), (0, 0, 200, 255)).save(p)
    return str(p)


def test_signed_artifact_keeps_word_body(
    db_session: Session,
    tmp_path: Path,
    word_version: BookVersion,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    # PDF conversion is environment-dependent — force the docx fallback path.
    monkeypatch.setattr(document_service, "convert_docx_to_pdf", lambda p: None)
    rel = document_service.render_signed_pdf(
        db_session, version=word_version, signer_signature_path=_sig(tmp_path)
    )
    from app.config import get_settings

    signed = Path(rel)
    if not signed.is_absolute():
        signed = get_settings().data_dir / signed
    assert signed.suffix == ".docx"  # conversion stubbed out
    text = docx_to_text(signed)
    assert BODY_LINE in text  # the authored body SURVIVED signing
    # and the signature image landed (anchored drawing present)
    with zipfile.ZipFile(signed) as z:
        assert b"<wp:anchor" in z.read("word/document.xml")
    signed.unlink()  # keep the shared output dir clean


def test_sign_raises_when_authored_docx_missing(
    db_session: Session,
    tmp_path: Path,
    word_version: BookVersion,
) -> None:
    """fields == {} with the docx gone must FAIL loudly — falling through to the
    template re-render would reproduce the blank signed paper (review M2)."""
    from app.api.errors import AppError

    doc = db_session.get(Document, word_version.document_id)
    assert doc is not None and doc.docx_path
    Path(doc.docx_path).unlink()
    with pytest.raises(AppError) as ei:
        document_service.render_signed_pdf(
            db_session, version=word_version, signer_signature_path=_sig(tmp_path)
        )
    assert ei.value.code == "SOURCE_DOCX_MISSING"


def test_sign_raises_when_stamp_fails(
    db_session: Session,
    tmp_path: Path,
    word_version: BookVersion,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A 'signed' artifact with no visible signature is the defect class this
    branch fixes — a failed stamp must abort the signing (review M1)."""
    import app.core.docx_engine as docx_engine_mod
    from app.api.errors import AppError

    monkeypatch.setattr(docx_engine_mod, "stamp_signature_above_name", lambda *a, **k: False)
    with pytest.raises(AppError) as ei:
        document_service.render_signed_pdf(
            db_session, version=word_version, signer_signature_path=_sig(tmp_path)
        )
    assert ei.value.code == "SIGNATURE_STAMP_FAILED"


def test_sign_falls_back_to_submitter_signature(db_session: Session, tmp_path: Path) -> None:
    from app.db.models import Employee, Submitter, User
    from app.services.book_service import _resolve_signer_signature

    sig = tmp_path / "emp-sig.png"
    from PIL import Image

    Image.new("RGBA", (40, 20), (0, 0, 0, 255)).save(sig)

    db_session.add(Employee(id="G7001", name_en="Signer Emp"))
    db_session.flush()
    user = User(
        email="signer@test.ae",
        password_hash="x",
        role="manager",
        status="active",
        employee_id="G7001",
        signature_path=None,
    )
    db_session.add(user)
    db_session.add(Submitter(employee_id="G7001", name="Signer Emp", stored_sig_path=str(sig)))
    db_session.commit()

    resolved = _resolve_signer_signature(db_session, user)
    assert resolved is not None and resolved.name == "emp-sig.png"


def test_sign_prefers_own_signature(db_session: Session, tmp_path: Path) -> None:
    from app.db.models import User
    from app.services.book_service import _resolve_signer_signature

    own = tmp_path / "own.png"
    from PIL import Image

    Image.new("RGBA", (40, 20), (0, 0, 0, 255)).save(own)
    user = User(
        email="own@test.ae",
        password_hash="x",
        role="manager",
        status="active",
        signature_path=str(own),
    )
    db_session.add(user)
    db_session.commit()
    resolved = _resolve_signer_signature(db_session, user)
    assert resolved is not None and resolved.name == "own.png"


def test_rich_versions_still_rerender(
    db_session: Session,
    tmp_path: Path,
    word_version: BookVersion,
    monkeypatch: pytest.MonkeyPatch,
) -> None:
    """A version WITH fields keeps the existing template re-render path."""
    word_version.fields = {"subject": "موضوع", "body": "نص"}
    db_session.commit()
    calls: list[str] = []

    def fake_fill(self: object, tid: str, data: dict[str, object], out: Path) -> Path:
        calls.append(tid)
        DocxFile().save(str(out))
        return Path(out)

    monkeypatch.setattr(document_service.DocxEngine, "fill", fake_fill)
    monkeypatch.setattr(document_service, "convert_docx_to_pdf", lambda p: None)
    monkeypatch.setattr("app.core.docx_engine._postprocess_general_book_footer", lambda p: None)
    monkeypatch.setattr(
        document_service.DocxEngine, "stamp_aztec_code", staticmethod(lambda *a, **k: True)
    )
    rel = document_service.render_signed_pdf(
        db_session, version=word_version, signer_signature_path=_sig(tmp_path)
    )
    assert calls == ["General Book"]
    from app.config import get_settings

    leftover = Path(rel)
    if not leftover.is_absolute():
        leftover = get_settings().data_dir / leftover
    leftover.unlink(missing_ok=True)
