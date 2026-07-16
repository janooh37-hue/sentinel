# backend/tests/test_arabic_rtl_comments.py
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.core.arabic_rtl import html_to_docx


def _all_text(doc: Any) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


def _page_break_count(doc: Any) -> int:
    brs = doc.element.body.findall(".//" + qn("w:br"))
    return sum(1 for b in brs if b.get(qn("w:type")) == "page")


def test_comment_text_never_rendered() -> None:
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx("<p>a</p><!--[if !supportLists]-->junk<!--[endif]--><p>b</p>", p)
    assert "supportLists" not in _all_text(doc)
    assert "endif" not in _all_text(doc)


def test_pagebreak_comment_at_top_level_emits_page_break() -> None:
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx("<p>one</p><!-- pagebreak --><p>two</p>", p)
    assert _page_break_count(doc) == 1
    assert "pagebreak" not in _all_text(doc)


def test_pagebreak_comment_inside_block_emits_page_break() -> None:
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx("<p>one<!-- pagebreak -->two</p>", p)
    assert _page_break_count(doc) == 1


def test_mce_pagebreak_div_still_works() -> None:
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx('<p>one</p><div class="mce-pagebreak"></div><p>two</p>', p)
    assert _page_break_count(doc) == 1
