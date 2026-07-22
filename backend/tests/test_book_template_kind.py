"""TDD: M4-9 Part B — kind tagging on template list, Part A — build_base_templates script."""

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document as DocxDocument

# ---------------------------------------------------------------------------
# Helpers to create a minimal retokenizable docx
# ---------------------------------------------------------------------------


def _write_retokenizable_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_paragraph("الرقم: 1/2026")
    doc.add_paragraph("التاريخ: 01-01-2026")
    doc.add_paragraph("السيد / اختبار")
    doc.add_paragraph("الموضوع: موضوع تجريبي للاختبار الوظيفي")
    doc.add_paragraph("نص الكتاب العام هنا للاختبار الوظيفي.")
    footer = doc.sections[0].footer
    run = footer.paragraphs[0].add_run("G-0000")
    from docx.shared import Pt

    run.font.size = Pt(9)
    doc.save(str(path))


# ---------------------------------------------------------------------------
# Part B tests: kind tagging via list_templates()
# ---------------------------------------------------------------------------


def test_list_templates_kind_base_and_custom(tmp_path, monkeypatch):
    """Only base_table.docx is built in; no-table uses the normal document."""
    from app.services import book_template_service

    # Point templates_dir() at tmp_path
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tmp_path)

    # Drop three files: two base names + one custom
    (tmp_path / "base_text.docx").write_bytes(b"x")
    (tmp_path / "base_table.docx").write_bytes(b"x")
    (tmp_path / "custom_foo.docx").write_bytes(b"x")

    items = book_template_service.list_templates()
    by_name = {i.name: i.kind for i in items}

    assert by_name["base_text.docx"] == "custom"
    assert by_name["base_table.docx"] == "base"
    assert by_name["custom_foo.docx"] == "custom"


def test_list_templates_kind_default_custom(tmp_path, monkeypatch):
    """Any template that is NOT in _BASE_TEMPLATE_NAMES gets kind='custom'."""
    from app.services import book_template_service

    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tmp_path)
    (tmp_path / "my_template.docx").write_bytes(b"x")

    items = book_template_service.list_templates()
    assert items[0].kind == "custom"


# ---------------------------------------------------------------------------
# Part B tests: kind flows through the API response
# ---------------------------------------------------------------------------


def test_list_word_templates_api_kind(tmp_path, monkeypatch):
    """GET /books/word-templates propagates kind to the JSON response."""
    import secrets

    from fastapi.testclient import TestClient
    from sqlalchemy import create_engine
    from sqlalchemy.orm import sessionmaker

    from app.api.deps import get_current_user
    from app.db import session as session_mod
    from app.db.models import Base, User
    from app.db.session import attach_sqlite_pragmas, get_db
    from app.main import create_app
    from app.services import book_template_service, perm_service

    # Minimal in-memory DB
    eng = create_engine("sqlite://", future=True)
    attach_sqlite_pragmas(eng, wal=False)
    Base.metadata.create_all(eng)
    TestSession = sessionmaker(bind=eng, autoflush=False, expire_on_commit=False, future=True)
    monkeypatch.setattr(session_mod, "engine", eng)
    monkeypatch.setattr(session_mod, "SessionLocal", TestSession)
    db = TestSession()
    perm_service.seed_role_defaults(db)

    user = User(
        email=f"{secrets.token_hex(4)}@test.ae",
        password_hash="x",
        role="admin",
        status="active",
    )
    db.add(user)
    db.commit()
    db.refresh(user)

    # Point templates_dir() at tmp_path
    tpl_dir = tmp_path / "book_templates"
    tpl_dir.mkdir()
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_dir)

    (tpl_dir / "base_text.docx").write_bytes(b"x")
    (tpl_dir / "custom_foo.docx").write_bytes(b"x")

    app = create_app()
    app.dependency_overrides[get_db] = lambda: db
    app.dependency_overrides[get_current_user] = lambda: user
    client = TestClient(app, raise_server_exceptions=True)

    resp = client.get("/api/v1/books/word-templates")
    assert resp.status_code == 200, resp.text
    by_name = {item["name"]: item["kind"] for item in resp.json()}
    assert by_name["base_text.docx"] == "custom"
    assert by_name["custom_foo.docx"] == "custom"


# ---------------------------------------------------------------------------
# Part A tests: build_base_templates script
# ---------------------------------------------------------------------------


def _import_build_script():
    """Import build_base_templates, adding backend/ to sys.path first."""
    backend = Path(__file__).resolve().parents[1]
    if str(backend) not in sys.path:
        sys.path.insert(0, str(backend))
    import importlib

    return importlib.import_module("scripts.build_base_templates")


def test_build_base_templates_creates_files(tmp_path):
    """build_templates(output_dir) produces only base_table.docx."""
    mod = _import_build_script()
    mod.build_templates(tmp_path)
    assert not (tmp_path / "base_text.docx").exists()
    assert (tmp_path / "base_table.docx").is_file()


def test_build_base_templates_validate(tmp_path):
    """The produced table template passes validation."""
    mod = _import_build_script()
    mod.build_templates(tmp_path)

    from app.core.book_template_retokenize import validate_book_template

    validate_book_template(tmp_path / "base_table.docx")


def test_build_base_templates_check_flag(tmp_path):
    """check_templates(dir) returns True when the table template validates."""
    mod = _import_build_script()
    mod.build_templates(tmp_path)
    assert mod.check_templates(tmp_path) is True


def test_build_base_templates_check_flag_missing(tmp_path):
    """check_templates(dir) returns False when files are absent."""
    mod = _import_build_script()
    assert mod.check_templates(tmp_path) is False
