# backend/tests/test_arabic_rtl_whitespace.py
from docx import Document

from app.core.arabic_rtl import html_to_docx


def _render(html):
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(html, p)
    return doc


def test_newlines_inside_text_collapse_to_single_space():
    doc = _render("<p>محمد\n  مشرف حسين</p>")
    assert doc.paragraphs[0].text == "محمد مشرف حسين"


def test_intertag_whitespace_not_rendered_in_cells():
    doc = _render("<table><tr>\n  <td>\n  <p>A</p>\n  </td>\n</tr></table>")
    cell = doc.tables[0].rows[0].cells[0]
    for q in cell.paragraphs:
        assert "\n" not in q.text
        assert q.text.strip(" ") in ("A", "")


def test_nbsp_is_preserved():
    doc = _render("<p>&nbsp;</p>")
    assert doc.paragraphs[0].text == " "  # noqa: RUF001


def test_pre_whitespace_preserved():
    doc = _render("<pre>a\n  b</pre>")
    assert "a\n  b" in doc.paragraphs[0].text


def test_single_spaces_between_inline_tags_kept():
    doc = _render("<p><b>A</b> <i>B</i></p>")
    assert doc.paragraphs[0].text == "A B"
