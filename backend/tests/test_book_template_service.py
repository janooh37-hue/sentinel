"""Library naming rules, listing, and the save-as-template flow."""

from __future__ import annotations

import unicodedata
from datetime import UTC, datetime
from pathlib import Path

import pytest

from app.api.errors import AppError
from app.services import book_template_service as svc

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------

_GENERAL_BOOK = "GSSG-GS_300-003_General_Book.docx"


def _settings(tmp_path: Path):
    from app.config import Settings

    return Settings(data_dir=tmp_path / "data", templates_dir=tmp_path / "templates")


def _write_minimal_docx(path: Path) -> None:
    """Write a minimal docx with tokens matching the General Book template."""
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("{%p if ref %}")
    doc.add_paragraph("الرقم: {{ ref }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("التاريخ: {{ date }}")
    doc.add_paragraph("السيد / {{ recipient_name }}")
    doc.add_paragraph("الموضوع: {{ subject }}")
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("{{ cc }}")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_title }}")
    doc.add_paragraph("{{ submitter_g }}")
    doc.save(str(path))


@pytest.fixture()
def finished_word_book(db_session, tmp_path, monkeypatch):
    """Create a finished General Book via word_book_service.

    Returns the Book row with a committed BookVersion + Document pointing
    at a real docx on disk.
    """
    from app.db.models import BookCategory, BookEditSession, User
    from app.services import word_book_service

    # Seed the GS category
    if db_session.get(BookCategory, "GS") is None:
        db_session.add(BookCategory(id="GS", prefix="GS"))
        db_session.commit()

    # Minimal template
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)

    # Monkeypatch both services to use the same tmp settings
    settings = _settings(tmp_path)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)

    # Also patch templates_dir to tmp_path for the template-service calls
    monkeypatch.setattr(svc, "templates_dir", lambda: tmp_path / "tpl_lib")
    (tmp_path / "tpl_lib").mkdir(parents=True, exist_ok=True)

    # Create user (employee_id=None — FK to employees table; submitter_g will be None)
    user = User(email="test_tpl@test.ae", password_hash="x", status="active")
    db_session.add(user)
    db_session.commit()
    db_session.refresh(user)

    # Create the word book session
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="Template test book",
        cc=None,
        manager_id=None,
    )

    # Simulate Word having saved: set last_put_at on the session
    edit_session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    edit_session.last_put_at = datetime.now(UTC).replace(tzinfo=None)
    db_session.commit()

    # Finish the session → BookVersion + Document
    book = word_book_service.finish_word_session(db_session, user=user, book_id=info.book_id)
    return book


# ---------------------------------------------------------------------------
# Tests: naming rules
# ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    "bad",
    ["", "..", "../evil", "a/b", "a\\b", "CON", "con.docx", "NUL", "COM3", "name."],
)
def test_bad_names_rejected(bad):
    with pytest.raises(AppError) as ei:
        svc.safe_template_name(bad)
    assert ei.value.code == "TEMPLATE_BAD_NAME"


def test_name_gets_docx_extension_and_nfc():
    decomposed = unicodedata.normalize("NFD", "قالب")
    assert svc.safe_template_name(decomposed) == "قالب.docx"
    assert svc.safe_template_name("التصاريح الأمنية") == "التصاريح الأمنية.docx"
    assert svc.safe_template_name("جاهز.docx") == "جاهز.docx"


# ---------------------------------------------------------------------------
# Tests: listing
# ---------------------------------------------------------------------------


def test_list_filters_to_docx(tmp_path, monkeypatch):
    monkeypatch.setattr(svc, "templates_dir", lambda: tmp_path)
    (tmp_path / "صيانة.docx").write_bytes(b"x")
    (tmp_path / "stray.tmp").write_bytes(b"x")
    names = [t.name for t in svc.list_templates()]
    assert names == ["صيانة.docx"]


# ---------------------------------------------------------------------------
# Tests: save flow
# ---------------------------------------------------------------------------


def test_save_book_as_template_roundtrip(db_session, finished_word_book):
    """finished_word_book: fixture creating a finished General Book via
    word_book_service (reuse/extract the pattern from test_word_book_service).
    Returns the Book row."""
    info = svc.save_book_as_template(db_session, book_id=finished_word_book.id, name="قالب التجربة")
    stored = svc.resolve_template_path(info.name)
    assert stored.exists()
    # stored file is tokenized and valid
    from app.core.book_template_retokenize import validate_book_template

    validate_book_template(stored)


def test_save_collision_409(db_session, finished_word_book):
    svc.save_book_as_template(db_session, book_id=finished_word_book.id, name="مكرر")
    with pytest.raises(AppError) as ei:
        svc.save_book_as_template(db_session, book_id=finished_word_book.id, name="مكرر")
    assert ei.value.http_status == 409
