"""Passport Release List: styled item rows + unfilled slots trimmed.

Both are page-geometry regressions — the form used to render 2 pages even with
zero employees, and only the first 3 of its 15 item rows carried the table
styling (the rest were raw Excel-paste leftovers: no borders, no centring, Times
New Roman 10pt). Page count itself needs Word COM, so it is checked by hand via
the render script; these assertions cover the causes.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn

from app.core.docx_engine import DocxEngine

TEMPLATES = Path(__file__).resolve().parents[1] / "templates"
FIRST_ITEM_ROW = 7
ITEM_SLOTS = 15
SIGNATURE_LABEL = "تم إستلام جوازات السفر"


def _items(n: int) -> list[dict[str, str]]:
    return [
        {
            "employee_id": f"G{3000 + i}",
            "name": "محمد عبدالله الهاشمي",
            "nationality": "الهند",
            "passport_no": f"P{1234567 + i}",
        }
        for i in range(n)
    ]


def _render(tmp_path: Path, n: int) -> Any:
    out = tmp_path / f"list_{n}.docx"
    DocxEngine(TEMPLATES).fill("Passport Release List", {"items": _items(n)}, out)
    return Document(str(out)).tables[0]


def _border(cell: Any, edge: str) -> str | None:
    borders = cell._tc.tcPr.find(qn("w:tcBorders"))
    if borders is None:
        return None
    side = borders.find(qn(f"w:{edge}"))
    return None if side is None else str(side.get(qn("w:val")))


def test_every_item_slot_is_styled_like_the_first() -> None:
    """Slots 3-14 used to be borderless, left-aligned and a different font, so
    the 4th employee onward rendered as naked text under the table."""
    table = Document(str(TEMPLATES / "GSSG-HR_300-004b_Passport_Release_List.docx")).tables[0]
    first = table.rows[FIRST_ITEM_ROW]
    for slot in range(ITEM_SLOTS):
        row = table.rows[FIRST_ITEM_ROW + slot]
        assert row.cells[0].text.strip() == f"{{{{ item({slot}, 'employee_id') }}}}"
        for col, reference in enumerate(first.cells):
            cell = row.cells[col]
            assert _border(cell, "left") == _border(reference, "left") != "nil"
            assert _border(cell, "bottom") == _border(reference, "bottom") != "nil"
            jc = cell.paragraphs[0]._p.get_or_add_pPr().find(qn("w:jc"))
            assert jc is not None and jc.get(qn("w:val")) == "center"


def test_unfilled_slots_are_removed(tmp_path: Path) -> None:
    """The blank tail rows are what pushed the signature block onto page 2."""
    for n in (0, 1, 8):
        table = _render(tmp_path, n)
        filled = [r for r in table.rows if r.cells[0].text.strip().startswith("G3")]
        assert len(filled) == n
        blank_after_header = [
            r
            for r in table.rows[FIRST_ITEM_ROW : FIRST_ITEM_ROW + ITEM_SLOTS]
            if not any(c.text.strip() for c in r.cells)
        ]
        assert blank_after_header == []


def test_full_list_keeps_all_fifteen_rows(tmp_path: Path) -> None:
    table = _render(tmp_path, ITEM_SLOTS)
    filled = [r for r in table.rows if r.cells[0].text.strip().startswith("G3")]
    assert len(filled) == ITEM_SLOTS


def test_signature_block_survives_the_trim(tmp_path: Path) -> None:
    """Deleting rows must not touch the block below them (same table)."""
    for n in (0, 8, ITEM_SLOTS):
        table = _render(tmp_path, n)
        assert any(SIGNATURE_LABEL in r.cells[0].text for r in table.rows)
