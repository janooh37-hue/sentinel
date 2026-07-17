"""Tests for backend/app/core/book_text.py (Task 15).

TDD: RED → implement → GREEN.

Covers:
- normalize_ar: alef variants, tatweel, taa-marbuta->ha, alef-maqsura->ya, diacritics, whitespace.
- docx_to_text: paragraphs AND table cells extracted.
- html_to_text: tags stripped, entities decoded.
- build_search_text: concatenate+normalise.
- Integration: finish_word_session populates book.search_text AND an
  alef-variant FTS query (via books_fts) finds the book.
"""

from __future__ import annotations

import importlib.util
import pathlib
import secrets
import sqlite3
from datetime import UTC, datetime
from pathlib import Path

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _get_migration_0057():
    p = pathlib.Path(__file__).parent.parent / "app/db/migrations/versions/0057_books_fts.py"
    spec = importlib.util.spec_from_file_location("migration_0057", p)
    assert spec and spec.loader
    mod = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(mod)  # type: ignore[union-attr]
    return mod


def _fts_conn_from_db_session(db_session) -> sqlite3.Connection:
    """Open a raw sqlite3 connection to the same in-memory DB used by db_session,
    then apply the FTS DDL from migration 0057 so MATCH queries work."""
    raw = db_session.get_bind().raw_connection()
    mod = _get_migration_0057()
    for sql in mod.UPGRADE_SQL:
        raw.executescript(sql)
    raw.commit()
    return raw  # type: ignore[return-value]


def _make_tiny_docx(tmp_path: Path, para_text: str = "الأقفال", cell_text: str = "مرحبا") -> Path:
    """Build a minimal .docx with one paragraph and one 1x1 table cell."""
    from docx import Document

    doc = Document()
    doc.add_paragraph(para_text)
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).text = cell_text
    p = tmp_path / "tiny.docx"
    doc.save(str(p))
    return p


def _settings(tmp_path: Path):
    from app.config import Settings

    return Settings(data_dir=tmp_path / "data", templates_dir=tmp_path / "templates")


def _user(db):
    from app.db.models import User

    u = User(email=f"{secrets.token_hex(4)}@test.ae", password_hash="x", status="active")
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


def _make_book_with_session(db, user, tmp_path, *, subject="test book"):
    from app.db.models import Book, BookCategory, BookEditSession
    from app.db.repos import refs_repo

    if db.get(BookCategory, "GS") is None:
        db.add(BookCategory(id="GS", prefix="GS"))
        db.commit()

    ref = refs_repo.allocate_ref_with_retry(db, "GS")
    book = Book(
        category_id="GS",
        ref_number=ref,
        subject=subject,
        approval_state="none",
        submitted_by_user_id=user.id,
    )
    db.add(book)
    db.flush()

    working_dir = tmp_path / "data" / "editing" / f"book-{book.id}"
    working_dir.mkdir(parents=True, exist_ok=True)
    working_path = working_dir / f"{ref.replace('/', '-')}.docx"

    # Write a real minimal docx so docx_to_text can parse it
    from docx import Document as DocxDocument

    d = DocxDocument()
    d.add_paragraph("الأقفال الأمنية")  # contains alef variant أ
    working_path_obj = working_path
    d.save(str(working_path_obj))

    token = secrets.token_urlsafe(32)
    session = BookEditSession(
        book_id=book.id,
        user_id=user.id,
        token=token,
        working_path=str(working_path),
        state="active",
        last_put_at=datetime.now(UTC).replace(tzinfo=None),
    )
    db.add(session)
    db.commit()
    db.refresh(book)
    db.refresh(session)
    return book, session


# ---------------------------------------------------------------------------
# normalize_ar
# ---------------------------------------------------------------------------


def test_normalize_ar_alef_variants():
    """Alef variants (U+0623/0625/0622/0671) all collapse to plain alef (U+0627)."""
    from app.core.book_text import normalize_ar

    # With diacritic on alef + alef variant in the word
    assert normalize_ar("الأَقْفال") == normalize_ar("الاقفال")


def test_normalize_ar_tatweel():
    from app.core.book_text import normalize_ar

    assert normalize_ar("الـكتاب") == normalize_ar("الكتاب")


def test_normalize_ar_taa_marbuta():
    from app.core.book_text import normalize_ar

    # taa marbuta (U+0629) -> ha (U+0647)
    assert normalize_ar("مدرسة") == "مدرسه"


def test_normalize_ar_alef_maqsura():
    from app.core.book_text import normalize_ar

    # alef maqsura (U+0649) -> ya (U+064A)
    assert normalize_ar("مستشفى") == "مستشفي"


def test_normalize_ar_whitespace():
    from app.core.book_text import normalize_ar

    assert normalize_ar("  كلمة   أخرى  ") == "كلمه اخري"


# ---------------------------------------------------------------------------
# docx_to_text
# ---------------------------------------------------------------------------


def test_docx_to_text_paragraphs_and_table_cells(tmp_path):
    from app.core.book_text import docx_to_text

    p = _make_tiny_docx(tmp_path, para_text="نص الفقرة", cell_text="نص الخلية")
    text = docx_to_text(p)
    assert "نص الفقرة" in text
    assert "نص الخلية" in text


# ---------------------------------------------------------------------------
# html_to_text
# ---------------------------------------------------------------------------


def test_html_to_text_strips_tags():
    from app.core.book_text import html_to_text

    result = html_to_text("<p>مرحبا <b>بك</b></p>")
    assert result == "مرحبا بك"


def test_html_to_text_decodes_entities():
    from app.core.book_text import html_to_text

    result = html_to_text("<p>&amp; &lt; &gt; &quot;</p>")
    assert "&" in result
    assert "<" in result


def test_html_to_text_plain_text_passthrough():
    from app.core.book_text import html_to_text

    # No tags → returns as-is (stripped)
    assert html_to_text("مجرد نص") == "مجرد نص"


# ---------------------------------------------------------------------------
# build_search_text
# ---------------------------------------------------------------------------


def test_build_search_text_concatenates_and_normalises():
    from app.core.book_text import build_search_text, normalize_ar

    result = build_search_text(subject="الأمن", ref="GS-0001", body="مستشفى")
    # subject الأمن → الامن (alef normalised); body مستشفى → مستشفي
    expected = normalize_ar("الأمن  GS-0001  مستشفى")
    assert result == expected


def test_build_search_text_none_subject():
    from app.core.book_text import build_search_text

    result = build_search_text(subject=None, ref="GS-0001", body="")
    assert "GS-0001" in result


# ---------------------------------------------------------------------------
# Integration: finish_word_session populates search_text + FTS MATCH works
# ---------------------------------------------------------------------------


def test_finish_populates_search_text_and_fts_match(db_session, tmp_path, monkeypatch):
    """After finish_word_session, book.search_text is set AND an alef-variant
    FTS MATCH on the raw sqlite3 connection finds the book row."""
    from app.services import word_book_service

    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))
    monkeypatch.setattr(
        "app.services.word_book_service.convert_docx_to_pdf",
        lambda p: None,
    )

    user = _user(db_session)
    subject = "الأقفال الأمنية"  # contains أ (alef with hamza above)
    book, _session = _make_book_with_session(db_session, user, tmp_path, subject=subject)

    # Set book.subject so build_search_text sees it after finish
    book.subject = subject
    db_session.commit()

    # Install the FTS DDL on the raw connection BEFORE finish so the triggers fire
    raw = _fts_conn_from_db_session(db_session)

    word_book_service.finish_word_session(db_session, user=user, book_id=book.id)

    db_session.refresh(book)

    # 1. search_text is populated
    assert book.search_text, "book.search_text must be set after finish"

    # 2. Normalised form: alef-with-hamza (U+0623) -> plain alef (U+0627) in search_text
    from app.core.book_text import normalize_ar

    assert normalize_ar(subject) in book.search_text or normalize_ar("الأقفال") in book.search_text

    # 3. FTS MATCH: query with unaccented/differently-spelt alef finds the book.
    # The FTS tokenizer's `remove_diacritics 2` + our alef normalisation means
    # "الاقفال" (plain alef) should match "الأقفال" stored as "الاقفال" in search_text.
    query_term = "الاقفال"  # plain alef, no hamza — variant of أ
    rows = raw.execute(
        "SELECT rowid FROM books_fts WHERE books_fts MATCH ?", (query_term,)
    ).fetchall()
    assert rows, (
        f"FTS MATCH for {query_term!r} should find book {book.id}; search_text={book.search_text!r}"
    )
