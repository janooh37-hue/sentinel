# backend/tests/test_material_request_signing.py
"""Regression tests for SC-0425-class bugs in the Material Request Form
signing path:

  1. The manager signature must NOT grow the signature row (which pushed the
     form onto a second page). It is floated *behind text* (zero layout
     height), the same technique the Leave Permit / Admin Leave forms use.
  2. Re-rendering a version at sign-time must reproduce the manager NAME that
     was chosen at generation (stored on ``Book.doc_manager_id``); the signed
     copy used to render it blank.
"""
from __future__ import annotations

import io
from pathlib import Path

import pytest
from docx import Document
from PIL import Image

from app.core.docx_engine import DocxEngine
from app.services.document_service import _TEMPLATES_DIR


def _make_sig_png(path: Path) -> Path:
    # A diagonal stroke so the ink bounding box has real width AND height
    # (a flat 1px line crops to ~zero height and can't be bottom-aligned).
    img = Image.new("RGBA", (400, 168), (255, 255, 255, 0))
    for x in range(40, 360):
        y = 20 + int((x - 40) * 120 / 320)
        for dy in (-1, 0, 1):
            img.putpixel((x, y + dy), (0, 0, 0, 255))
    img.save(path)
    return path


def _mrf_sig_cell(doc):
    """Table 4 / row 5 / cell 3 holds {{ manager_sig }} (the verifier
    Signature cell)."""
    return doc.tables[4].rows[5].cells[3]


def test_manager_signature_is_floated_behind_text(tmp_path):
    sig = _make_sig_png(tmp_path / "sig.png")
    out = tmp_path / "mrf.docx"
    data = {
        "manager_name": "SAEED RASHED SANAD KHALFAN ALYAHYAEE",
        "manager_sig_path": str(sig),
        "items": [{"sno": "1", "description": "Widget", "qty": "2"}],
        "_sig_size_mm": 45,
    }
    DocxEngine(_TEMPLATES_DIR).fill("Material Request Form", data, out)

    doc = Document(str(out))
    cell_xml = _mrf_sig_cell(doc)._tc.xml
    # Manager name still renders in its own cell.
    assert "SAEED RASHED" in doc.tables[4].rows[4].cells[1].text
    # The signature image must be present...
    assert "w:drawing" in cell_xml, "signature image missing from signature cell"
    # ...and FLOATING (behind text), not inline — inline would add row height
    # and bump the table onto a second page.
    assert 'behindDoc="1"' in cell_xml, "manager signature is not floated behind text"
    assert "wp:inline" not in cell_xml, "manager signature is still inline (adds height)"
    # ...and lifted UP (negative vertical offset) so it rests on the signature
    # line and rises into the empty Contact cell above, instead of hanging down
    # across the "Supply Chain" divider below.
    import re

    voffsets = re.findall(r"<wp:positionV[^>]*>\s*<wp:posOffset>(-?\d+)</wp:posOffset>", cell_xml)
    assert voffsets and int(voffsets[0]) < 0, f"signature not lifted up (vOffset={voffsets})"
