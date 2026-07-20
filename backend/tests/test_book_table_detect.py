from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from app.core.book_table import detect_table_schema


def _tbl(doc, headers, data_rows):
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for r, row in enumerate(data_rows, start=1):
        for c, val in enumerate(row):
            t.cell(r, c).text = val
    return t


def test_clean_table_returns_headers():
    doc = Document()
    _tbl(doc, ["الاسم", "الرقم", "الجهة"], [["أحمد", "101", "الأمن"]])
    assert detect_table_schema(doc) == ["الاسم", "الرقم", "الجهة"]


def test_no_table_returns_none():
    doc = Document()
    doc.add_paragraph("نص بدون جدول")
    assert detect_table_schema(doc) is None


def test_two_tables_returns_none():
    doc = Document()
    for _ in range(2):
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "هـ"
    assert detect_table_schema(doc) is None


def test_vmerge_data_row_returns_none():
    doc = Document()
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "هـ"
    t.cell(0, 1).text = "و"
    tcPr = t.cell(1, 0)._tc.get_or_add_tcPr()
    etree.SubElement(tcPr, qn("w:vMerge")).set(qn("w:val"), "restart")
    assert detect_table_schema(doc) is None


def test_gridspan_data_cell_returns_none():
    doc = Document()
    t = doc.add_table(rows=2, cols=3)
    for i, h in enumerate(["أ", "ب", "ج"]):
        t.cell(0, i).text = h
    tcPr = t.cell(1, 0)._tc.get_or_add_tcPr()
    etree.SubElement(tcPr, qn("w:gridSpan")).set(qn("w:val"), "2")
    assert detect_table_schema(doc) is None


def test_header_only_table_returns_headers():
    doc = Document()
    t = doc.add_table(rows=1, cols=2)
    t.cell(0, 0).text = "الاسم"
    t.cell(0, 1).text = "التاريخ"
    assert detect_table_schema(doc) == ["الاسم", "التاريخ"]


def test_empty_body_returns_none():
    assert detect_table_schema(Document()) is None
