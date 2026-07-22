# backend/tests/test_arabic_rtl_table_width.py
from docx import Document
from docx.oxml.ns import qn

from app.core.arabic_rtl import html_to_docx


def _tblW(html: str) -> int:
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(html, p)
    tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
    return int(tblPr.find(qn("w:tblW")).get(qn("w:w")))


def _content_twips() -> int:
    # Default python-docx Letter section: 8.5in - 2.5in margins = 6.0in = 8640.
    return 8640


def test_px_width_attr_honored() -> None:
    # Word paste: <table width=410 ...> -> 410px = 6150 twips < content 8640.
    assert _tblW('<table width="410"><tr><td>A</td></tr></table>') == 6150


def test_pt_style_width_honored() -> None:
    # 300pt * 20 = 6000.
    assert _tblW('<table style="width:300pt"><tr><td>A</td></tr></table>') == 6000


def test_percent_width_of_content() -> None:
    assert _tblW('<table style="width:50%"><tr><td>A</td></tr></table>') == _content_twips() // 2


def test_width_capped_at_content() -> None:
    assert _tblW('<table width="2000"><tr><td>A</td></tr></table>') == _content_twips()


def test_no_width_stays_full_content() -> None:
    assert _tblW("<table><tr><td>A</td></tr></table>") == _content_twips()


def test_width_auto_enables_autofit() -> None:
    # style="width:auto" -> Word AutoFit-to-Contents: tblW type=auto (w=0),
    # tblLayout=autofit, jc=center, and cell tcW type=auto (columns hug text).
    doc = Document()
    html_to_docx(
        '<table style="width:auto" dir="rtl"><tr><td>الاسم</td><td>1</td></tr></table>',
        doc.add_paragraph(),
    )
    tbl = doc.tables[0]
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    tblW = tblPr.find(qn("w:tblW"))
    assert tblW.get(qn("w:type")) == "auto" and tblW.get(qn("w:w")) == "0"
    assert tblPr.find(qn("w:tblLayout")).get(qn("w:type")) == "autofit"
    assert tblPr.find(qn("w:jc")).get(qn("w:val")) == "center"
    tcW = tbl.rows[0].cells[0]._tc.find(qn("w:tcPr")).find(qn("w:tcW"))
    assert tcW.get(qn("w:type")) == "auto"


def test_default_table_stays_fixed_right() -> None:
    # No width:auto -> unchanged: fixed layout, jc=right (regression guard).
    doc = Document()
    html_to_docx('<table dir="rtl"><tr><td>A</td></tr></table>', doc.add_paragraph())
    tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
    assert tblPr.find(qn("w:tblLayout")).get(qn("w:type")) == "fixed"
    assert tblPr.find(qn("w:jc")).get(qn("w:val")) == "right"
