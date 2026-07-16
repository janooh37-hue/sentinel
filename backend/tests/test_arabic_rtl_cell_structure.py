# backend/tests/test_arabic_rtl_cell_structure.py
from docx import Document
from docx.shared import Pt

from app.core.arabic_rtl import html_to_docx


def _cell(html, r=0, c=0):
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(html, p)
    return doc.tables[0].rows[r].cells[c]


def test_word_paste_cell_renders_single_paragraph():
    cell = _cell("<table><tr><td><p>محمد مشرف</p></td></tr></table>")
    assert len(cell.paragraphs) == 1
    assert cell.paragraphs[0].text == "محمد مشرف"


def test_word_paste_cell_with_intertag_whitespace_single_paragraph():
    cell = _cell("<table><tr>\n <td>\n  <p>A</p>\n </td>\n</tr></table>")
    assert len(cell.paragraphs) == 1
    assert cell.paragraphs[0].text == "A"


def test_cell_with_two_paragraphs_keeps_both():
    cell = _cell("<table><tr><td><p>A</p><p>B</p></td></tr></table>")
    assert [q.text for q in cell.paragraphs] == ["A", "B"]


def test_every_cell_paragraph_is_hugged():
    cell = _cell("<table><tr><td><p>A</p><p>B</p></td></tr></table>")
    for q in cell.paragraphs:
        assert q.paragraph_format.space_before == Pt(0)
        assert q.paragraph_format.space_after == Pt(0)
        assert q.paragraph_format.line_spacing == 1.0


def test_cell_explicit_line_height_wins_over_hug():
    cell = _cell('<table><tr><td><p style="line-height: 2">A</p></td></tr></table>')
    assert cell.paragraphs[0].paragraph_format.line_spacing == 2.0


def test_bare_text_cell_unchanged():
    cell = _cell("<table><tr><td>A</td></tr></table>")
    assert len(cell.paragraphs) == 1
    assert cell.paragraphs[0].text == "A"


def test_nbsp_only_cell_keeps_one_paragraph():
    # The GSSG insert-table button fills body cells with &nbsp;.
    cell = _cell("<table><tr><td> </td></tr></table>")
    assert len(cell.paragraphs) == 1
