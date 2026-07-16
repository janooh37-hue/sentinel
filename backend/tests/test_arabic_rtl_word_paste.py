# backend/tests/test_arabic_rtl_word_paste.py
"""Word-paste integration pin: markup shape taken from a real Word filtered-
HTML export of the operator's security-permits letter (mso classes, per-cell
<p class=MsoNormal>, pt heights/widths, inter-tag newlines). If this breaks,
pasted Word tables regressed."""

from docx import Document
from docx.oxml.ns import qn

from app.core.arabic_rtl import html_to_docx

WORD_PASTE = """
<p class=MsoNormal dir=RTL style='text-align:right;direction:rtl'><b><i>
<span lang=AR-SA style='font-size:13.0pt'>الرقم:1/ 5 /GSSG/ 140</span></i></b></p>

<div align=right>
<table class=MsoNormalTable dir=rtl border=0 cellspacing=0 cellpadding=0
 width=614 style='width:460.7pt;border-collapse:collapse'>
 <tr style='height:7.55pt'>
  <td width=76 nowrap style='width:56.7pt;border:solid windowtext 1.0pt;
  background:#004F88;padding:0in 5.4pt 0in 5.4pt;height:7.55pt'>
  <p class=MsoNormal align=center dir=RTL style='text-align:center;direction:
  rtl'><b><span lang=AR-SA style='color:white'>م</span></b></p>
  </td>
  <td width=228 nowrap style='width:171.05pt;border:solid windowtext 1.0pt;
  background:#004F88;padding:0in 5.4pt 0in 5.4pt;height:7.55pt'>
  <p class=MsoNormal align=center dir=RTL style='text-align:center;direction:
  rtl'><b><span lang=AR-SA style='color:white'>الاســــــم
  </span></b></p>
  </td>
 </tr>
 <tr style='height:15.75pt'>
  <td width=76 nowrap style='width:56.7pt;padding:0in 5.4pt 0in 5.4pt;
  height:15.75pt'>
  <p class=MsoNormal align=center dir=RTL style='text-align:center;direction:
  rtl'><span lang=AR-SA style='color:black'>1</span></p>
  </td>
  <td width=228 nowrap style='width:171.05pt;padding:0in 5.4pt 0in 5.4pt;
  height:15.75pt'>
  <p class=MsoNormal align=center dir=RTL style='text-align:center;direction:
  rtl'><span lang=AR-SA style='color:black'>محمد
  مشرف حسين محمد حسن </span></p>
  </td>
 </tr>
</table>
</div>

<p class=MsoNormal dir=RTL style='text-align:justify;direction:rtl'>
<span lang=AR-SA style='font-size:15.0pt'>للتفضل بالعلم وإجراءاتكم لطفاً،،،</span></p>
"""


def _render():
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(WORD_PASTE, p)
    return doc


def test_every_cell_is_exactly_one_clean_paragraph():
    doc = _render()
    t = doc.tables[0]
    for row in t.rows:
        for cell in row.cells:
            assert len(cell.paragraphs) == 1
            assert "\n" not in cell.paragraphs[0].text


def test_arabic_name_collapsed_to_single_spaced_text():
    doc = _render()
    assert doc.tables[0].rows[1].cells[1].text == "محمد مشرف حسين محمد حسن"


def test_row_heights_and_cantsplit_stamped():
    doc = _render()
    t = doc.tables[0]
    for r, expected in enumerate(("151", "315")):  # 7.55pt*20, 15.75pt*20
        trPr = t.rows[r]._tr.find(qn("w:trPr"))
        assert trPr.find(qn("w:cantSplit")) is not None
        assert trPr.find(qn("w:trHeight")).get(qn("w:val")) == expected


def test_table_width_matches_word_pt_width():
    doc = _render()
    tblPr = doc.tables[0]._tbl.find(qn("w:tblPr"))
    # 460.7pt = 9214 twips, capped at the test doc's 8640-twip content width.
    # (In the real General Book template the A4 content width is ~10477 twips,
    # so the Word width is honored uncapped there.)
    assert int(tblPr.find(qn("w:tblW")).get(qn("w:w"))) == 8640


def test_narrative_order_preserved_around_table():
    doc = _render()
    body = doc.element.body
    kinds = [e.tag.split("}")[1] for e in body if e.tag.split("}")[1] in ("p", "tbl")]
    first_tbl = kinds.index("tbl")
    assert "p" in kinds[:first_tbl]  # الرقم line before the table
    assert "p" in kinds[first_tbl + 1 :]  # closing line after the table


def test_header_cell_shading_survives():
    doc = _render()
    tcPr = doc.tables[0].rows[0].cells[0]._tc.find(qn("w:tcPr"))
    shd = tcPr.find(qn("w:shd"))
    assert shd is not None and shd.get(qn("w:fill")) == "004F88"
