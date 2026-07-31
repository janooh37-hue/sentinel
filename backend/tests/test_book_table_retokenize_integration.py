from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml.parser import parse_xml
from docx.oxml.xmlchemy import BaseOxmlElement

from app.core.book_template_retokenize import retokenize_general_book
from app.core.book_text import docx_to_text


def _book_with_table(tmp_path: Path, headers: list[str], data_rows: list[list[str]]) -> Path:
    p = tmp_path / "book_tbl.docx"
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: قائمة الموظفين في الإدارة المحترمة")
    doc.add_paragraph("تفاصيل القائمة:")
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for r, row in enumerate(data_rows, start=1):
        for c, v in enumerate(row):
            t.cell(r, c).text = v
    doc.save(str(p))
    return p


def test_retokenize_normalizes_table(tmp_path: Path) -> None:
    p = _book_with_table(tmp_path, ["الاسم", "الرقم"], [["أحمد", "G-001"], ["خالد", "G-002"]])
    retokenize_general_book(p)
    text = docx_to_text(p)
    assert "{%tr for row in table_rows %}" in text
    assert "{{ row.c0 }}" in text
    assert "G-001" not in text  # PII stripped


def test_retokenize_table_idempotent(tmp_path: Path) -> None:
    p = _book_with_table(tmp_path, ["أ", "ب"], [["x", "y"]])
    retokenize_general_book(p)
    t1 = docx_to_text(p)
    retokenize_general_book(p)
    t2 = docx_to_text(p)
    assert t1 == t2


def _loop_row_cell(path: Path) -> BaseOxmlElement:
    """First cell of the {{ row.cN }} loop row of the single normalized table."""
    tbl = Document(str(path)).element.body.findall(qn("w:tbl"))[0]
    for tr in tbl.findall(qn("w:tr")):
        tc = tr.findall(qn("w:tc"))[0]
        if "row.c0" in "".join(t.text or "" for t in tc.findall(f".//{qn('w:t')}")):
            return tc
    raise AssertionError("no loop row found")


def test_loop_row_keeps_data_row_styling_not_headers(tmp_path: Path) -> None:
    """The repeating row must look like the operator's BODY row: header shading
    (a coloured title band) must not bleed into it, and the body cells' own
    alignment must survive."""
    p = _book_with_table(tmp_path, ["الاسم", "الرقم"], [["أحمد", "G-001"]])
    doc = Document(str(p))
    t = doc.tables[0]
    # header: red band, no explicit alignment; body: light grey, centred
    for cell, fill in ((t.cell(0, 0), "C00000"), (t.cell(1, 0), "D9D9D9")):
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:fill="{fill}"/>')
        )
    t.cell(1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.save(str(p))

    retokenize_general_book(p)

    tc = _loop_row_cell(p)
    shd = tc.find(f".//{qn('w:shd')}")
    assert shd is not None and shd.get(qn("w:fill")) == "D9D9D9", "loop row took the header's fill"
    jc = tc.find(f".//{qn('w:jc')}")
    assert jc is not None and jc.get(qn("w:val")) == "center", "data-row alignment lost"


def test_retokenize_plain_book_no_table_tokens(tmp_path: Path) -> None:
    p = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("السيد / مدير الإدارة المحترم")
    doc.add_paragraph("الموضوع: موضوع الكتاب في نص طويل نسبياً هنا")
    doc.add_paragraph("نص الكتاب هنا")
    doc.save(str(p))
    retokenize_general_book(p)
    text = docx_to_text(p)
    assert "{%tr" not in text
    assert "{{ ref }}" in text


def test_retokenize_two_table_book_no_tokens(tmp_path: Path) -> None:
    p = tmp_path / "two.docx"
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: كتاب عادي بجدولين بيانيين للاختبار")
    doc.add_paragraph("نص الكتاب")
    for _ in range(2):
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "هـ"
    doc.save(str(p))
    retokenize_general_book(p)
    text = docx_to_text(p)
    assert "{%tr" not in text
