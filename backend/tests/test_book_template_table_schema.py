"""Tests for book_template_service.table_schema_for and delete_template."""

import pytest
from docx import Document

from app.api.errors import AppError
from app.core.book_template_retokenize import retokenize_general_book
from app.services import book_template_service as svc


def _patch_templates_dir(monkeypatch, tmp_path):
    """Point book_template_service.templates_dir at tmp_path."""
    monkeypatch.setattr(svc, "templates_dir", lambda: tmp_path)


def _plain(path):
    """Retokenized General Book docx with NO table."""
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: خطاب نصّي بدون جدول للاختبار")
    doc.add_paragraph("نص")
    doc.save(str(path))
    retokenize_general_book(path)


def _table(path, headers):
    """Retokenized General Book docx WITH a table (headers normalized to loop)."""
    doc = Document()
    doc.add_paragraph("الرقم: 1/5/141")
    doc.add_paragraph("التاريخ: 20-07-2026")
    doc.add_paragraph("الموضوع: كتاب جدول بيانات للاختبار المحترم")
    t = doc.add_table(rows=2, cols=len(headers))
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for i in range(len(headers)):
        t.cell(1, i).text = f"بيانات {i}"
    doc.save(str(path))
    retokenize_general_book(path)


def test_table_schema_plain(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    _plain(tmp_path / "plain.docx")
    has, cols = svc.table_schema_for("plain.docx")
    assert has is False
    assert cols == []


def test_table_schema_table(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    headers = ["الاسم", "الرقم الوظيفي", "الجهة"]
    _table(tmp_path / "table.docx", headers)
    has, cols = svc.table_schema_for("table.docx")
    assert has is True
    assert cols == headers


def test_table_schema_missing_404(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    with pytest.raises(AppError) as exc_info:
        svc.table_schema_for("ghost.docx")
    assert exc_info.value.http_status == 404
    assert exc_info.value.code == "TEMPLATE_NOT_FOUND"


def test_delete_removes_file(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    p = tmp_path / "del.docx"
    _plain(p)
    assert p.exists()
    svc.delete_template("del.docx")
    assert not p.exists()


def test_delete_missing_404(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    with pytest.raises(AppError) as exc_info:
        svc.delete_template("ghost.docx")
    assert exc_info.value.http_status == 404
    assert exc_info.value.code == "TEMPLATE_NOT_FOUND"


def test_delete_traversal_rejected(tmp_path, monkeypatch):
    _patch_templates_dir(monkeypatch, tmp_path)
    with pytest.raises(AppError) as exc_info:
        svc.delete_template("../evil.docx")
    assert exc_info.value.http_status == 422
    assert exc_info.value.code == "TEMPLATE_BAD_NAME"
