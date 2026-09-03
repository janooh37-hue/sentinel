"""The canonical General Book template renders an Arabic ref line above the
date when ``ref`` is provided, and omits it (three guard paragraphs collapse)
when it is not. Asserts the ARABIC string per the i18n lesson."""

from pathlib import Path

from app.core.book_text import docx_to_text
from app.core.docx_engine import DocxEngine
from app.services.document_service import GENERAL_BOOK_BODY_SENTINEL

TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"

_BASE_DATA = {
    "subject": "اختبار",
    "body": GENERAL_BOOK_BODY_SENTINEL,
    "body_html": "<p>نص تجريبي</p>",
    "recipient_name": "السيد المدير",
    "cc": [],
    "submitter_g": "G-1234",
}


def test_ref_line_renders_above_date(tmp_path):
    out = tmp_path / "out.docx"
    DocxEngine(TEMPLATES_DIR).fill("General Book", {**_BASE_DATA, "ref": "1/5/141"}, out)
    text = docx_to_text(out)
    assert "الرقم: 1/5/141" in text
    # above the date: الرقم line appears before التاريخ in document order
    assert text.index("الرقم:") < text.index("التاريخ:")


def test_ref_line_absent_without_ref(tmp_path):
    out = tmp_path / "out.docx"
    DocxEngine(TEMPLATES_DIR).fill("General Book", dict(_BASE_DATA), out)
    text = docx_to_text(out)
    assert "الرقم" not in text  # preview/serial-free renders show no ref line


def test_ref_run_marked_ltr(tmp_path):
    """The dynamic reference stays in stored order in an explicit LTR run."""
    from docx import Document

    out = tmp_path / "out.docx"
    DocxEngine(TEMPLATES_DIR).fill("General Book", {**_BASE_DATA, "ref": "1/5/141"}, out)
    doc = Document(str(out))
    ref_para = next(p for p in doc.paragraphs if "1/5/141" in p.text)
    ref_runs = [r for r in ref_para.runs if r.text.startswith("1/")]
    assert ref_runs, "ref value must be in its own run"
    assert all(r.font.rtl is False for r in ref_runs)


def _header_text(docx_path) -> str:
    from docx import Document

    doc = Document(str(docx_path))
    parts = []
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header):
            parts.extend(p.text for p in hdr.paragraphs)
    return "\n".join(parts)


def test_ref_renders_ltr_segment_order_directly_after_label(tmp_path):
    from docx import Document

    out = tmp_path / "out.docx"
    DocxEngine(TEMPLATES_DIR).fill("General Book", {**_BASE_DATA, "ref": "1/15/141"}, out)
    doc = Document(str(out))
    ref_para = next(p for p in doc.paragraphs if "1/15/141" in p.text)
    non_empty = [r for r in ref_para.runs if r.text.strip()]
    label_idx = next(i for i, r in enumerate(non_empty) if "الرقم" in r.text)
    value_idx = next(i for i, r in enumerate(non_empty) if "1/15/141" in r.text)
    assert value_idx == label_idx + 1
    value_run = non_empty[value_idx]
    assert value_run.text == "1/15/141"  # dynamic stored order; Word must not reverse it
    assert value_run.font.rtl is False


def test_ref_line_preserves_canonical_template_size_with_rtl_label(tmp_path):
    from docx import Document
    from docx.oxml.ns import qn

    out = tmp_path / "formatted.docx"
    DocxEngine(TEMPLATES_DIR).fill("General Book", {**_BASE_DATA, "ref": "1/5/142"}, out)
    doc = Document(str(out))
    paragraph = next(p for p in doc.paragraphs if "1/5/142" in p.text)
    date_paragraph = next(p for p in doc.paragraphs if "التاريخ:" in p.text)
    runs = [r for r in paragraph.runs if r.text]
    date_size = next(r.font.size for r in date_paragraph.runs if r.text)

    assert [r.text for r in runs] == ["الرقم: ", "1/5/142"]
    assert paragraph._p.pPr.find(qn("w:bidi")) is not None
    assert runs[0].font.rtl is True
    assert runs[1].font.rtl is False
    assert date_size is not None and date_size.pt == 10
    assert all(r.font.size == date_size for r in runs)
    assert all(not r.font.italic for r in runs)

def test_library_template_preserves_ref_run_format(tmp_path):
    from docx import Document
    from docx.shared import Pt

    template = tmp_path / "library.docx"
    doc = Document()
    paragraph = doc.add_paragraph()
    template_run = paragraph.add_run("الرقم: {{ ref }}")
    template_run.font.name = "Arial"
    template_run.font.size = Pt(9)
    template_run.font.bold = True
    doc.add_paragraph(GENERAL_BOOK_BODY_SENTINEL)
    doc.save(template)

    out = tmp_path / "out.docx"
    DocxEngine(TEMPLATES_DIR).fill_general_book_path(
        template, {**_BASE_DATA, "ref": "1/5/142"}, out
    )
    rendered = Document(out)
    paragraph = next(p for p in rendered.paragraphs if "1/5/142" in p.text)
    runs = [r for r in paragraph.runs if r.text]

    assert [r.text for r in runs] == ["الرقم: ", "1/5/142"]
    assert runs[1].font.rtl is False
    assert all(r.font.name == "Arial" and r.font.size.pt == 9 and r.font.bold for r in runs)


def test_word_template_keeps_intro_without_large_blank_gap(tmp_path):
    from docx import Document

    out = tmp_path / "word-book.docx"
    data = {**_BASE_DATA, "body_html": ""}
    DocxEngine(TEMPLATES_DIR).fill("General Book", {**data, "ref": "1/5/142"}, out)
    paragraphs = [p.text.strip() for p in Document(out).paragraphs]

    subject_idx = next(i for i, text in enumerate(paragraphs) if text.startswith("الموضوع:"))
    intro_idx = paragraphs.index(
        "يطيب لنا أن نتقدم لكم بأطيب التحيات والتقدير، "
        "ننوه على الموضوع أعلاه انه"
    )
    closing_idx = paragraphs.index("للتفضل بالعلم وإجراءاتكم")

    assert intro_idx - subject_idx <= 2
    assert closing_idx - intro_idx <= 2


def test_word_book_has_ref_line_and_no_header_stamp(db_session, admin_user):
    """Word-path create: body الرقم line present, English Ref: stamp gone."""
    from app.services import word_book_service

    info = word_book_service.create_word_book(
        db_session,
        user=admin_user,
        classification_code="5/1",
        recipient_id=None,
        subject="اختبار القالب",
        cc=[],
        manager_id=None,
    )
    from app.db.models import BookEditSession

    session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    text = docx_to_text(Path(session.working_path))
    assert f"الرقم: {info.ref_number}" in text
    assert "Ref:" not in _header_text(session.working_path)
