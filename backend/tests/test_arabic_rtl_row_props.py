# backend/tests/test_arabic_rtl_row_props.py
from docx import Document
from docx.oxml.ns import qn

from app.core.arabic_rtl import html_to_docx


def _table(html):
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx(html, p)
    return doc.tables[0]


def _trPr(table, r):
    return table.rows[r]._tr.find(qn("w:trPr"))


def test_explicit_pt_row_height_becomes_atleast_trheight():
    t = _table('<table><tr style="height:15.75pt"><td>A</td></tr></table>')
    trPr = _trPr(t, 0)
    h = trPr.find(qn("w:trHeight"))
    assert h is not None
    assert h.get(qn("w:val")) == "315"  # 15.75pt * 20
    assert h.get(qn("w:hRule")) == "atLeast"


def test_px_height_attr_supported():
    t = _table('<table><tr height="21"><td>A</td></tr></table>')
    h = _trPr(t, 0).find(qn("w:trHeight"))
    assert h is not None
    assert h.get(qn("w:val")) == "315"  # 21px * 15


def test_no_height_no_trheight():
    t = _table("<table><tr><td>A</td></tr></table>")
    trPr = _trPr(t, 0)
    assert trPr is None or trPr.find(qn("w:trHeight")) is None


def test_every_row_gets_cantsplit():
    t = _table("<table><tr><td>A</td></tr><tr><td>B</td></tr></table>")
    for r in range(2):
        assert _trPr(t, r).find(qn("w:cantSplit")) is not None


def test_thead_row_repeats_as_header():
    t = _table(
        "<table><thead><tr><th>H</th></tr></thead><tbody><tr><td>B</td></tr></tbody></table>"
    )
    assert _trPr(t, 0).find(qn("w:tblHeader")) is not None
    trPr1 = _trPr(t, 1)
    assert trPr1 is None or trPr1.find(qn("w:tblHeader")) is None
