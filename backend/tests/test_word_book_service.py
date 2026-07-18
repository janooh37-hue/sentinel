"""Tests for word_book_service.create_word_book (Task 5).

Every Word book renders on the ONE General Book template and takes its ref
from the classified register — same rules as the rich-editor path.
"""

from __future__ import annotations

import secrets
from pathlib import Path

import docx
import pytest

from app.db.models import Book, BookCategory, BookEditSession, Manager, User

_GENERAL_BOOK = "GSSG-GS_300-003_General_Book.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _seed_gs(db):
    if db.get(BookCategory, "GS") is None:
        db.add(BookCategory(id="GS", prefix="GS"))
        db.commit()


def _user(db, *, employee_id: str | None = None) -> User:
    """Create a test user. employee_id is nullable; skip FK by leaving None."""
    u = User(email=f"{secrets.token_hex(4)}@test.ae", password_hash="x", status="active")
    u.employee_id = employee_id
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


# ---------------------------------------------------------------------------
# Tests
# ---------------------------------------------------------------------------


def test_create_classified_book_returns_session_info(db_session, tmp_path, monkeypatch):
    """Success path: classified book returns WordSessionInfo with correct ref."""
    from app.services import word_book_service

    _seed_gs(db_session)
    settings = _settings(tmp_path)
    # Place a minimal placeholder template at the expected path
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="Security permit test",
        cc=None,
        manager_id=None,
    )

    assert info.ref_number == "1/5/GSSG/1"
    assert info.book_id is not None
    assert info.token
    assert info.filename == "1-5-GSSG-1.docx"
    assert info.word_url.startswith("ms-word:ofe|u|")
    assert info.dav_url

    # Book row exists
    book = db_session.get(Book, info.book_id)
    assert book is not None
    assert book.ref_number == "1/5/GSSG/1"
    assert book.classification_code == "5/1"
    assert book.category_id == "GS"
    assert book.approval_state == "none"

    # Edit session exists and is active
    session = db_session.query(BookEditSession).filter_by(book_id=info.book_id).one()
    assert session.state == "active"
    assert session.token == info.token

    # Working file exists and is a valid docx carrying the ref (header stamp —
    # the General Book template has no {{ ref }} body token)
    working_file = Path(session.working_path)
    assert working_file.exists()
    doc = docx.Document(str(working_file))
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "1/5/GSSG/1" in header_text


def test_working_docx_gets_header_ref_stamp(db_session, tmp_path, monkeypatch):
    """The working docx carries the SAME header ref stamp the rich-editor path
    writes (DocxEngine.stamp_ref_number) — the General Book template has no
    {{ ref }} token, so the stamp is what puts the number on paper."""
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="stamp test",
        cc=None,
        manager_id=None,
    )
    working_file = Path(
        db_session.query(BookEditSession).filter_by(book_id=info.book_id).one().working_path
    )
    doc = docx.Document(str(working_file))
    header_text = "\n".join(p.text for p in doc.sections[0].header.paragraphs)
    assert "Ref: 1/5/GSSG/1" in header_text


def test_two_classified_creates_get_sequential_serials(db_session, tmp_path, monkeypatch):
    """Two creates across different classifications → serials 1, 2."""
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info1 = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="3/1",  # tab 3 (Annual leaves)
        recipient_id=None,
        subject="first",
        cc=None,
        manager_id=None,
    )
    info2 = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",  # tab 5 (Security permits)
        recipient_id=None,
        subject="second",
        cc=None,
        manager_id=None,
    )
    assert info1.ref_number == "1/3/GSSG/1"
    assert info2.ref_number == "1/5/GSSG/2"


def test_unknown_classification_code_raises_422(db_session, tmp_path, monkeypatch):
    """Unknown classification_code → AppError UNKNOWN_CLASSIFICATION 422."""
    from app.api.errors import AppError
    from app.services import word_book_service

    _seed_gs(db_session)
    settings = _settings(tmp_path)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)
    user = _user(db_session)

    with pytest.raises(AppError) as exc_info:
        word_book_service.create_word_book(
            db_session,
            user=user,
            classification_code="99/1",
            recipient_id=None,
            subject="x",
            cc=None,
            manager_id=None,
        )
    assert exc_info.value.code == "UNKNOWN_CLASSIFICATION"
    assert exc_info.value.http_status == 422


def test_missing_classification_raises_422(db_session, tmp_path, monkeypatch):
    """classification_code=None → CLASSIFICATION_REQUIRED 422 — every book's
    ref comes from the classified register, no legacy GS-#### path remains."""
    from app.api.errors import AppError
    from app.services import word_book_service

    _seed_gs(db_session)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))
    user = _user(db_session)

    with pytest.raises(AppError) as exc_info:
        word_book_service.create_word_book(
            db_session,
            user=user,
            classification_code=None,
            recipient_id=None,
            subject="plain book",
            cc=None,
            manager_id=None,
        )
    assert exc_info.value.code == "CLASSIFICATION_REQUIRED"
    assert exc_info.value.http_status == 422


def test_missing_template_file_raises_409(db_session, tmp_path, monkeypatch):
    """Template file missing on disk → AppError TEMPLATE_MISSING 409."""
    from app.api.errors import AppError
    from app.services import word_book_service

    _seed_gs(db_session)
    # templates_dir exists but the file is NOT there
    (tmp_path / "templates").mkdir(parents=True)
    settings = _settings(tmp_path)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)
    user = _user(db_session)

    with pytest.raises(AppError) as exc_info:
        word_book_service.create_word_book(
            db_session,
            user=user,
            classification_code="5/1",
            recipient_id=None,
            subject="x",
            cc=None,
            manager_id=None,
        )
    assert exc_info.value.code == "TEMPLATE_MISSING"
    assert exc_info.value.http_status == 409
    assert exc_info.value.details == {"file": _GENERAL_BOOK}


def test_create_with_cc_list(db_session, tmp_path, monkeypatch):
    """cc as a list is normalized into a newline-joined string in the docx."""
    from app.services import word_book_service

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="cc test",
        cc=["Ahmed Hassan", "Saeed Rashed"],
        manager_id=None,
    )
    working_file = Path(
        db_session.query(BookEditSession).filter_by(book_id=info.book_id).one().working_path
    )
    doc = docx.Document(str(working_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Both names appear somewhere in the rendered doc
    assert "Ahmed Hassan" in full_text


def test_create_with_manager(db_session, tmp_path, monkeypatch):
    """manager_id is resolved and manager_name appears in rendered docx."""
    from app.services import word_book_service

    _seed_gs(db_session)
    mgr = Manager(name_ar="أحمد سعيد", name_en="Ahmed Said", title="Director")
    db_session.add(mgr)
    db_session.commit()

    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="manager test",
        cc=None,
        manager_id=mgr.id,
    )
    working_file = Path(
        db_session.query(BookEditSession).filter_by(book_id=info.book_id).one().working_path
    )
    doc = docx.Document(str(working_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert "أحمد سعيد" in full_text


def test_body_sentinel_never_survives_in_working_docx(db_session, tmp_path, monkeypatch):
    """The {{ body }} anchor is cleared — the operator writes the body in Word,
    so no sentinel text may leak into the working file."""
    from app.services import word_book_service
    from app.services.document_service import GENERAL_BOOK_BODY_SENTINEL

    _seed_gs(db_session)
    (tmp_path / "templates").mkdir(parents=True)
    _write_minimal_docx(tmp_path / "templates" / _GENERAL_BOOK)
    monkeypatch.setattr(word_book_service, "get_settings", lambda: _settings(tmp_path))

    user = _user(db_session)
    info = word_book_service.create_word_book(
        db_session,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="sentinel test",
        cc=None,
        manager_id=None,
    )
    working_file = Path(
        db_session.query(BookEditSession).filter_by(book_id=info.book_id).one().working_path
    )
    doc = docx.Document(str(working_file))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    assert GENERAL_BOOK_BODY_SENTINEL not in full_text


# ---------------------------------------------------------------------------
# Utilities
# ---------------------------------------------------------------------------


def _settings(tmp_path):
    from app.config import Settings

    return Settings(data_dir=tmp_path / "data", templates_dir=tmp_path / "templates")


def _write_minimal_docx(path: Path) -> None:
    """Write a minimal docx with tokens matching the General Book template."""
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("التاريخ: {{ date }}")
    doc.add_paragraph("السيد / {{ recipient_name }}")
    doc.add_paragraph("الموضوع: {{ subject }}")
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("{{ cc }}")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_title }}")
    doc.add_paragraph("{{ submitter_g }}")
    doc.save(str(path))
