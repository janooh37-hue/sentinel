# backend/tests/test_report_render.py
from pathlib import Path

from docx import Document

from app.core.docx_engine import DocxEngine
from app.services.document_service import GENERAL_BOOK_BODY_SENTINEL

TEMPLATE = Path("backend/templates/GSSG-GS_300-004_Report.docx")

ADVERSARIAL_BODY = """
<h2 style="color:#1F4E79;">أولاً: ملخص / Summary</h2>
<p>نص عربي مع English mid-sentence لاختبار الـ bidi. </p>
<ol><li>البند الأول</li><li>البند الثاني</li></ol>
<table><thead><tr><th>الرقم</th><th>الوصف</th></tr></thead>
<tbody><tr><td>1</td><td>الليوان A3</td></tr></tbody></table>
<div class="mce-pagebreak"></div>
<p style="text-align:center;">— نهاية / End —</p>
"""


def test_report_template_tokens_and_no_ref():
    assert TEMPLATE.exists(), "run scripts/build_report_template.py first"
    doc = Document(str(TEMPLATE))
    text = "\n".join(p.text for p in doc.paragraphs)
    # ref line is gone entirely
    assert "{{ ref }}" not in text
    assert "الرقم" not in text
    # body + author tokens present, author labelled and ordered name→title→sig
    assert "{{ body }}" in text
    assert "{{ date }}" in text
    assert "{{ recipient_name }}" in text
    assert "الموضوع" in text
    assert "الاسم: {{ manager_name }}" in text
    assert "المسمى الوظيفي: {{ manager_title }}" in text
    assert "التوقيع: {{ manager_sig }}" in text
    # closing formula present
    assert "وتفضلوا بقبول فائق الاحترام والتقدير" in text
    # name paragraph appears before the signature paragraph
    names = [p.text for p in doc.paragraphs]
    assert names.index("الاسم: {{ manager_name }}") < names.index("التوقيع: {{ manager_sig }}")


def test_report_render_end_to_end(tmp_path):
    data = {
        "date": "23-07-2026",
        "subject": "اختبار",
        "recipient_name": "مدير المركز",
        "body": GENERAL_BOOK_BODY_SENTINEL,
        "body_html": ADVERSARIAL_BODY,
        "manager_name": "مهند أل علي",
        "manager_title": "مسؤول وحدة الإرساليات",
        "cc": "",
        "submitter_g": "G-2001",
    }
    out = tmp_path / "report.docx"
    DocxEngine(TEMPLATE.parent).fill("Report", data, out)
    from docx import Document as _D

    doc = _D(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert GENERAL_BOOK_BODY_SENTINEL not in text  # sentinel replaced
    assert "الرقم" not in text
    assert "نهاية" in text  # page-2 content present
    assert any("الوصف" in c.text for t in doc.tables for r in t.rows for c in r.cells)
