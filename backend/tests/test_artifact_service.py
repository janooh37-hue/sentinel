"""Public behavior of generated-artifact production."""

from __future__ import annotations

import hashlib
from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from pathlib import Path

from docx import Document

from app.core.constants import TEMPLATE_FILES
from app.services import artifact_service


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def test_produce_from_docx_preserves_authoritative_source_when_pdf_is_not_requested(
    tmp_path: Path,
) -> None:
    source = tmp_path / "authored.docx"
    document = Document()
    document.add_paragraph("P6-AUTHORED-CONTENT-20260905")
    document.add_table(rows=1, cols=2).cell(0, 1).text = "نص مؤلف للمرحلة السادسة"
    document.save(source)
    source_hash = _sha256(source)

    result = artifact_service.produce_from_docx(
        source_path=source,
        destination=tmp_path / "finished.docx",
        convert_pdf=False,
        collision="exact",
    )

    assert result.docx_path.read_bytes() == source.read_bytes()
    assert _sha256(source) == source_hash
    assert result.conversion.status == "skipped"
    assert result.conversion.pdf_path is None
    assert result.created_paths == (result.docx_path,)


def test_suffix_production_reserves_distinct_paths_concurrently(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)

    def produce() -> Path:
        return artifact_service.produce_from_docx(
            source_path=source,
            destination=tmp_path / "record.docx",
            convert_pdf=False,
        ).docx_path

    with ThreadPoolExecutor(max_workers=2) as pool:
        paths = list(pool.map(lambda _index: produce(), range(2)))

    assert {path.name for path in paths} == {"record.docx", "record_1.docx"}


def test_converter_outcomes_do_not_hide_failure_policy(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)

    unavailable = artifact_service.produce_from_docx(
        source_path=source,
        destination=tmp_path / "unavailable.docx",
        converter=lambda _path: None,
        collision="exact",
    )

    def crash(_path: Path) -> Path | None:
        raise RuntimeError("WORD-CRASH")

    error = artifact_service.produce_from_docx(
        source_path=source,
        destination=tmp_path / "error.docx",
        converter=crash,
        collision="exact",
    )

    assert unavailable.conversion.status == "unavailable"
    assert error.conversion.status == "error"
    assert error.conversion.error == "RuntimeError: WORD-CRASH"
    assert unavailable.docx_path.is_file()
    assert error.docx_path.is_file()


def test_produce_from_template_uses_configured_root_and_copies_input_data(
    tmp_path: Path, monkeypatch
) -> None:
    template_root = tmp_path / "templates"
    template_root.mkdir()
    template = Document()
    template.add_paragraph("{{ subject }}")
    template.save(template_root / TEMPLATE_FILES["General Book"])
    data: dict[str, object] = {"subject": "Phase Six"}

    result = artifact_service.produce_from_template(
        template_id="General Book",
        data=data,
        destination=tmp_path / "rendered.docx",
        stamps=artifact_service.StampPlan(reference="P6/1", header_reference=True),
        convert_pdf=False,
        collision="exact",
        template_root=template_root,
    )

    assert data == {"subject": "Phase Six"}
    rendered = Document(result.docx_path)
    assert rendered.paragraphs[0].text == "Phase Six"
    assert "Ref: P6/1" in rendered.sections[0].header.paragraphs[0].text


def test_signature_failure_removes_only_new_destination(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)
    destination = tmp_path / "signed.docx"

    try:
        artifact_service.produce_from_docx(
            source_path=source,
            destination=destination,
            stamps=artifact_service.StampPlan(
                signature=artifact_service.SignatureStamp(
                    image_path=tmp_path / "missing.png",
                    anchor_names=("Phase Six Manager",),
                    size_mm=30,
                    boldness=0,
                )
            ),
            convert_pdf=False,
            collision="exact",
        )
    except artifact_service.ArtifactStampError as exc:
        assert exc.operation == "signature"
    else:
        raise AssertionError("missing signature must fail")

    assert source.is_file()
    assert not destination.exists()


def test_build_docx_filename_preserves_existing_literal_policy() -> None:
    assert (
        artifact_service.build_docx_filename(
            "General Book", "A Very Long Employee Name", datetime(2026, 9, 5, 16, 7)
        )
        == "GeneralBook_A_Very_Long_Employee_20260905_1607.docx"
    )
    assert (
        artifact_service.build_docx_filename("Unknown Form", "", datetime(2026, 9, 5, 16, 7))
        == "UnknownForm_General_20260905_1607.docx"
    )


def test_produce_from_template_accepts_sandboxed_library_docx(tmp_path: Path) -> None:
    library = tmp_path / "library.docx"
    template = Document()
    template.add_paragraph("{{ subject }}")
    template.save(library)

    result = artifact_service.produce_from_template(
        template_id="General Book",
        template_path=library,
        data={"subject": "Library source"},
        destination=tmp_path / "library-output.docx",
        convert_pdf=False,
        collision="exact",
    )

    assert Document(result.docx_path).paragraphs[0].text == "Library source"


def test_converter_partial_pdf_is_removed_while_docx_fallback_is_retained(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)

    def partial_then_crash(docx_path: Path) -> Path | None:
        docx_path.with_suffix(".pdf").write_bytes(b"partial")
        raise RuntimeError("conversion interrupted")

    result = artifact_service.produce_from_docx(
        source_path=source,
        destination=tmp_path / "output" / "record.docx",
        converter=partial_then_crash,
        collision="exact",
    )

    partial_pdf = result.docx_path.with_suffix(".pdf")
    assert result.conversion.status == "error"
    assert result.created_paths == (result.docx_path,)
    assert not partial_pdf.exists()
    assert source.is_file()


def test_converter_rejects_and_removes_truncated_pdf(tmp_path: Path) -> None:
    source = tmp_path / "source.docx"
    Document().save(source)

    def truncated(docx_path: Path) -> Path | None:
        pdf = docx_path.with_suffix(".pdf")
        pdf.write_bytes(b"%PDF-truncated")
        return pdf

    result = artifact_service.produce_from_docx(
        source_path=source,
        destination=tmp_path / "output.docx",
        converter=truncated,
        collision="exact",
    )

    assert result.conversion.status == "unavailable"
    assert result.conversion.pdf_path is None
    assert result.created_paths == (result.docx_path,)
    assert not result.docx_path.with_suffix(".pdf").exists()
