"""TDD: Task 6 — classifications list + word-session route + BookRead additions.

RED → implement → GREEN.
"""

from __future__ import annotations

import secrets
from pathlib import Path

import pytest
from fastapi.testclient import TestClient
from sqlalchemy import create_engine
from sqlalchemy.orm import Session, sessionmaker

from app.api.deps import get_current_user
from app.db import session as session_mod
from app.db.models import Base, BookCategory, User
from app.db.session import attach_sqlite_pragmas, get_db
from app.main import create_app
from app.services import perm_service

# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


@pytest.fixture()
def api_db(monkeypatch, tmp_path) -> Session:
    db_file = tmp_path / "routes.db"
    eng = create_engine(
        f"sqlite:///{db_file}",
        future=True,
        connect_args={"check_same_thread": False},
    )
    attach_sqlite_pragmas(eng, wal=False)
    Base.metadata.create_all(eng)
    TestSession = sessionmaker(bind=eng, autoflush=False, expire_on_commit=False, future=True)
    monkeypatch.setattr(session_mod, "engine", eng)
    monkeypatch.setattr(session_mod, "SessionLocal", TestSession)
    db = TestSession()
    perm_service.seed_role_defaults(db)
    yield db
    db.close()


def _make_user(db: Session, role: str = "operator", email: str | None = None) -> User:
    u = User(
        email=email or f"{secrets.token_hex(4)}@test.ae",
        password_hash="x",
        role=role,
        status="active",
    )
    db.add(u)
    db.commit()
    db.refresh(u)
    return u


def _client(db: Session, user: User, monkeypatch, tmp_path: Path) -> TestClient:
    """Build a test client wired to db + user, with settings pointing at tmp_path."""
    from app.config import Settings
    from app.services import word_book_service

    settings = Settings(
        data_dir=tmp_path / "data",
        templates_dir=tmp_path / "templates",
    )
    # One template for every General Book — classified or not
    (tmp_path / "templates").mkdir(parents=True, exist_ok=True)
    _write_minimal_docx(tmp_path / "templates" / "GSSG-GS_300-003_General_Book.docx")

    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)

    app = create_app()
    app.dependency_overrides[get_db] = lambda: db
    app.dependency_overrides[get_current_user] = lambda: user
    return TestClient(app, raise_server_exceptions=True)


def _write_minimal_docx(path: Path) -> None:
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("{{ subject }}")
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("{{ recipient_name }}")
    doc.add_paragraph("{{ cc }}")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_title }}")
    doc.add_paragraph("{{ submitter_g }}")
    doc.add_paragraph("{{ date }}")
    doc.save(str(path))


def _write_retokenizable_docx(path: Path) -> None:
    """Full token set accepted by retokenize_general_book (superset of _write_minimal_docx)."""
    import docx as _docx

    doc = _docx.Document()
    doc.add_paragraph("{%p if ref %}")
    doc.add_paragraph("الرقم: {{ ref }}")
    doc.add_paragraph("{%p endif %}")
    doc.add_paragraph("التاريخ: {{ date }}")
    doc.add_paragraph("السيد / {{ recipient_name }}")
    doc.add_paragraph("الموضوع: {{ subject }}")
    doc.add_paragraph("{{ body }}")
    doc.add_paragraph("{{ cc }}")
    doc.add_paragraph("{{ manager_name }}")
    doc.add_paragraph("{{ manager_title }}")
    doc.add_paragraph("{{ submitter_g }}")
    doc.save(str(path))


def _seed_gs(db: Session) -> None:
    if db.get(BookCategory, "GS") is None:
        db.add(BookCategory(id="GS", prefix="GS"))
        db.commit()


# ---------------------------------------------------------------------------
# GET /books/classifications
# ---------------------------------------------------------------------------


def test_classifications_returns_15_items(api_db, monkeypatch, tmp_path):
    user = _make_user(api_db)
    c = _client(api_db, user, monkeypatch, tmp_path)
    resp = c.get("/api/v1/books/classifications")
    assert resp.status_code == 200, resp.text
    body = resp.json()
    assert "items" in body
    assert len(body["items"]) == 15


def test_classifications_shape(api_db, monkeypatch, tmp_path):
    user = _make_user(api_db)
    c = _client(api_db, user, monkeypatch, tmp_path)
    items = c.get("/api/v1/books/classifications").json()["items"]
    first = items[0]
    assert set(first.keys()) >= {"code", "tab", "name_ar", "name_en", "unit_ar"}


def test_classifications_requires_auth(api_db):
    """No session cookie → get_current_user raises 401."""
    app = create_app()
    app.dependency_overrides[get_db] = lambda: api_db
    # No get_current_user override — real dependency runs, no cookie → 401.
    c = TestClient(app, raise_server_exceptions=True)
    resp = c.get("/api/v1/books/classifications")
    assert resp.status_code == 401


# ---------------------------------------------------------------------------
# POST /books/word-sessions
# ---------------------------------------------------------------------------


def test_post_word_session_returns_201(api_db, monkeypatch, tmp_path):
    _seed_gs(api_db)
    user = _make_user(api_db, role="manager")
    c = _client(api_db, user, monkeypatch, tmp_path)
    resp = c.post(
        "/api/v1/books/word-sessions",
        json={"classification_code": "5/1", "subject": "Security permit test"},
    )
    assert resp.status_code == 201, resp.text
    body = resp.json()
    assert body["word_url"].startswith("ms-word:ofe|u|")
    assert body["book_id"] is not None
    assert body["ref_number"]
    assert body["token"]
    assert body["filename"].endswith(".docx")
    assert body["dav_url"]


def test_post_word_session_no_classification_rejected(api_db, monkeypatch, tmp_path):
    """Every book's ref comes from the classified register — no classification,
    no book (422 CLASSIFICATION_REQUIRED)."""
    _seed_gs(api_db)
    user = _make_user(api_db, role="manager")
    c = _client(api_db, user, monkeypatch, tmp_path)
    resp = c.post(
        "/api/v1/books/word-sessions",
        json={"classification_code": None, "subject": "Plain book"},
    )
    assert resp.status_code == 422, resp.text
    assert resp.json()["error"]["code"] == "CLASSIFICATION_REQUIRED"


def test_post_word_session_requires_books_manage(api_db, monkeypatch, tmp_path):
    """Plain operator (no books.manage) → 403."""
    _seed_gs(api_db)
    user = _make_user(api_db, role="operator")
    c = _client(api_db, user, monkeypatch, tmp_path)
    resp = c.post(
        "/api/v1/books/word-sessions",
        json={"classification_code": "5/1", "subject": "Should fail"},
    )
    assert resp.status_code == 403, resp.text


# ---------------------------------------------------------------------------
# BookRead additions: is_draft + classification_code + edit_session
# ---------------------------------------------------------------------------


def test_book_read_draft_fields(api_db, monkeypatch, tmp_path):
    """New word-session book has is_draft=True, classification_code set, edit_session set."""
    _seed_gs(api_db)
    user = _make_user(api_db, role="manager")
    c = _client(api_db, user, monkeypatch, tmp_path)

    # Create a word-session book
    create_resp = c.post(
        "/api/v1/books/word-sessions",
        json={"classification_code": "5/1", "subject": "Draft test"},
    )
    assert create_resp.status_code == 201, create_resp.text
    book_id = create_resp.json()["book_id"]

    # Fetch via GET /books/{id}
    # Need books.view permission — give user admin role which has all caps
    admin_user = _make_user(api_db, role="admin", email="admin@test.ae")
    c2 = _client(api_db, admin_user, monkeypatch, tmp_path)
    detail = c2.get(f"/api/v1/books/{book_id}").json()

    assert detail["is_draft"] is True
    assert detail["classification_code"] == "5/1"
    assert detail["voided_at"] is None
    assert detail["edit_session"] is not None
    assert detail["edit_session"]["user_id"] == user.id
    assert detail["edit_session"]["state"] == "active"


def test_list_books_batches_edit_sessions(api_db, monkeypatch, tmp_path):
    """GET /books returns correct is_draft + edit_session for each row via batch load.

    Two word-session books created by the same user (both get active sessions).
    List must show edit_session populated for both, is_draft True for both.
    """
    _seed_gs(api_db)
    admin_user = _make_user(api_db, role="admin", email="admin_batch@test.ae")
    c = _client(api_db, admin_user, monkeypatch, tmp_path)

    ids = []
    for code, subject in (("3/1", "Batch book A"), ("5/1", "Batch book B")):
        r = c.post(
            "/api/v1/books/word-sessions",
            json={"classification_code": code, "subject": subject},
        )
        assert r.status_code == 201, r.text
        ids.append(r.json()["book_id"])

    resp = c.get("/api/v1/books")
    assert resp.status_code == 200, resp.text
    items = {item["id"]: item for item in resp.json()["items"]}

    for book_id in ids:
        assert book_id in items, f"book {book_id} missing from list"
        row = items[book_id]
        assert row["is_draft"] is True
        assert row["edit_session"] is not None
        assert row["edit_session"]["state"] == "active"


# ---------------------------------------------------------------------------
# Fixtures for template routes
# ---------------------------------------------------------------------------


@pytest.fixture()
def finished_word_book(api_db, monkeypatch, tmp_path):
    """A finished General Book (BookVersion + Document on disk) via api_db."""
    # Fixture user != request user: save-as-template deliberately does not check book ownership (books.manage is the gate).
    from datetime import UTC, datetime

    from app.config import Settings
    from app.db.models import BookEditSession, User
    from app.services import book_template_service, word_book_service

    _seed_gs(api_db)

    settings = Settings(
        data_dir=tmp_path / "data",
        templates_dir=tmp_path / "templates",
    )
    (tmp_path / "templates").mkdir(parents=True, exist_ok=True)
    # Use the full token set required by retokenize_general_book (needs {{ date }})
    _write_retokenizable_docx(tmp_path / "templates" / "GSSG-GS_300-003_General_Book.docx")
    monkeypatch.setattr(word_book_service, "get_settings", lambda: settings)

    # Point template library at an isolated tmp dir
    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    user = User(email="tpl_route@test.ae", password_hash="x", status="active")
    api_db.add(user)
    api_db.commit()
    api_db.refresh(user)

    info = word_book_service.create_word_book(
        api_db,
        user=user,
        classification_code="5/1",
        recipient_id=None,
        subject="Template route test book",
        cc=None,
        manager_id=None,
    )
    edit_session = api_db.query(BookEditSession).filter_by(book_id=info.book_id).one()
    edit_session.last_put_at = datetime.now(UTC).replace(tzinfo=None)
    api_db.commit()
    return word_book_service.finish_word_session(api_db, user=user, book_id=info.book_id)


def _admin_client(api_db, monkeypatch, tmp_path):
    """books.manage-capable client (admin role)."""
    user = _make_user(api_db, role="admin", email=f"adm_{id(tmp_path)}@test.ae")
    return _client(api_db, user, monkeypatch, tmp_path), user


def _plain_client(api_db, monkeypatch, tmp_path):
    """No books.manage (operator role)."""
    user = _make_user(api_db, role="operator", email=f"op_{id(tmp_path)}@test.ae")
    return _client(api_db, user, monkeypatch, tmp_path), user


# ---------------------------------------------------------------------------
# GET /books/word-templates
# ---------------------------------------------------------------------------


def test_list_word_templates_empty(api_db, monkeypatch, tmp_path):
    from app.services import book_template_service

    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates")
    assert r.status_code == 200
    assert r.json() == []


def test_list_word_templates_requires_books_manage(api_db, monkeypatch, tmp_path):
    from app.services import book_template_service

    tpl_lib = tmp_path / "tpl_lib"
    tpl_lib.mkdir(parents=True, exist_ok=True)
    monkeypatch.setattr(book_template_service, "templates_dir", lambda: tpl_lib)

    c, _ = _plain_client(api_db, monkeypatch, tmp_path)
    r = c.get("/api/v1/books/word-templates")
    assert r.status_code == 403


# ---------------------------------------------------------------------------
# POST /books/{book_id}/save-as-template
# ---------------------------------------------------------------------------


def test_save_as_template_bad_name_422(api_db, monkeypatch, tmp_path, finished_word_book):
    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.post(
        f"/api/v1/books/{finished_word_book.id}/save-as-template",
        json={"name": "../evil"},
    )
    assert r.status_code == 422
    assert r.json()["error"]["code"] == "TEMPLATE_BAD_NAME"


def test_save_as_template_and_list(api_db, monkeypatch, tmp_path, finished_word_book):
    c, _ = _admin_client(api_db, monkeypatch, tmp_path)
    r = c.post(
        f"/api/v1/books/{finished_word_book.id}/save-as-template",
        json={"name": "قالب المسار"},
    )
    assert r.status_code == 201
    assert r.json()["name"] == "قالب المسار.docx"
    listed = c.get("/api/v1/books/word-templates").json()
    assert [t["name"] for t in listed] == ["قالب المسار.docx"]
