# backend/tests/test_arabic_rtl_table_spacing.py
from docx import Document
from docx.shared import Pt

from app.core.arabic_rtl import html_to_docx


def test_table_cell_paragraphs_have_zero_spacing():
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx("<p>before</p><table><tr><td>A</td><td>B</td></tr></table>", p)

    assert len(doc.tables) == 1
    cell_para = doc.tables[0].rows[0].cells[0].paragraphs[0]
    assert cell_para.paragraph_format.space_before == Pt(0)
    assert cell_para.paragraph_format.space_after == Pt(0)
    assert cell_para.paragraph_format.line_spacing == 1.0


def test_narrative_paragraph_spacing_untouched():
    doc = Document()
    p = doc.add_paragraph()
    html_to_docx("<p>before</p><table><tr><td>A</td></tr></table>", p)
    # The first block reuses the passed-in paragraph and must NOT be zeroed.
    assert p.paragraph_format.space_after is None
