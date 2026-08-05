"""The Inmate Conduct Violations template renders: row loop, bullets, tokens."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from app.core.constants import TEMPLATE_FILES
from app.core.docx_render import render

TEMPLATE_ID = "Inmate Conduct Violations"
TEMPLATES_DIR = Path(__file__).resolve().parents[1] / "templates"


def _table_texts(doc: Document) -> list[list[str]]:
    out: list[list[str]] = []
    for row in doc.tables[1].rows:
        seen: set[int] = set()
        cells: list[str] = []
        for c in row.cells:
            if id(c._tc) in seen:
                continue
            seen.add(id(c._tc))
            cells.append(c.text.strip())
        out.append(cells)
    return out


def _footer_text(doc: Document) -> str:
    """Join every cell of the first-page footer (word/footer3.xml)."""
    footer = doc.sections[0].first_page_footer
    return "\n".join(c.text for t in footer.tables for r in t.rows for c in r.cells)


@pytest.fixture
def rendered(tmp_path: Path) -> Document:
    src = TEMPLATES_DIR / TEMPLATE_FILES[TEMPLATE_ID]
    out = tmp_path / "out.docx"
    render(
        src,
        {
            "today": "05/08/2026",
            "now_time": "12:43 م",
            "inmates": [
                {
                    "name": "محمد سالم ياسر",
                    "nationality": "الامارات",
                    "wing": "1A",
                    "uid": "159809450",
                    "holding_no": "1565118",
                },
                {
                    "name": "خالد عبدالله",
                    "nationality": "مصر",
                    "wing": "3B",
                    "uid": "778112",
                    "holding_no": "990211",
                },
            ],
            "violation_details": "قام النزيل بالاعتداء داخل الليوان",
            "actions": ["تم ابلاغ مدير فرع شؤون النزلاء", "تم نقل النزيل الى قسم B وتقييده"],
            "reporter_name": "عبدالله سيف المنصوري",
            "reporter_g": "G-2001",
            "manager_name": "ناصر فاضل الساعدي",
            "submitter_g": "G-0312",
        },
        out,
    )
    return Document(str(out))


def test_registered_in_template_files() -> None:
    assert TEMPLATE_ID in TEMPLATE_FILES
    assert (TEMPLATES_DIR / TEMPLATE_FILES[TEMPLATE_ID]).exists()


def test_one_row_per_inmate(rendered: Document) -> None:
    rows = _table_texts(rendered)
    # R0 = date/day/time, R1 = column headers, then one row per inmate.
    assert rows[2][0] == "1"
    assert rows[2][1] == "محمد سالم ياسر"
    assert rows[2][3] == "1A"
    assert rows[3][0] == "2"
    assert rows[3][1] == "خالد عبدالله"
    assert rows[3][5] == "990211"


def test_date_day_time_filled(rendered: Document) -> None:
    header = rendered.tables[1].rows[0].cells
    joined = " ".join(c.text for c in header)
    assert "05/08/2026" in joined
    assert "الأربعاء" in joined  # weekday_ar derived from `today`
    assert "12:43 م" in joined
    assert "Auto_edit" not in joined  # placeholder text is gone


def test_actions_render_one_bullet_each(rendered: Document) -> None:
    body = "\n".join(p.text for p in rendered.tables[1].rows[8].cells[0].paragraphs)
    assert "تم ابلاغ مدير فرع شؤون النزلاء" in body
    assert "تم نقل النزيل الى قسم B وتقييده" in body
    # The unchosen action must NOT print.
    assert "تم كتابة مخالفة مسلكية" not in body


def test_reporter_and_manager_tokens(rendered: Document) -> None:
    whole = "\n".join(p.text for p in rendered.paragraphs)
    whole += "\n".join(c.text for r in rendered.tables[1].rows for c in r.cells)
    assert "عبدالله سيف المنصوري" in whole
    assert "G-2001" in whole
    assert "ناصر فاضل الساعدي" in whole


def test_no_malformed_tokens_survive(rendered: Document) -> None:
    whole = "\n".join(p.text for p in rendered.paragraphs)
    whole += "\n".join(c.text for r in rendered.tables[1].rows for c in r.cells)
    for bad in ("{{", "}}", "{%", "MANAGER-SIGN", "persionar"):
        assert bad not in whole


def test_footer_renders_submitter_g(rendered: Document) -> None:
    """The inherited first-page footer originally carried a non-framework
    token, {{ submitter_id }} — nothing ever filled it. Task 3's build script
    rewrote it to the canonical {{ submitter_g }}; guard the render end to
    end so a future template regeneration can't silently reintroduce the
    stray token or drop the fix.
    """
    footer = _footer_text(rendered)
    assert "G-0312" in footer
    assert "submitter_id" not in footer
    # The sibling paragraph in the same footer cell must survive the rewrite.
    assert "www.gss-group.net" in footer
