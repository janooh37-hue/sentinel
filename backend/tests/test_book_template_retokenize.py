"""Retokenize surgery: a finished General Book docx becomes a library
template with exactly three live tokens; all foreign Jinja is inert."""

from pathlib import Path

import pytest
from docx import Document

from app.core.book_template_retokenize import (
    retokenize_general_book,
    validate_book_template,
)
from app.core.book_text import docx_to_text
from app.core.docx_render import render


def _finished_book(tmp_path: Path, *, ref_line: bool = True, spacing: str = "") -> Path:
    """Minimal stand-in for a finished book: date + optional ref + body."""
    p = tmp_path / "book.docx"
    doc = Document()
    if ref_line:
        doc.add_paragraph(f"الرقم:{spacing}1/{spacing}5{spacing}/GSSG/{spacing}140")
    doc.add_paragraph("التاريخ: 13/07/2026")
    doc.add_paragraph("السيد / مدير الإدارة المحترم")
    doc.add_paragraph("الموضوع: التصاريح الأمنية بتاريخ 01/07/2026")
    doc.add_paragraph("نص الكتاب هنا")
    doc.save(str(p))
    return p


def _rendered_text(tpl: Path, tmp_path: Path, **data) -> str:
    out = tmp_path / "rendered.docx"
    render(tpl, data, out, sandboxed=True)
    return docx_to_text(out)


def test_ref_and_date_retokenized(tmp_path):
    p = _finished_book(tmp_path)
    retokenize_general_book(p)
    text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
    assert "الرقم: 9/9/GSSG/999" in text
    assert "التاريخ: 31-12-2099" in text
    assert "140" not in text  # old baked ref gone


def test_legacy_spacing_handled(tmp_path):
    p = _finished_book(tmp_path, spacing=" ")
    retokenize_general_book(p)
    text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
    assert "الرقم: 9/9/GSSG/999" in text


def test_missing_ref_line_inserted_above_date(tmp_path):
    p = _finished_book(tmp_path, ref_line=False)
    retokenize_general_book(p)
    text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
    assert text.index("الرقم: 9/9/GSSG/999") < text.index("التاريخ:")


def test_prose_date_untouched(tmp_path):
    p = _finished_book(tmp_path)
    retokenize_general_book(p)
    text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
    assert "بتاريخ 01/07/2026" in text  # date inside الموضوع prose survives


def test_foreign_jinja_neutralized(tmp_path):
    p = tmp_path / "book.docx"
    doc = Document()
    doc.add_paragraph("التاريخ: 13/07/2026")
    doc.add_paragraph("خصم {{ 7*7 }} بالمئة {% if x %}شرط{% endif %}")
    doc.save(str(p))
    retokenize_general_book(p)
    text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
    assert "49" not in text  # never executed
    assert "7*7" in text  # visible text preserved
    assert "شرط" in text  # {% if %} inert, content kept literal


def test_split_delimiter_fails_closed(tmp_path):
    """Known ceiling: per-w:t neutralization misses a Jinja delimiter split
    across two runs (run1 ends '{', run2 starts '{') — docxtpl's patch_xml can
    reassemble it at render time. Save-time validation is the fail-closed
    backstop (verified experimentally in review): the template is either
    rejected or renders without executing. The invariant is "never executes",
    not a specific failure mode."""
    p = tmp_path / "book.docx"
    doc = Document()
    doc.add_paragraph("التاريخ: 13/07/2026")
    split = doc.add_paragraph()
    split.add_run("خصم {")
    split.add_run("{ 7*7 }} بالمئة")
    doc.save(str(p))
    retokenize_general_book(p)
    try:
        validate_book_template(p)
    except ValueError:
        pass  # rejected at save time — fail-closed holds
    else:
        text = _rendered_text(p, tmp_path, ref="9/9/GSSG/999", date="31-12-2099")
        assert "49" not in text  # survived validation, but never executed


def test_ref_run_marked_ltr(tmp_path):
    p = _finished_book(tmp_path)
    retokenize_general_book(p)
    doc = Document(str(p))
    ref_para = next(pp for pp in doc.paragraphs if "{{ ref }}" in pp.text)
    run = next(r for r in ref_para.runs if "{{ ref }}" in r.text)
    assert run.font.rtl is False


def test_date_token_run_marked_ltr(tmp_path):
    p = _finished_book(tmp_path)
    retokenize_general_book(p)
    doc = Document(str(p))
    date_para = next(pp for pp in doc.paragraphs if "{{ date }}" in pp.text)
    token_run = next(r for r in date_para.runs if "{{ date }}" in r.text)
    assert token_run.font.rtl is False


def test_validate_accepts_good_template(tmp_path):
    p = _finished_book(tmp_path)
    retokenize_general_book(p)
    validate_book_template(p)  # no raise


def test_validate_rejects_unretokenized_doc(tmp_path):
    p = _finished_book(tmp_path)
    with pytest.raises(ValueError):
        validate_book_template(p)  # no tokens → dummy values never render
