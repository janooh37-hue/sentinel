from docx import Document
from docx.oxml.ns import qn

from app.core.book_table import normalize_data_table
from app.core.book_text import docx_to_text
from app.core.docx_render import render


def _mk(tmp_path, headers, data_rows):
    p = tmp_path / "t.docx"
    doc = Document()
    doc.add_paragraph("التاريخ: 01/01/2026")
    t = doc.add_table(rows=1 + len(data_rows), cols=len(headers))
    for i, h in enumerate(headers):
        t.cell(0, i).text = h
    for r, row in enumerate(data_rows, start=1):
        for c, v in enumerate(row):
            t.cell(r, c).text = v
    doc.save(str(p))
    return doc, p


def test_injects_loop_tokens_and_strips_pii(tmp_path):
    doc, p = _mk(tmp_path, ["الاسم", "الرقم"], [["أحمد", "101"], ["خالد", "202"]])
    normalize_data_table(doc)
    doc.save(str(p))
    text = docx_to_text(p)
    assert "{%tr for row in table_rows %}" in text
    assert "{{ row.c0 }}" in text and "{{ row.c1 }}" in text
    assert "{%tr endfor %}" in text
    assert "أحمد" not in text and "خالد" not in text


def test_idempotent(tmp_path):
    doc, p = _mk(tmp_path, ["أ", "ب", "ج"], [["x", "y", "z"]])
    normalize_data_table(doc)
    doc.save(str(p))
    xml1 = doc.element.body.xml
    doc2 = Document(str(p))
    normalize_data_table(doc2)
    doc2.save(str(p))
    xml2 = doc2.element.body.xml
    assert xml1 == xml2


def test_noop_no_table(tmp_path):
    doc = Document()
    doc.add_paragraph("التاريخ: 01/01/2026")
    before = doc.element.body.xml
    normalize_data_table(doc)
    assert doc.element.body.xml == before


def test_noop_two_tables(tmp_path):
    doc = Document()
    for _ in range(2):
        t = doc.add_table(rows=2, cols=2)
        t.cell(0, 0).text = "أ"
    before = doc.element.body.xml
    normalize_data_table(doc)
    assert doc.element.body.xml == before


def test_tblheader_set(tmp_path):
    doc, _ = _mk(tmp_path, ["الاسم", "الرقم"], [["أحمد", "101"]])
    normalize_data_table(doc)
    hdr = doc.element.body.findall(qn("w:tbl"))[0].findall(qn("w:tr"))[0]
    trPr = hdr.find(qn("w:trPr"))
    assert trPr is not None and trPr.find(qn("w:tblHeader")) is not None


def test_renders_table_rows(tmp_path):
    doc, p = _mk(tmp_path, ["الاسم", "الجهة"], [["علي", "الأمن"]])
    normalize_data_table(doc)
    doc.save(str(p))
    out = tmp_path / "r.docx"
    render(
        p,
        {
            "ref": "1/5/141",
            "date": "20-07-2026",
            "submitter_g": "G-1",
            "table_rows": [{"c0": "محمد", "c1": "المعلومات"}, {"c0": "ياسر", "c1": "الدعم"}],
        },
        out,
        sandboxed=True,
    )
    text = docx_to_text(out)
    assert "محمد" in text and "ياسر" in text and "المعلومات" in text
    assert "{%tr" not in text
