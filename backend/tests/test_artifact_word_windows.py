"""Opt-in real Microsoft Word acceptance for the Phase 6 artifact boundary."""

from __future__ import annotations

import asyncio
import hashlib
import json
import os
import shutil
import sys
import zipfile
from pathlib import Path

import fitz
import pytest
from docx import Document as DocxFile
from docx.enum.text import WD_BREAK
from PIL import Image, ImageDraw
from sqlalchemy import select
from starlette.requests import Request

from app.config import get_settings
from app.db.models import BookCategory, BookEditSession, BookVersion, Document, Employee, User
from app.db.session import SessionLocal
from app.services import document_service, perm_service, word_book_service


def _sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def _real_dispatch_ex_converter(source: Path) -> Path | None:
    """Delegate to the repository's actual fresh-DispatchEx production method."""
    from app.core.pdf_chain import PdfChain

    destination = source.with_suffix(".pdf")
    PdfChain()._via_win32com(source.resolve(), destination.resolve())
    return destination


def _put_request(body: bytes) -> Request:
    sent = False

    async def receive() -> dict[str, object]:
        nonlocal sent
        if sent:
            return {"type": "http.disconnect"}
        sent = True
        return {"type": "http.request", "body": body, "more_body": False}

    return Request(
        {"type": "http", "method": "PUT", "path": "/dav", "headers": []},
        receive,
    )


def _dav_put(db, session: BookEditSession, body: bytes) -> None:
    from app.api.dav import dav_handler

    response = asyncio.run(
        dav_handler(
            token=session.token,
            filename=Path(session.working_path).name,
            request=_put_request(body),
            db=db,
        )
    )
    assert response.status_code == 204


def _word_save(path: Path) -> int:
    import pythoncom
    import win32com.client
    import win32process

    pythoncom.CoInitialize()
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        _thread_id, process_id = win32process.GetWindowThreadProcessId(word.Hwnd)
        document = word.Documents.Open(str(path.resolve()), ReadOnly=False)
        document.Save()
        document.Close(False)
        document = None
        word.Quit()
        word = None
        return int(process_id)
    finally:
        if document is not None:
            document.Close(False)
        if word is not None:
            word.Quit()
        pythoncom.CoUninitialize()


def _manifest(path: Path) -> dict[str, object]:
    document = DocxFile(path)
    with zipfile.ZipFile(path) as archive:
        media = {
            name: hashlib.sha256(archive.read(name)).hexdigest()
            for name in archive.namelist()
            if name.startswith("word/media/")
        }
        names = sorted(archive.namelist())
        xml = archive.read("word/document.xml")
    return {
        "sha256": _sha256(path),
        "paragraphs": [paragraph.text for paragraph in document.paragraphs],
        "tables": [
            [[cell.text for cell in row.cells] for row in table.rows] for table in document.tables
        ],
        "headers": [
            paragraph.text
            for section in document.sections
            for paragraph in section.header.paragraphs
        ],
        "footers": [
            paragraph.text
            for section in document.sections
            for paragraph in section.footer.paragraphs
        ],
        "zip_members": names,
        "media": media,
        "drawing_count": xml.count(b"<w:drawing"),
    }


def _copy_pdf_pages(pdf_path: Path, pages_dir: Path, prefix: str) -> None:
    pages_dir.mkdir(parents=True, exist_ok=True)
    with fitz.open(pdf_path) as pdf:
        assert pdf.page_count > 0
        for index, page in enumerate(pdf):
            page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False).save(
                pages_dir / f"{prefix}-{index + 1}.png"
            )


@pytest.mark.skipif(sys.platform != "win32", reason="requires Microsoft Word on Windows")
def test_phase6_windows_word_artifact_smoke() -> None:
    if os.environ.get("GSSG_RUN_WORD_ARTIFACT_SMOKE") != "1":
        pytest.skip("set GSSG_RUN_WORD_ARTIFACT_SMOKE=1 for the isolated Word gate")
    evidence_dir = Path(os.environ["GSSG_WORD_EVIDENCE_DIR"]).resolve(strict=True)
    settings = get_settings()
    assert settings.data_dir.resolve() != Path("C:/Users/Admin/sentinel/data").resolve()

    signature = evidence_dir / "synthetic-signature.png"
    image = Image.new("RGBA", (300, 100), (255, 255, 255, 0))
    draw = ImageDraw.Draw(image)
    draw.line([(20, 75), (80, 20), (145, 75), (220, 25), (280, 65)], fill="black", width=6)
    image.save(signature)
    embedded = evidence_dir / "synthetic-embedded.png"
    Image.new("RGB", (80, 40), "navy").save(embedded)

    word_pids: list[int] = []
    with SessionLocal() as db:
        perm_service.seed_role_defaults(db)
        for category in (BookCategory(id="GS", prefix="GS"), BookCategory(id="HR", prefix="HR")):
            if db.get(BookCategory, category.id) is None:
                db.add(category)
        employee = db.get(Employee, "P6-9001")
        if employee is None:
            employee = Employee(
                id="P6-9001", name_en="Phase Six Employee", name_ar="موظف المرحلة السادسة"
            )
            db.add(employee)
        user = db.scalar(select(User).where(User.email == "p6-word-smoke@example.invalid"))
        if user is None:
            user = User(
                email="p6-word-smoke@example.invalid",
                password_hash="synthetic",
                role="admin",
                status="active",
                display_name="Phase Six Operator",
                employee_id=employee.id,
            )
            db.add(user)
        db.commit()

        leave = document_service.generate_document(
            db,
            employee_id=employee.id,
            template_id="Leave Application Form",
            fields={
                "leave_type": "Sick Leave",
                "start_date": "05/09/2026",
                "end_date": "05/09/2026",
                "total_days": 1,
            },
            commit=True,
            current_user=user,
            converter=_real_dispatch_ex_converter,
        )
        leave_docx = leave.docx_path
        leave_pdf = leave.pdf_path
        assert leave_pdf is not None

        info = word_book_service.create_word_book(
            db,
            user=user,
            classification_code="5/1",
            recipient_id=None,
            subject="P6-AUTHORED-CONTENT-20260905",
            cc=None,
            manager_id=None,
        )
        session = db.scalar(
            select(BookEditSession).where(
                BookEditSession.book_id == info.book_id, BookEditSession.state == "active"
            )
        )
        assert session is not None
        working = Path(session.working_path)
        authored = DocxFile(working)
        english = authored.add_paragraph().add_run("P6-AUTHORED-CONTENT-20260905")
        english.bold = True
        arabic = authored.add_paragraph().add_run("نص مؤلف للمرحلة السادسة")
        arabic.italic = True
        table = authored.add_table(rows=1, cols=2)
        table.cell(0, 0).text = "P6-TABLE-A"
        table.cell(0, 1).text = "P6-TABLE-B"
        authored.add_picture(str(embedded))
        authored.sections[0].header.paragraphs[0].text += " P6-HEADER"
        authored.sections[0].footer.paragraphs[0].text += " P6-FOOTER"
        authored.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        authored.add_paragraph("P6-PREVIEW-OLD")
        authored.add_paragraph("")
        authored.add_paragraph("Phase Six Manager")
        authored.save(working)
        word_pids.append(_word_save(working))
        _dav_put(db, session, working.read_bytes())
        source_manifest = _manifest(working)
        shutil.copy2(working, evidence_dir / "authored-source.docx")

        preview_old = word_book_service.render_session_preview(
            db, book_id=info.book_id, converter=_real_dispatch_ex_converter
        )
        shutil.copy2(preview_old, evidence_dir / "preview-old.pdf")

        changed = DocxFile(working)
        changed.add_paragraph("P6-PREVIEW-NEW")
        changed.save(working)
        word_pids.append(_word_save(working))
        _dav_put(db, session, working.read_bytes())
        preview_new = word_book_service.render_session_preview(
            db, book_id=info.book_id, converter=_real_dispatch_ex_converter
        )
        shutil.copy2(preview_new, evidence_dir / "preview-new.pdf")

        book = word_book_service.finish_word_session(
            db, user=user, book_id=info.book_id, converter=_real_dispatch_ex_converter
        )
        version = db.scalar(
            select(BookVersion)
            .where(BookVersion.book_id == book.id)
            .order_by(BookVersion.version_no.desc())
        )
        assert version is not None and version.document_id is not None and version.fields == {}
        row = db.get(Document, version.document_id)
        assert row is not None and row.docx_path
        finished = Path(row.docx_path)
        if not finished.is_absolute():
            finished = settings.data_dir / finished
        finished_manifest = _manifest(finished)
        for key in ("tables", "headers", "footers", "media"):
            assert finished_manifest[key] == source_manifest[key]
        assert "P6-AUTHORED-CONTENT-20260905" in finished_manifest["paragraphs"]
        assert not working.exists()
        shutil.copy2(finished, evidence_dir / "authored-finished.docx")

        signed_value = document_service.render_signed_pdf(
            db,
            version=version,
            signer_signature_path=str(signature),
            signer_names=("Phase Six Manager", "مدير المرحلة السادسة"),
            converter=_real_dispatch_ex_converter,
        )
        signed = Path(signed_value)
        if not signed.is_absolute():
            signed = settings.data_dir / signed
        assert signed.suffix.lower() == ".pdf"

    shutil.copy2(leave_docx, evidence_dir / "leave.docx")
    shutil.copy2(leave_pdf, evidence_dir / "leave.pdf")
    shutil.copy2(signed, evidence_dir / "authored-signed.pdf")
    _copy_pdf_pages(leave_pdf, evidence_dir / "pages", "leave")
    _copy_pdf_pages(signed, evidence_dir / "pages", "signed")
    _copy_pdf_pages(evidence_dir / "preview-new.pdf", evidence_dir / "pages", "preview-new")
    (evidence_dir / "word-smoke.json").write_text(
        json.dumps(
            {
                "converter_adapter": "PdfChain._via_win32com",
                "converter_claim": "fresh DispatchEx only",
                "default_chain_verified": False,
                "process_pool_verified": False,
                "word_pids": word_pids,
                "leave_docx_sha256": _sha256(leave_docx),
                "leave_pdf_sha256": _sha256(leave_pdf),
                "authored_source_manifest": source_manifest,
                "finished_manifest": finished_manifest,
                "signed_pdf_sha256": _sha256(signed),
            },
            ensure_ascii=False,
            indent=2,
            sort_keys=True,
        ),
        encoding="utf-8",
    )
