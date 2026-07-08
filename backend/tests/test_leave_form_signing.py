# backend/tests/test_leave_form_signing.py
"""Regression tests for the Leave Permit (301-004) and Administrative Leave
(301-005) signer block:

  1. The manager signature floats *behind text* (zero layout height) on the
     signature line just ABOVE the manager name at the foot of the form — not
     inside the project-manager notes box, and not in the branch-manager
     التوقيع cell.
  2. On Admin Leave the branch-manager التوقيع cell is left blank (no image, no
     literal "x" placeholder) for a separate manual signature.
  3. The manager name + title block is LEFT-aligned and consistent.
  4. With NO manager selected, the standalone designation (DEFAULT_MANAGER_TITLE)
     is not left stranded at the foot of the form, and no signature is placed.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

from app.core.constants import DEFAULT_MANAGER_TITLE
from app.core.docx_engine import DocxEngine
from app.services.document_service import _TEMPLATES_DIR

_MANAGER = "سعيد راشد اليحيائي"
_TITLE = "مدير مشروع مركز الإصلاح والتأهيل – الوثبة 2"  # noqa: RUF001


def _make_sig_png(path: Path) -> Path:
    img = Image.new("RGBA", (400, 168), (255, 255, 255, 0))
    for x in range(40, 360):
        y = 20 + int((x - 40) * 120 / 320)
        for dy in (-1, 0, 1):
            img.putpixel((x, y + dy), (0, 0, 0, 255))
    img.save(path)
    return path


def _base() -> dict:
    return {
        "employee_id": "G3082",
        "employee_name_ar": "محمد أحمد",
        "date": "08/07/2026",
        "unit_branch": "مركز الوثبة",
        "leave_reason": "ظرف عائلي",
        "duration_hours": "4",
        "duration": "3 أيام",
        "leave_date_range": "08/07 - 10/07",
        "admin_leaves_this_month": "2",
    }


def _with_manager(sig: Path) -> dict:
    return {
        **_base(),
        "manager_name": _MANAGER,
        "manager_title": _TITLE,
        "manager_sig_path": str(sig),
        "_sig_size_mm": 45,
    }


def _render(form: str, data: dict, tmp_path: Path):
    out = tmp_path / "form.docx"
    DocxEngine(_TEMPLATES_DIR).fill(form, data, out)
    return Document(str(out))


def _manager_block_paragraphs(doc):
    return [p for p in doc.paragraphs if p.text.strip() in (_MANAGER, _TITLE)]


def _is_left_aligned(paragraph) -> bool:
    if paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT:
        return False
    pPr = paragraph._p.find(qn("w:pPr"))
    return pPr is None or pPr.find(qn("w:bidi")) is None


def _para_above_name(doc):
    """The body paragraph directly above the manager-name paragraph (the
    signature-line gap where the float is anchored)."""
    body = list(doc.paragraphs)
    for i, p in enumerate(body):
        if p.text.strip() == _MANAGER and i > 0:
            return body[i - 1]
    return None


FORMS = ("Leave Permit Form", "Administrative Leave Form")


# --- signature lands above the name, floated behind text -------------------


def test_signature_floats_above_name_not_in_notes_box(tmp_path):
    for form in FORMS:
        doc = _render(form, _with_manager(_make_sig_png(tmp_path / "s.png")), tmp_path)
        anchor = _para_above_name(doc)
        assert anchor is not None, f"{form}: name paragraph not found"
        axml = anchor._p.xml
        assert "w:drawing" in axml, f"{form}: signature not on the line above the name"
        assert 'behindDoc="1"' in axml, f"{form}: signature not floated behind text"
        assert "wp:inline" not in axml, f"{form}: signature still inline"
        # The notes box (last table) must NOT hold the signature any more.
        assert "w:drawing" not in doc.tables[-1]._tbl.xml, f"{form}: sig still in notes box"


def test_manager_block_left_aligned(tmp_path):
    for form in FORMS:
        doc = _render(form, _with_manager(_make_sig_png(tmp_path / "s.png")), tmp_path)
        paras = _manager_block_paragraphs(doc)
        assert len(paras) == 2, f"{form}: expected name + title, got {len(paras)}"
        assert all(_is_left_aligned(p) for p in paras), f"{form}: block not left-aligned"


def test_admin_leave_branch_signature_cell_is_blank(tmp_path):
    doc = _render(
        "Administrative Leave Form", _with_manager(_make_sig_png(tmp_path / "s.png")), tmp_path
    )
    cell = doc.tables[2].rows[3].cells[1]
    assert "w:drawing" not in cell._tc.xml, "signature must not sit in the branch التوقيع cell"
    assert "x" not in cell.text.lower(), "literal 'x' placeholder not cleared"


# --- no manager selected → no stranded designation, no signature -----------


def test_no_manager_hides_designation_and_signature(tmp_path):
    for form in FORMS:
        doc = _render(form, _base(), tmp_path)  # no manager_* fields at all
        body_text = "\n".join(p.text for p in doc.paragraphs)
        assert DEFAULT_MANAGER_TITLE not in body_text, f"{form}: designation shown with no manager"
        assert "w:drawing" not in doc.element.body.xml, f"{form}: stray signature with no manager"


# --- manager name renders in Arabic (prefer_arabic wiring) -----------------


def test_leave_forms_use_arabic_manager_name(db_session, tmp_path, monkeypatch):
    """Leave Permit / Admin Leave must resolve the manager's Arabic name, while a
    non-leave form keeps the English-first default (scoping proof)."""
    from app.config import Settings
    from app.db.models import Manager
    from app.services import document_service
    from app.services.document_service import _build_template_data

    settings = Settings(data_dir=tmp_path, templates_dir=_TEMPLATES_DIR)
    monkeypatch.setattr(document_service, "get_settings", lambda: settings)

    mgr = Manager(name_en="Saeed Rashed", name_ar="سعيد راشد اليحيائي", title=None)
    db_session.add(mgr)
    db_session.flush()

    def _name(template_id: str) -> str:
        data = _build_template_data(
            db_session,
            template_id=template_id,
            employee=None,
            employee_id=None,
            fields={},
            manager_id=mgr.id,
            submitter_id=None,
            embed_signature=None,
            current_user=None,
        )
        return data["manager_name"]

    assert _name("Leave Permit Form") == "سعيد راشد اليحيائي"
    assert _name("Administrative Leave Form") == "سعيد راشد اليحيائي"
    # A form outside the Arabic-preferring set keeps English-first behaviour.
    assert _name("Material Request Form") == "Saeed Rashed"
