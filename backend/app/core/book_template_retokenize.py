"""Turn a finished General Book docx into a library boilerplate template.

The per-book fields are (re)injected as tokens -- ``{{ ref }}``, ``{{ date }}``,
``{{ submitter_g }}``, ``{{ recipient_name }}``, ``{{ subject }}``, ``{{ cc }}``
-- so a template is a reusable shell, not a snapshot frozen to the addressee and
subject of the book it was saved from. ALL pre-existing Jinja delimiters in the document
are neutralized (a zero-width space inside each delimiter) so operator-typed
text can never execute server-side (SSTI defense; stored templates are
untrusted). Validation test-renders under StrictUndefined + sandbox and
fails closed.
"""

from __future__ import annotations

import copy
import re
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.text.paragraph import Paragraph

from app.core.book_table import normalize_data_table
from app.core.book_text import docx_to_text
from app.core.docx_render import render

_ZWSP = "​"  # zero-width space — invisible, breaks Jinja delimiters
_JINJA_DELIM = re.compile(r"\{\{|\}\}|\{%|%\}|\{#|#\}")
# The Aztec stamp's anchor carries this fixed relativeHeight (see
# _docx_helpers.insert_floating_image_in_header) — the letterhead images do
# not, so this selector can never remove the letterhead.
_AZTEC_RELATIVE_HEIGHT = "251670000"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"

_REF_LABEL = re.compile(r"^\s*الرقم\s*[:：]")  # noqa: RUF001 — full-width colon is a legitimate Arabic-text variant
_DATE_LABEL = re.compile(r"^\s*التاريخ\s*[:：]")  # noqa: RUF001 — full-width colon is a legitimate Arabic-text variant
_SUBJECT_LABEL = re.compile(r"^\s*الموضوع\s*[:：]")  # noqa: RUF001 — full-width colon is a legitimate Arabic-text variant
# The paper's addressee line: «السيد \ {name} المحترم» (the separator is a
# backslash on the current template, a slash on the older hand-typed books).
_ADDRESSEE = re.compile(r"^\s*السيد\s*([\\/])")
_G_NUMBER = re.compile(r"\bG[-\s]?\d{1,6}\b")

_DUMMY = {
    "ref": "9/9/9999",
    "date": "31-12-2099",
    "submitter_g": "G-9999",
    # Truthy so the {%p if %} guards render — asserted only via the body check.
    "recipient_name": "DUMMY_RECIPIENT",
    "subject": "DUMMY_SUBJECT",
    "cc": "DUMMY_CC",
}

# w:t namespace for walking raw XML runs
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_W_T = f"{{{_W_NS}}}t"


def _neutralize_wt(text: str) -> str:
    """Insert ZWSP inside every Jinja delimiter found in *text*."""
    return _JINJA_DELIM.sub(lambda m: m.group(0)[0] + _ZWSP + m.group(0)[1], text)


def _neutralize_part_runs(container: Any) -> None:
    """Walk all paragraphs (and table cells recursively) in *container*,
    neutralizing Jinja delimiters at the w:t element level so cross-run
    delimiters are also caught.

    python-docx `para.runs` groups consecutive w:r elements but may miss
    delimiters split across run boundaries. Walking w:t directly covers both
    cases: each w:t text node is replaced atomically.
    """
    for para in container.paragraphs:
        # Fast path: skip paragraphs without any delimiter characters
        if "{{" not in para.text and "{%" not in para.text and "{#" not in para.text:
            continue
        for wt in para._p.iter(_W_T):
            if wt.text and _JINJA_DELIM.search(wt.text):
                wt.text = _neutralize_wt(wt.text)
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                _neutralize_part_runs(cell)


def _clear_runs(para: Paragraph) -> None:
    for r in list(para.runs):
        r._element.getparent().remove(r._element)


def _first_run_style(para: Paragraph) -> Any | None:
    return para.runs[0] if para.runs else None


def _guard_body(para: Paragraph) -> str:
    """Paragraph text with ZWSP and Jinja delimiters stripped — "p if ref" for a
    (possibly neutralized) ``{%p if ref %}`` guard."""
    return _JINJA_DELIM.sub("", (para.text or "").replace(_ZWSP, "")).strip()


def _guard_para(para: Paragraph, directive: str, *, after: bool) -> None:
    """Ensure a guard paragraph carrying *directive* sits immediately before
    (or after) *para*. Reuses a guard that is already there — possibly with
    ZWSP-broken delimiters from the neutralize pass — so a second retokenize
    call doesn't accumulate duplicates."""
    body = _JINJA_DELIM.sub("", directive).strip()
    sibling = para._p.getnext() if after else para._p.getprevious()
    if sibling is not None and _guard_body(Paragraph(sibling, para._parent)) == body:
        guard = Paragraph(sibling, para._parent)
    else:
        clone = copy.deepcopy(para._p)
        (para._p.addnext if after else para._p.addprevious)(clone)
        guard = Paragraph(clone, para._parent)
    _clear_runs(guard)
    guard.add_run(directive)


def _write_ref_block(anchor: Paragraph, *, replace: bool, style_src: Any | None = None) -> None:
    """Write {%p if ref %} / الرقم: {{ ref }} / {%p endif %} at *anchor*.

    replace=True: anchor IS the old الرقم paragraph (reuse it for the label
    line, keeping its formatting). replace=False: insert all three before
    anchor (the التاريخ paragraph).

    style_src overrides the run-style source; default is anchor itself."""
    src = _first_run_style(style_src if style_src is not None else anchor)

    def styled(run: Any) -> Any:
        if src is not None:
            run.font.name = src.font.name
            run.font.size = src.font.size
            run.font.bold = src.font.bold
        return run

    if replace:
        label_para = anchor
    else:
        new_p = copy.deepcopy(anchor._p)
        anchor._p.addprevious(new_p)
        label_para = Paragraph(new_p, anchor._parent)

    _guard_para(label_para, "{%p if ref %}", after=False)

    _clear_runs(label_para)
    styled(label_para.add_run("الرقم: "))
    # <w:rtl/> on the ref run — the EXACT encoding of the hand-typed legacy
    # books (verified by XML dump: their digit runs are RTL-marked). Word
    # then orders the segments right-to-left so the bumping serial reads
    # LAST on the line. Both no-mark and a forced <w:rtl w:val="0"/> made
    # Word lay the value as one LTR unit with the serial landing right next
    # to الرقم: (operator-reported twice).
    ref_run = styled(label_para.add_run("{{ ref }}"))
    ref_run.font.rtl = True

    _guard_para(label_para, "{%p endif %}", after=True)


def _retokenize_labeled_line(para: Paragraph, *parts: str) -> None:
    """Replace *para*'s runs with *parts*, all carrying the original first
    run's font. Token runs inherit the paragraph's RTL context (same rationale
    as the ref run — match the legacy books' natural bidi flow)."""
    src = _first_run_style(para)
    _clear_runs(para)
    for part in parts:
        run = para.add_run(part)
        if src is not None:
            run.font.name = src.font.name
            run.font.size = src.font.size
            run.font.bold = src.font.bold


def _strip_header_artifacts(doc: Any) -> None:
    """Remove the old Aztec anchor (by its unique relativeHeight) and any
    legacy English 'Ref:' stamp text from both header parts."""
    for section in doc.sections:
        for hdr in (section.header, section.first_page_header):
            for para in hdr.paragraphs:
                for anchor in para._p.findall(f".//{{{_WP_NS}}}anchor"):
                    if anchor.get("relativeHeight") == _AZTEC_RELATIVE_HEIGHT:
                        drawing = anchor.getparent()
                        drawing.getparent().remove(drawing)
                if para.text.strip().startswith("Ref:"):
                    _clear_runs(para)


def _retokenize_footers(doc: Any, submitter_g: str | None) -> None:
    """Both footers (footer2 is a synced copy of footer3): the baked G-number
    becomes {{ submitter_g }} so a new author's G renders at create. Hand-made
    templates (the Desktop imports) carry their own footers with NO G at all —
    then the token is INSERTED into the default footer so the author's G still
    renders on every book (9pt, matching the canonical footer's size)."""
    replaced = False
    for section in doc.sections:
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for para in footer.paragraphs:
                for run in para.runs:
                    if submitter_g and submitter_g in run.text:
                        run.text = run.text.replace(submitter_g, "{{ submitter_g }}")
                        replaced = True
                    elif _G_NUMBER.search(run.text):
                        run.text = _G_NUMBER.sub("{{ submitter_g }}", run.text, count=1)
                        replaced = True
    if not replaced:
        from docx.shared import Pt

        footer = doc.sections[0].footer
        if footer.paragraphs and not footer.paragraphs[-1].text.strip():
            para = footer.paragraphs[-1]
        else:
            para = footer.add_paragraph()
        run = para.add_run("{{ submitter_g }}")
        run.font.size = Pt(9)


def retokenize_general_book(docx_path: Path, *, submitter_g: str | None = None) -> None:
    doc = Document(str(docx_path))

    # 1. Neutralize FIRST — everything currently in the doc is untrusted.
    _neutralize_part_runs(doc)
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            _neutralize_part_runs(part)

    # 1b. Normalize a clean data table (if present) — AFTER neutralize so the
    # injected {%tr%}/{{ row.cN }} tokens are inserted fresh and never ZWSP-broken.
    # normalize_data_table handles all cases internally (no-op when no single
    # clean table; strips ZWSP-broken directive rows on re-run for idempotency).
    normalize_data_table(doc)

    # 2/3. Ref + date lines (first labeled body paragraph each; prose ignored).
    date_para = next((p for p in doc.paragraphs if _DATE_LABEL.match(p.text)), None)
    if date_para is None:
        raise ValueError("لا يحتوي المستند على سطر التاريخ — لا يمكن حفظه كقالب")
    ref_para = next((p for p in doc.paragraphs if _REF_LABEL.match(p.text)), None)
    if ref_para is not None:
        _write_ref_block(ref_para, replace=True, style_src=date_para)
    else:
        _write_ref_block(date_para, replace=False)
    _retokenize_labeled_line(date_para, "التاريخ: ", "{{ date }}")

    # 3b. Addressee / subject / CC — the source book's literal values would
    # otherwise be frozen into every book made from this template (the form's
    # pickers had nothing to fill). Subject and CC keep the base paper's
    # {%p if %} guard so an empty value hides the whole line, label and all.
    addressee = next((p for p in doc.paragraphs if _ADDRESSEE.match(p.text)), None)
    if addressee is not None:
        sep = _ADDRESSEE.match(addressee.text).group(1)  # type: ignore[union-attr]
        # The recipient picker is optional, so guard the line too — an unpicked
        # recipient would otherwise print a bare «السيد \ المحترم» with a hole
        # in the middle (the base paper's own behaviour, unguarded there).
        _guard_para(addressee, "{%p if recipient_name %}", after=False)
        _retokenize_labeled_line(addressee, f"السيد {sep} ", "{{ recipient_name }}", " المحترم ")
        _guard_para(addressee, "{%p endif %}", after=True)
    subject_para = next((p for p in doc.paragraphs if _SUBJECT_LABEL.match(p.text)), None)
    if subject_para is not None:
        _guard_para(subject_para, "{%p if subject %}", after=False)
        _retokenize_labeled_line(subject_para, "الموضوع: ", "{{ subject }}")
        _guard_para(subject_para, "{%p endif %}", after=True)
    # The CC line is post-processed into "• نسخة إلى: X" bullets, so match on
    # the label alone rather than anchoring to the paragraph start.
    cc_para = next((p for p in doc.paragraphs if "نسخة إلى" in (p.text or "")), None)
    if cc_para is not None:
        _guard_para(cc_para, "{%p if cc %}", after=False)
        _retokenize_labeled_line(cc_para, "نسخة إلى: ", "{{ cc }}")
        _guard_para(cc_para, "{%p endif %}", after=True)

    # 4. Footer G-number → token (both footers).
    _retokenize_footers(doc, submitter_g)

    # 5. Old Aztec + English header stamp out.
    _strip_header_artifacts(doc)

    doc.save(str(docx_path))


def _body_text_no_tables(docx_path: Path) -> str:
    """Return body paragraph text only — table cells excluded.

    doc.paragraphs already excludes table cells in python-docx, so this
    is the right source for the body-preservation check (table cell text
    is removed during normalize_data_table and would cause false positives).
    """
    return "\n".join(p.text for p in Document(str(docx_path)).paragraphs if p.text)


def validate_book_template(docx_path: Path) -> None:
    """Fail-closed check: dummy render must succeed under sandbox+strict and
    place each dummy value exactly once. Raises ValueError (operator-safe
    message, no paths/tracebacks)."""
    # Discover column count from the normalized template text: after
    # retokenize, data cells contain {{ row.c0 }}, {{ row.c1 }}, …
    # detect_table_schema cannot be used here because the normalized table
    # (with directive rows) no longer satisfies its "clean table" criteria.
    tpl_text = docx_to_text(docx_path)
    col_indices = [int(m) for m in re.findall(r"row\.c(\d+)", tpl_text)]
    has_table = bool(col_indices)
    dummy: dict[str, object] = dict(_DUMMY)
    if has_table:
        n = max(col_indices) + 1
        row: dict[str, str] = {f"c{i}": f"DUMMY_CELL_{i}" for i in range(n)}
        # SSTI probe: last cell carries a literal Jinja expression; it must
        # appear verbatim in the output (cell values are data, not re-expanded).
        row[f"c{n - 1}"] = "{{ ref }}"
        dummy["table_rows"] = [row]

    source_text = _body_text_no_tables(docx_path)
    with tempfile.TemporaryDirectory() as td:
        out = Path(td) / "check.docx"
        try:
            render(docx_path, dummy, out, strict=True, sandboxed=True)
        except Exception as exc:  # sandbox/strict/syntax — reason stays generic
            raise ValueError("تعذر التحقق من القالب — فشل عرض تجريبي") from exc
        text = docx_to_text(out)
    # submitter_g is deliberately NOT asserted — it is optional-inject (books
    # without a footer G-number are valid templates).
    if text.count(_DUMMY["ref"]) != 1 or text.count(_DUMMY["date"]) != 1:
        raise ValueError("سطر الرقم أو التاريخ لم يُستبدل بشكل صحيح")
    # Body preserved: every substantial source line (minus token lines)
    # must survive the render. Uses paragraph-only text to exclude table cells
    # (which are transformed/removed by normalize_data_table).
    for line in source_text.splitlines():
        line = line.strip()
        if (
            len(line) >= 15
            and "{{" not in line
            and "{%" not in line
            and line.replace(_ZWSP, "") not in text.replace(_ZWSP, "")
        ):
            raise ValueError("نص القالب تغيّر أثناء العرض التجريبي")
    # Table-specific assertions: ≥1 data row rendered + no double-expansion.
    if has_table:
        # When n > 1, c0 carries "DUMMY_CELL_0" (distinct from the SSTI probe
        # in the last cell) — assert it rendered.  With n == 1, c0 IS the SSTI
        # probe cell, so there is no separate DUMMY_CELL_0 to check.
        if n > 1 and "DUMMY_CELL_0" not in text:
            raise ValueError("الجدول لم يُصيَّر بشكل صحيح — لم تظهر صفوف البيانات")
        if "{{ ref }}" not in text:
            raise ValueError("قيمة خلية الجدول تعرّضت لإعادة تفسير غير مسموح بها")
