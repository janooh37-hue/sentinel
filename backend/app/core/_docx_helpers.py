"""Low-level OOXML helpers — kept for the post-process hooks in
`docx_engine.py` that can't live in Jinja templates.

This file used to host v3.5.4's full X-mark fill toolkit (`fill_x_mark`,
`set_cell_text`, `set_cell_image`, `fill_sig_and_date_on_x_mark`, etc.).
After the move to docxtpl token rendering, only two helpers survive:

* `fill_image_behind_text_on_x_mark` — used by the 301-004 / 301-005 admin
  forms so the manager-signature image anchors as a floating "behind text"
  shape and doesn't grow the cell row height. python-docx has no public
  API for this; we build the OOXML by hand.
* `replace_paragraph_text` — used by the Resignation Letter post-process
  to swap dotted-line paragraphs for the reason text when the reason is
  too long to fit in the inline X slot.

`fill_image_on_x_mark` is kept private (only the behind-text helper uses
it internally as a fallback).
"""

from __future__ import annotations

import contextlib
import io
from pathlib import Path
from typing import Any

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor

from app.core.signature_render import DEFAULT_SIG_BOLDNESS, prepare_signature

_ARIAL = "Arial"
_BLACK = RGBColor(0, 0, 0)
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS = "http://schemas.openxmlformats.org/drawingml/2006/main"


def _find_x_index(text: str) -> int:
    """Index of the first 'X' or 'x' in `text`, or -1 if absent."""
    idx_upper = text.find("X")
    idx_lower = text.find("x")
    if idx_upper == -1:
        return idx_lower
    if idx_lower == -1:
        return idx_upper
    return min(idx_upper, idx_lower)


def _fill_image_on_x_mark(
    cell: Any,
    image_path: Path | str | None,
    *,
    width_inches: float = 1.4,
    dilate_radius_px: int = DEFAULT_SIG_BOLDNESS,
) -> bool:
    """Replace the first X mark with an inline image (fallback path).

    Internal — only called by `fill_image_behind_text_on_x_mark` when the
    anchor swap fails. Returns True iff an X was found.
    """
    if not image_path or not Path(image_path).exists():
        return False
    for para in cell.paragraphs:
        for run in para.runs:
            if not run.text:
                continue
            idx = _find_x_index(run.text)
            if idx < 0:
                continue
            t = run.text
            run.text = t[:idx] + t[idx + 1 :]
            img_run = para.add_run()
            with contextlib.suppress(OSError, ValueError):
                img_run.add_picture(
                    io.BytesIO(
                        prepare_signature(
                            Path(image_path).read_bytes(),
                            dilate_radius_px=dilate_radius_px,
                        )
                    ),
                    width=Inches(width_inches),
                )
            return True
    with contextlib.suppress(OSError, ValueError):
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(
            io.BytesIO(
                prepare_signature(
                    Path(image_path).read_bytes(),
                    dilate_radius_px=dilate_radius_px,
                )
            ),
            width=Inches(width_inches),
        )
    return False


def fill_image_behind_text_on_x_mark(
    cell: Any,
    image_path: Path | str | None,
    *,
    width_inches: float = 1.4,
    dilate_radius_px: int = DEFAULT_SIG_BOLDNESS,
) -> bool:
    """Insert a 'behind text' floating image at the X mark.

    Used on 301-004 / 301-005 admin forms so the manager-signature image
    doesn't grow the cell and bump the date row below it. Falls back to
    inline placement on any XML error.
    """
    if not image_path or not Path(image_path).exists():
        return False
    try:
        target_run = None
        target_para = None
        for para in cell.paragraphs:
            for run in para.runs:
                if run.text and _find_x_index(run.text) >= 0:
                    target_run, target_para = run, para
                    break
            if target_run is not None:
                break
        if target_run is None or target_para is None:
            return _fill_image_on_x_mark(
                cell, image_path, width_inches=width_inches, dilate_radius_px=dilate_radius_px
            )
        idx = _find_x_index(target_run.text)
        t = target_run.text
        target_run.text = t[:idx] + t[idx + 1 :]

        img_run = target_para.add_run()
        img_run.add_picture(
            io.BytesIO(
                prepare_signature(
                    Path(image_path).read_bytes(),
                    dilate_radius_px=dilate_radius_px,
                )
            ),
            width=Inches(width_inches),
        )
        drawing = img_run._element.find(qn("w:drawing"))
        if drawing is None:
            return True
        inline = drawing.find(f"{{{_WP_NS}}}inline")
        if inline is None:
            return True
        extent = inline.find(f"{{{_WP_NS}}}extent")
        cx = extent.get("cx") if extent is not None else None
        cy = extent.get("cy") if extent is not None else None

        anchor_xml = (
            f'<wp:anchor xmlns:wp="{_WP_NS}" '
            'distT="0" distB="0" distL="0" distR="0" simplePos="0" '
            'relativeHeight="251660000" behindDoc="1" locked="0" '
            'layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="column">'
            "<wp:posOffset>0</wp:posOffset></wp:positionH>"
            '<wp:positionV relativeFrom="paragraph">'
            "<wp:posOffset>0</wp:posOffset></wp:positionV>"
            f'<wp:extent cx="{cx or 1143000}" cy="{cy or 457200}"/>'
            '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            "<wp:wrapNone/>"
            "</wp:anchor>"
        )
        anchor = parse_xml(anchor_xml)
        ns_uri = {"wp": _WP_NS, "a": _A_NS}
        for prefixed in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
            prefix, local = prefixed.split(":")
            child = inline.find(f"{{{ns_uri[prefix]}}}{local}")
            if child is not None:
                anchor.append(child)
        drawing.remove(inline)
        drawing.append(anchor)
        return True
    except (OSError, ValueError, AttributeError):
        try:
            return _fill_image_on_x_mark(
                cell, image_path, width_inches=width_inches, dilate_radius_px=dilate_radius_px
            )
        except (OSError, ValueError):
            return False


def insert_floating_image_in_header(
    header: Any,
    image_bytes: bytes,
    *,
    x_emu: int,
    y_emu: int,
    size_emu: int,
    behind: bool = False,
) -> bool:
    """Place *image_bytes* as a page-relative floating image in *header*.

    Used for the ref QR: pinned to a fixed page corner (offset ``x_emu``/``y_emu``
    from the top-left of the page), ``size_emu`` square, so it never reflows the
    letterhead. ``behind=False`` keeps it in front of content (its white
    quiet-zone means it never hides under a banner). Returns ``False`` on any
    OOXML error so the caller can proceed without the image.
    """
    if not image_bytes:
        return False
    try:
        para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = para.add_run()
        run.add_picture(io.BytesIO(image_bytes), width=Emu(size_emu))
        drawing = run._element.find(qn("w:drawing"))
        if drawing is None:
            return False
        inline = drawing.find(f"{{{_WP_NS}}}inline")
        if inline is None:
            return False
        extent = inline.find(f"{{{_WP_NS}}}extent")
        cx = extent.get("cx") if extent is not None else str(size_emu)
        cy = extent.get("cy") if extent is not None else str(size_emu)

        anchor_xml = (
            f'<wp:anchor xmlns:wp="{_WP_NS}" '
            'distT="0" distB="0" distL="0" distR="0" simplePos="0" '
            f'relativeHeight="251670000" behindDoc="{1 if behind else 0}" '
            'locked="0" layoutInCell="1" allowOverlap="1">'
            '<wp:simplePos x="0" y="0"/>'
            '<wp:positionH relativeFrom="page">'
            f"<wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>"
            '<wp:positionV relativeFrom="page">'
            f"<wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>"
            f'<wp:extent cx="{cx}" cy="{cy}"/>'
            '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
            "<wp:wrapNone/>"
            "</wp:anchor>"
        )
        anchor = parse_xml(anchor_xml)
        ns_uri = {"wp": _WP_NS, "a": _A_NS}
        for prefixed in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
            prefix, local = prefixed.split(":")
            child = inline.find(f"{{{ns_uri[prefix]}}}{local}")
            if child is not None:
                anchor.append(child)
        drawing.remove(inline)
        drawing.append(anchor)
        return True
    except Exception:
        return False


def replace_paragraph_text(
    para: Any,
    new_text: str,
    *,
    size: int = 11,
    bold: bool = False,
    align_center: bool = False,
) -> None:
    """Overwrite all runs in `para` with a single new run.

    Used by the Resignation Letter post-process to write the reason into
    the body cell's dotted-line paragraphs.
    """
    for run in para.runs:
        run.text = ""
    for run in list(para.runs):
        parent = run._element.getparent()
        if parent is not None:
            parent.remove(run._element)
    if align_center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(str(new_text))
    run.font.name = _ARIAL
    run.font.size = Pt(size)
    run.font.color.rgb = _BLACK
    run.font.bold = bold


__all__ = ["fill_image_behind_text_on_x_mark", "insert_floating_image_in_header", "replace_paragraph_text"]


# Ensure the OxmlElement re-export from the old module isn't accidentally
# needed by stale callers — surface a useful error if so.
_ = OxmlElement
