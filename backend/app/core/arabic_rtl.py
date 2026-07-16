"""Arabic / RTL run stamping ported from `arabic_editor.py` (v3.5.4 line 522+).

Word renders runs that carry both ``<w:rtl/>`` *and* ``<w:cs/>`` using their
complex-script properties (``rFonts/@cs``, ``<w:szCs>``, ``<w:bCs>``,
``<w:iCs>``) — and silently ignores the Latin properties python-docx writes
by default (``<w:sz>``, ``<w:b>``, ``<w:i>``). Without mirroring those into
the Cs slots Word falls back to the document default (Arial 11pt body),
which is the "Ashkhari" Arabic-name truncation bug v3 hit repeatedly.

This module exposes pure-OOXML helpers — no Qt, no PySide6. Higher layers
(Phase 04 General Book HTML pipeline, the React Arabic editor) feed in
already-parsed runs/paragraphs.

Public contract (per `plans/01-core-port.md`):

    stamp_run(run, family: str) -> None
        Mark a python-docx ``Run`` as RTL/complex-script, mirroring its
        Latin font properties into Cs equivalents.

    stamp_paragraph(paragraph) -> None
        Apply ``<w:bidi/>`` so the paragraph reads right-to-left.

    set_run_shading(run, rgb: tuple[int, int, int]) -> None
        Add ``<w:shd/>`` (background color) — used for TinyMCE highlights.

    html_to_docx(html, paragraph, **kwargs) -> None
        Render TinyMCE-style HTML after ``paragraph``. Ported from v3's
        ``arabic_editor.html_to_docx`` (lxml-based, no Qt). Renders real RTL
        Word tables, hyperlinks, line-height, and page breaks.
"""

from __future__ import annotations

import base64
import contextlib
import logging
import re
from io import BytesIO
from typing import Any, TypedDict

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

log = logging.getLogger(__name__)

_HTML_WS_RE = re.compile(r"[ \t\r\n\f]+")


class _WalkState(TypedDict, total=False):
    """State threaded through ``_walk_*`` and ``_state_new_paragraph``."""

    anchor: Any  # original paragraph passed in (required at runtime)
    current: Any  # paragraph currently being filled (required at runtime)
    first_used: bool  # True once the anchor paragraph has been used
    parent_obj: Any  # parent container of ``anchor`` (cell body, etc.)


def stamp_run(run: Any, family: str) -> None:
    """Mark `run` as Arabic/RTL and mirror Latin font props into Cs slots."""
    rPr = run._element.get_or_add_rPr()

    # w:rtl / w:cs flags — both are required for Word to pick the cs props.
    for tag in ("w:rtl", "w:cs"):
        if rPr.find(qn(tag)) is None:
            rPr.append(OxmlElement(tag))

    # rFonts: set complex-script + Latin font families to the same value.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), family)
    rFonts.set(qn("w:ascii"), family)
    rFonts.set(qn("w:hAnsi"), family)

    # Mirror sz → szCs (Word ignores w:sz on RTL runs).
    sz = rPr.find(qn("w:sz"))
    if sz is not None and sz.get(qn("w:val")):
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), sz.get(qn("w:val")))

    # Mirror b → bCs and i → iCs.
    if rPr.find(qn("w:b")) is not None and rPr.find(qn("w:bCs")) is None:
        rPr.append(OxmlElement("w:bCs"))
    if rPr.find(qn("w:i")) is not None and rPr.find(qn("w:iCs")) is None:
        rPr.append(OxmlElement("w:iCs"))


def stamp_paragraph(paragraph: Any) -> None:
    """Mark `paragraph` as bidirectional (right-to-left)."""
    pPr = paragraph._p.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def set_run_shading(run: Any, rgb: tuple[int, int, int]) -> None:
    """Set a background-fill colour on `run` via ``<w:shd/>``.

    Accepts a ``(r, g, b)`` tuple of 0-255 integers; mirrors v3's
    ``_set_run_shading`` which took a ``QColor`` and packed it to hex.
    """
    r, g, b = rgb
    if not all(0 <= v <= 255 for v in (r, g, b)):
        raise ValueError(f"rgb must be 0-255 per channel, got {rgb}")
    rPr = run._element.get_or_add_rPr()
    shd = rPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        rPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")


def stamp_arabic_runs(paragraph: Any, family: str = "Arial") -> int:
    """Stamp every run in `paragraph` that contains Arabic codepoints.

    Returns the number of runs touched. Convenience wrapper for the most
    common case: paragraph already filled, just needs RTL marking. Avoids
    cluttering call sites with ``for run in p.runs`` boilerplate.
    """
    count = 0
    for run in paragraph.runs:
        if _contains_arabic(run.text or ""):
            stamp_run(run, family)
            count += 1
    if count:
        stamp_paragraph(paragraph)
    return count


def _contains_arabic(text: str) -> bool:
    """True if `text` has at least one Arabic-block codepoint.

    Covers basic Arabic (U+0600-U+06FF), Arabic Supplement (U+0750-U+077F),
    Extended-A (U+08A0-U+08FF), Presentation Forms-A (U+FB50-U+FDFF), and
    Presentation Forms-B (U+FE70-U+FEFF). Misses Arabic Mathematical
    Alphabetic Symbols (U+1EE00+) — acceptable: those don't appear on forms.
    """
    for ch in text:
        code = ord(ch)
        if (
            0x0600 <= code <= 0x06FF
            or 0x0750 <= code <= 0x077F
            or 0x08A0 <= code <= 0x08FF
            or 0xFB50 <= code <= 0xFDFF
            or 0xFE70 <= code <= 0xFEFF
        ):
            return True
    return False


# ---------------------------------------------------------------------------
# html_to_docx — TinyMCE HTML → python-docx (lxml-based, no Qt).
# Ported from v3's arabic_editor.html_to_docx (line 1795+).
# Tables render as real RTL Word tables (see _render_table).
# ---------------------------------------------------------------------------

_PT_PER_PX: float = 0.75  # 1 px @ 96 dpi = 0.75 pt
_PT_PER_EM: float = 12.0  # body default; em maps relative to this
_TWIPS_PER_PX: float = 15.0  # 1 px @ 96 dpi = 15 twips

_HTML_NAMED_COLORS: dict[str, str | None] = {
    "black": "000000",
    "white": "FFFFFF",
    "red": "FF0000",
    "green": "008000",
    "blue": "0000FF",
    "yellow": "FFFF00",
    "cyan": "00FFFF",
    "magenta": "FF00FF",
    "gray": "808080",
    "grey": "808080",
    "silver": "C0C0C0",
    "maroon": "800000",
    "olive": "808000",
    "purple": "800080",
    "teal": "008080",
    "navy": "000080",
    "lime": "00FF00",
    "aqua": "00FFFF",
    "fuchsia": "FF00FF",
    "orange": "FFA500",
    "pink": "FFC0CB",
    "transparent": None,
}


def _hparse_color(value: str) -> str | None:
    """Return a 6-char hex (no #) or None for transparent/unknown."""
    if not value:
        return None
    v = value.strip().lower()
    if v in ("inherit", "initial", "currentcolor", "none"):
        return None
    if v in _HTML_NAMED_COLORS:
        return _HTML_NAMED_COLORS[v]
    if v.startswith("#"):
        h = v[1:]
        if len(h) == 3:
            return (h[0] * 2 + h[1] * 2 + h[2] * 2).upper()
        if len(h) == 6:
            return h.upper()
        return None
    m = re.match(r"rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", v)
    if m:
        r, g, b = (max(0, min(255, int(x))) for x in m.groups())
        am = re.match(r"rgba\(\s*\d+\s*,\s*\d+\s*,\s*\d+\s*,\s*([\d.]+)\s*\)", v)
        if am:
            try:
                if float(am.group(1)) <= 0.0:
                    return None
            except ValueError:
                pass
        return f"{r:02X}{g:02X}{b:02X}"
    return None


def _hparse_size_pt(value: str) -> float | None:
    """Parse a CSS font-size or HTML size attr to points."""
    if not value:
        return None
    v = value.strip().lower()
    m = re.match(r"([\d.]+)\s*(pt|px|em|rem|%)?$", v)
    if not m:
        return None
    try:
        n = float(m.group(1))
    except ValueError:
        return None
    unit = m.group(2) or "px"
    if unit == "pt":
        return n
    if unit == "px":
        return n * _PT_PER_PX
    if unit in ("em", "rem"):
        return n * _PT_PER_EM
    if unit == "%":
        return n / 100.0 * _PT_PER_EM
    return None


def _parse_inline_style(style: str) -> dict[str, str]:
    """Parse a CSS ``style="..."`` attribute into a dict (lower-case keys)."""
    out: dict[str, str] = {}
    if not style:
        return out
    for chunk in style.split(";"):
        if ":" not in chunk:
            continue
        k, _, v = chunk.partition(":")
        k = k.strip().lower()
        v = v.strip()
        if k:
            out[k] = v
    return out


class _Fmt:
    """Inline character format inherited down the HTML tree."""

    __slots__ = (
        "bold",
        "color",
        "family",
        "highlight",
        "italic",
        "size_pt",
        "strike",
        "underline",
    )

    def __init__(self) -> None:
        self.family: str | None = None
        self.size_pt: float | None = None
        self.bold: bool = False
        self.italic: bool = False
        self.underline: bool = False
        self.strike: bool = False
        self.color: str | None = None
        self.highlight: str | None = None

    def copy(self) -> _Fmt:
        c = _Fmt()
        c.family = self.family
        c.size_pt = self.size_pt
        c.bold = self.bold
        c.italic = self.italic
        c.underline = self.underline
        c.strike = self.strike
        c.color = self.color
        c.highlight = self.highlight
        return c


class _BlockFmt:
    """Block-level format inherited at paragraph creation time."""

    __slots__ = ("align", "indent_left_px", "indent_right_px", "line_height", "rtl")

    def __init__(self) -> None:
        self.align: str | None = None
        self.rtl: bool = True  # default RTL — this editor is for Arabic
        self.indent_left_px: int = 0
        self.indent_right_px: int = 0
        self.line_height: float | None = None  # unit-less CSS line-height multiplier

    def copy(self) -> _BlockFmt:
        c = _BlockFmt()
        c.align = self.align
        c.rtl = self.rtl
        c.indent_left_px = self.indent_left_px
        c.indent_right_px = self.indent_right_px
        c.line_height = self.line_height
        return c


_INLINE_TAGS_BOLD = {"b", "strong"}
_INLINE_TAGS_ITALIC = {"i", "em"}
_INLINE_TAGS_UNDERLINE = {"u", "ins"}
_INLINE_TAGS_STRIKE = {"s", "strike", "del"}
_BLOCK_TAGS = {
    "p",
    "div",
    "h1",
    "h2",
    "h3",
    "h4",
    "h5",
    "h6",
    "blockquote",
    "pre",
    "li",
}
_HEADING_SIZE_PT: dict[str, float] = {
    "h1": 24.0,
    "h2": 18.0,
    "h3": 14.0,
    "h4": 12.0,
    "h5": 10.0,
    "h6": 8.0,
}


def _merge_fmt_from_attrs(fmt: _Fmt, tag: str, attrs: dict[str, str]) -> _Fmt:
    """Apply tag-implied + style/attr-driven formatting to a copy of ``fmt``."""
    out = fmt.copy()
    t = tag.lower()
    if t in _INLINE_TAGS_BOLD:
        out.bold = True
    if t in _INLINE_TAGS_ITALIC:
        out.italic = True
    if t in _INLINE_TAGS_UNDERLINE:
        out.underline = True
    if t in _INLINE_TAGS_STRIKE:
        out.strike = True
    if t in _HEADING_SIZE_PT:
        out.bold = True
        out.size_pt = _HEADING_SIZE_PT[t]

    if t == "font":
        face = attrs.get("face")
        if face:
            out.family = face.split(",")[0].strip().strip("'\"")
        col = _hparse_color(attrs.get("color", ""))
        if col:
            out.color = col
        sz = attrs.get("size")
        if sz and sz.isdigit():
            sizes_pt = {1: 8, 2: 10, 3: 12, 4: 14, 5: 18, 6: 24, 7: 36}
            out.size_pt = float(sizes_pt.get(int(sz), 12))

    style = _parse_inline_style(attrs.get("style", ""))
    if "font-family" in style:
        ff = style["font-family"].split(",")[0].strip().strip("'\"")
        if ff:
            out.family = ff
    if "font-size" in style:
        s = _hparse_size_pt(style["font-size"])
        if s is not None and s > 0:
            out.size_pt = s
    if "font-weight" in style:
        w = style["font-weight"].lower()
        if w in ("bold", "bolder") or (w.isdigit() and int(w) >= 600):
            out.bold = True
        elif w == "normal" or (w.isdigit() and int(w) < 600):
            out.bold = False
    if "font-style" in style:
        out.italic = style["font-style"].lower() in ("italic", "oblique")
    if "text-decoration" in style or "text-decoration-line" in style:
        td = (style.get("text-decoration") or style.get("text-decoration-line") or "").lower()
        if "underline" in td:
            out.underline = True
        if "line-through" in td:
            out.strike = True
        if td == "none":
            out.underline = False
            out.strike = False
    if "color" in style:
        col = _hparse_color(style["color"])
        if col:
            out.color = col
    if "background-color" in style or "background" in style:
        bg = _hparse_color(style.get("background-color") or style.get("background", ""))
        if bg:
            out.highlight = bg

    if t == "mark" and not out.highlight:
        out.highlight = "FFFF00"

    return out


def _merge_block_fmt_from_attrs(blk: _BlockFmt, tag: str, attrs: dict[str, str]) -> _BlockFmt:
    out = blk.copy()
    style = _parse_inline_style(attrs.get("style", ""))
    align = attrs.get("align") or style.get("text-align")
    if align:
        a = align.lower().strip()
        if a in ("left", "right", "center", "justify"):
            out.align = a
    direction = (attrs.get("dir") or style.get("direction") or "").lower()
    if direction == "rtl":
        out.rtl = True
    elif direction == "ltr":
        out.rtl = False
    for k in ("margin-left", "padding-left"):
        v = style.get(k)
        if v:
            try:
                px = float(re.match(r"([\d.]+)", v).group(1))  # type: ignore[union-attr]
                if "pt" in v:
                    px = px / _PT_PER_PX
                out.indent_left_px = max(out.indent_left_px, int(px))
            except (AttributeError, ValueError):
                pass
    for k in ("margin-right", "padding-right"):
        v = style.get(k)
        if v:
            try:
                px = float(re.match(r"([\d.]+)", v).group(1))  # type: ignore[union-attr]
                if "pt" in v:
                    px = px / _PT_PER_PX
                out.indent_right_px = max(out.indent_right_px, int(px))
            except (AttributeError, ValueError):
                pass
    lh = style.get("line-height")
    if lh:
        out.line_height = _parse_line_height(lh)
    return out


def _parse_line_height(value: str) -> float | None:
    """Parse a CSS ``line-height`` to a unit-less multiplier (e.g. 1.5).

    Word's ``line_spacing`` float is a multiple of single spacing. We map a
    unit-less value directly; ``%`` divides by 100; ``pt``/``px`` lengths are
    converted relative to the 12pt body default. Returns None when unparseable.
    """
    v = value.strip().lower()
    if v in ("normal", "inherit", "initial"):
        return None
    m = re.match(r"([\d.]+)\s*(pt|px|em|rem|%)?$", v)
    if not m:
        return None
    try:
        n = float(m.group(1))
    except ValueError:
        return None
    unit = m.group(2)
    if n <= 0:
        return None
    if unit in (None, "", "em", "rem"):
        return n
    if unit == "%":
        return n / 100.0
    if unit == "pt":
        return n / _PT_PER_EM
    if unit == "px":
        return (n * _PT_PER_PX) / _PT_PER_EM
    return None


def _apply_block_fmt(paragraph: Any, blk: _BlockFmt) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if blk.align:
        paragraph.alignment = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }[blk.align]
    if blk.rtl:
        stamp_paragraph(paragraph)
    if blk.line_height and blk.line_height > 0:
        paragraph.paragraph_format.line_spacing = blk.line_height
    if blk.indent_left_px or blk.indent_right_px:
        pPr = paragraph._p.get_or_add_pPr()
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        if blk.indent_left_px:
            ind.set(qn("w:left"), str(int(blk.indent_left_px * _TWIPS_PER_PX)))
        if blk.indent_right_px:
            ind.set(qn("w:right"), str(int(blk.indent_right_px * _TWIPS_PER_PX)))


def _apply_run_fmt(
    run: Any,
    fmt: _Fmt,
    default_family: str,
    default_size: float,
) -> None:
    from docx.shared import Pt, RGBColor

    family = fmt.family or default_family
    run.font.name = family
    run.font.size = Pt(fmt.size_pt if fmt.size_pt and fmt.size_pt > 0 else default_size)
    if fmt.bold:
        run.bold = True
    if fmt.italic:
        run.italic = True
    if fmt.underline:
        run.underline = True
    if fmt.strike:
        run.font.strike = True
    if fmt.color:
        with contextlib.suppress(ValueError, IndexError):
            run.font.color.rgb = RGBColor(
                int(fmt.color[0:2], 16),
                int(fmt.color[2:4], 16),
                int(fmt.color[4:6], 16),
            )
    if fmt.highlight and fmt.highlight not in ("FFFFFF",):
        with contextlib.suppress(ValueError, IndexError):
            r = int(fmt.highlight[0:2], 16)
            g = int(fmt.highlight[2:4], 16)
            b = int(fmt.highlight[4:6], 16)
            set_run_shading(run, (r, g, b))
    stamp_run(run, family)


def _embed_html_image(paragraph: Any, src: str, width_px: int) -> bool:
    """Embed a TinyMCE ``<img src=...>``. Supports data: URLs and local paths."""
    from docx.shared import Inches

    if not src:
        return False
    img_bytes: bytes | None = None
    if src.startswith("data:"):
        try:
            head, _, b64 = src.partition(",")
            if "base64" in head and b64:
                img_bytes = base64.b64decode(b64)
        except Exception as e:
            log.warning("_embed_html_image: base64 decode failed for src=%r: %s", src[:80], e)
            return False
    else:
        # Try as a local filesystem path (absolute path or no-scheme URL).
        try:
            from pathlib import Path as _P

            p = _P(src)
            if p.exists():
                img_bytes = p.read_bytes()
        except (OSError, ValueError):
            pass
    if not img_bytes:
        return False
    try:
        run = paragraph.add_run()
        bio = BytesIO(img_bytes)
        if width_px > 0:
            run.add_picture(bio, width=Inches(min(width_px / 96.0, 6.5)))
        else:
            run.add_picture(bio)
        return True
    except Exception as e:
        log.warning("_embed_html_image: add_picture failed for src=%r: %s", src[:80], e)
        return False


def _add_hyperlink(paragraph: Any, url: str, text: str, family: str) -> bool:
    """Append a real Word hyperlink (``<w:hyperlink>``) carrying ``text``.

    Creates an external relationship on the paragraph's part and a blue,
    underlined run inside the hyperlink wrapper. Returns ``True`` on success;
    ``False`` lets the caller fall back to a plain (blue/underlined) run so the
    text is never dropped.
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.shared import Pt, RGBColor

    if not url or not text:
        return False
    try:
        part = paragraph.part
        r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    except Exception as e:  # pragma: no cover - defensive
        log.warning("_add_hyperlink: relate_to failed for url=%r: %s", url[:80], e)
        return False

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    run.append(rPr)
    wt = OxmlElement("w:t")
    wt.set(qn("xml:space"), "preserve")
    wt.text = text
    run.append(wt)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)

    # Stamp the inner run RTL/cs + mirror size/family so Word renders Arabic
    # link text correctly. Wrap the raw <w:r> in a python-docx Run to reuse
    # the helpers.
    from docx.text.run import Run

    pyrun = Run(run, paragraph)
    pyrun.font.size = Pt(_PT_PER_EM)
    pyrun.font.name = family
    pyrun.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    pyrun.underline = True
    stamp_run(pyrun, family)
    return True


def _is_page_break_node(tag: str, attrs: dict[str, str]) -> tuple[bool, bool]:
    """Detect a TinyMCE page break. Returns ``(before, after)`` booleans.

    Handles ``<div class="mce-pagebreak">`` (a *standalone* page break, mapped
    to break-before) and ``page-break-before/after: always`` CSS on any block.
    """
    classes = (attrs.get("class") or "").lower().split()
    if "mce-pagebreak" in classes:
        return True, False
    style = _parse_inline_style(attrs.get("style", ""))
    before = "always" in (style.get("page-break-before") or style.get("break-before") or "")
    after = "always" in (style.get("page-break-after") or style.get("break-after") or "")
    return before, after


def _emit_page_break(state: _WalkState) -> None:
    """Emit a page break in its own paragraph, preserving document order."""
    from docx.enum.text import WD_BREAK

    p = _state_new_paragraph(state)
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def _state_new_paragraph(state: _WalkState) -> Any:
    """Allocate a new ``w:p`` after state['current'], or reuse the anchor first."""
    from docx.text.paragraph import Paragraph

    parent_obj = state.get("parent_obj") or state["anchor"]._parent
    if not state["first_used"]:
        state["first_used"] = True
        for run in list(state["anchor"].runs):
            run.text = ""
        state["current"] = state["anchor"]
        return state["anchor"]
    cur = state["current"]
    anchor_elem = _cur_elem(cur)
    new_p = OxmlElement("w:p")
    anchor_elem.addnext(new_p)
    pyp = Paragraph(new_p, parent_obj)
    state["current"] = pyp
    return pyp


def _cur_elem(cur: Any) -> Any:
    """Return the underlying lxml element for a ``state['current']`` value.

    ``current`` is a python-docx ``Paragraph`` (``._p``), a ``Table``
    (``._tbl``), or a raw ``<w:tbl>`` element (used as the cursor right after a
    table is inserted). Falls through to ``cur`` itself for the raw case.
    """
    if hasattr(cur, "_p"):
        return cur._p
    if hasattr(cur, "_tbl"):
        return cur._tbl
    return cur


def _state_insert_table(state: _WalkState, tbl_elem: Any) -> None:
    """Insert a built ``<w:tbl>`` after the cursor and advance the cursor.

    Threading the cursor onto the raw table element means the NEXT block (or
    the trailing narrative) is created *after* the table — preserving document
    order around it. If the anchor hasn't been used yet (leading-table
    fragment) it is consumed (runs cleared) and the table placed right after
    it, mirroring how ``_state_new_paragraph`` reuses the anchor.
    """
    if not state["first_used"]:
        state["first_used"] = True
        anchor = state["anchor"]
        for run in list(anchor.runs):
            run.text = ""
        anchor._p.addnext(tbl_elem)
        state["current"] = tbl_elem
        return
    cur_elem = _cur_elem(state["current"])
    cur_elem.addnext(tbl_elem)
    state["current"] = tbl_elem


def _collapse_html_whitespace(root: Any) -> None:
    """Collapse runs of ASCII whitespace in text/tail nodes to a single space
    (HTML rendering semantics). NBSP (\\u00a0) is intentionally preserved —
    the editor uses it for deliberate spacing. ``<pre>`` subtrees keep their
    whitespace verbatim (the editor's "Preformatted" block format).
    """
    for el in root.iter():
        if not isinstance(el.tag, str):
            if el.tail:
                el.tail = _HTML_WS_RE.sub(" ", el.tail)
            continue
        tag = el.tag.lower()
        in_pre = tag == "pre" or any(
            isinstance(a.tag, str) and a.tag.lower() == "pre" for a in el.iterancestors()
        )
        if not in_pre and el.text:
            el.text = _HTML_WS_RE.sub(" ", el.text)
        # The tail sits OUTSIDE the element: only an enclosing <pre> protects it.
        tail_in_pre = any(
            isinstance(a.tag, str) and a.tag.lower() == "pre" for a in el.iterancestors()
        )
        if not tail_in_pre and el.tail:
            el.tail = _HTML_WS_RE.sub(" ", el.tail)


def _emit_text_into_paragraph(
    paragraph: Any,
    text: str,
    fmt: _Fmt,
    default_family: str,
    default_size: float,
) -> None:
    if not text:
        return
    if not text.strip(" ") and not paragraph.runs:
        # Space-only text at paragraph start renders as nothing in HTML.
        return
    run = paragraph.add_run(text)
    _apply_run_fmt(run, fmt, default_family, default_size)


def _walk_inline(
    node: Any,
    paragraph: Any,
    fmt: _Fmt,
    blk: _BlockFmt,
    state: _WalkState,
    default_family: str,
    default_size: float,
) -> Any:
    """Walk an inline subtree, adding runs to ``paragraph``. Returns current paragraph."""
    if node.text:
        _emit_text_into_paragraph(paragraph, node.text, fmt, default_family, default_size)
    for child in node:
        if not isinstance(child.tag, str):
            # Comment / processing-instruction node — never narrative text.
            # HugeRTE's pagebreak plugin serializes breaks as the literal
            # comment `<!-- pagebreak -->` (its default pagebreak_separator).
            if "pagebreak" in ((child.text or "").lower()):
                _emit_page_break(state)
                paragraph = _state_new_paragraph(state)
                _apply_block_fmt(paragraph, blk)
            if child.tail:
                _emit_text_into_paragraph(paragraph, child.tail, fmt, default_family, default_size)
            continue
        tag = (child.tag or "").lower() if isinstance(child.tag, str) else ""
        attrs: dict[str, str] = dict(child.attrib) if hasattr(child, "attrib") else {}
        if tag == "br":
            run = paragraph.add_run()
            run.add_break()
        elif tag == "a" and attrs.get("href"):
            link_text = "".join(child.itertext())
            href = attrs["href"]
            if not _add_hyperlink(paragraph, href, link_text, fmt.family or default_family):
                # Fallback: blue underlined run so the text is never dropped.
                link_fmt = fmt.copy()
                link_fmt.color = "0563C1"
                link_fmt.underline = True
                _emit_text_into_paragraph(
                    paragraph, link_text, link_fmt, default_family, default_size
                )
        elif tag == "img":
            try:
                w = int(attrs.get("width", "0") or 0)
            except ValueError:
                w = 0
            style = _parse_inline_style(attrs.get("style", ""))
            if not w and "width" in style:
                m = re.match(r"([\d.]+)", style["width"])
                if m:
                    try:
                        w = int(float(m.group(1)))
                    except ValueError:
                        w = 0
            _embed_html_image(paragraph, attrs.get("src", ""), w)
        elif tag in _BLOCK_TAGS or tag in ("ul", "ol", "table"):
            _walk_block(child, blk, state, default_family, default_size)
            paragraph = _state_new_paragraph(state)
            _apply_block_fmt(paragraph, blk)
        else:
            new_fmt = _merge_fmt_from_attrs(fmt, tag, attrs)
            paragraph = _walk_inline(
                child, paragraph, new_fmt, blk, state, default_family, default_size
            )
        if child.tail:
            _emit_text_into_paragraph(paragraph, child.tail, fmt, default_family, default_size)
    return paragraph


def _table_rtl(node: Any, attrs: dict[str, str]) -> bool:
    """True unless the table explicitly opts out of RTL via ``dir="ltr"``."""
    style = _parse_inline_style(attrs.get("style", ""))
    direction = (attrs.get("dir") or style.get("direction") or "").lower()
    return direction != "ltr"


def _collect_table_rows(node: Any) -> list[list[Any]]:
    """Return the table's rows (``<tr>`` lxml elements grouped per row) from
    direct children and any ``<thead>/<tbody>/<tfoot>`` section wrappers."""
    rows: list[list[Any]] = []
    for child in node:
        ctag = (child.tag or "").lower() if isinstance(child.tag, str) else ""
        if ctag == "tr":
            rows.append([c for c in child if _is_cell(c)])
        elif ctag in ("thead", "tbody", "tfoot"):
            for tr in child:
                tr_tag = (tr.tag or "").lower() if isinstance(tr.tag, str) else ""
                if tr_tag == "tr":
                    rows.append([c for c in tr if _is_cell(c)])
    return rows


def _is_cell(node: Any) -> bool:
    t = (node.tag or "").lower() if isinstance(node.tag, str) else ""
    return t in ("td", "th")


def _table_content_twips(state: _WalkState) -> int:
    """Page content width (twips) of the anchor's document section.

    ``page_width - left - right`` margins in EMU, EMU/635 -> twips. Falls back
    to 9360 twips (6.5" Letter w/ 1" margins) when section dims are missing.
    """
    try:
        doc = state["anchor"].part.document
        sect = doc.sections[0]
        content_emu = sect.page_width - sect.left_margin - sect.right_margin
        content_twips = int(content_emu / 635)
    except (IndexError, AttributeError, TypeError):
        return 9360
    return content_twips if content_twips > 0 else 9360


def _col_fractions(node: Any, rows: list[list[Any]], n: int) -> list[float]:
    """Return ``n`` column fractions (sum 1.0) from ``<colgroup><col>`` widths,
    falling back to the first row's cell widths, then to an even split."""
    raw: list[str] = []
    for child in node:
        ctag = (child.tag or "").lower() if isinstance(child.tag, str) else ""
        if ctag == "colgroup":
            for col in child:
                col_tag = (col.tag or "").lower() if isinstance(col.tag, str) else ""
                if col_tag == "col":
                    a = dict(col.attrib)
                    style = _parse_inline_style(a.get("style", ""))
                    raw.append(style.get("width") or a.get("width", ""))
        elif ctag == "col":
            a = dict(child.attrib)
            style = _parse_inline_style(a.get("style", ""))
            raw.append(style.get("width") or a.get("width", ""))
    if len([w for w in raw if w]) != n:
        first = rows[0] if rows else []
        raw = [_parse_inline_style(dict(c.attrib).get("style", "")).get("width", "") for c in first]
    nums: list[float] = []
    for w in raw:
        m = re.match(r"\s*([\d.]+)", str(w))
        nums.append(float(m.group(1)) if m else 0.0)
    if len(nums) != n or sum(nums) <= 0:
        return [1.0 / n] * n
    total = sum(nums)
    return [x / total for x in nums]


def _set_table_rtl_and_width(tbl: Any, content_twips: int, col_twips: list[int], rtl: bool) -> None:
    """Stamp RTL bidiVisual + jc=right (when ``rtl``), full fixed width, and
    per-column grid widths on a freshly created ``Table``."""
    tblPr = tbl._tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl._tbl.insert(0, tblPr)
    if rtl:
        if tblPr.find(qn("w:bidiVisual")) is None:
            tblPr.append(OxmlElement("w:bidiVisual"))
        jc = tblPr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tblPr.append(jc)
        jc.set(qn("w:val"), "right")

    for existing in tblPr.findall(qn("w:tblW")):
        tblPr.remove(existing)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(content_twips))
    tblPr.append(tblW)
    for existing in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(existing)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    grid = tbl._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w in zip(grid.findall(qn("w:gridCol")), col_twips, strict=False):
            gc.set(qn("w:w"), str(w))


def _paragraph_is_visually_empty(p: Any) -> bool:
    """True when the paragraph would render as blank: no visible text
    (NBSP counts as visible — str.strip() removes it, so check explicitly),
    no images, no breaks."""
    text = "".join(r.text or "" for r in p.runs)
    if text.strip() or chr(0xA0) in text:
        return False
    return not p._p.xpath(".//w:drawing | .//w:br | .//w:pict")


def _render_table(
    node: Any,
    blk: _BlockFmt,
    state: _WalkState,
    default_family: str,
    default_size: float,
) -> None:
    """Render an lxml ``<table>`` node as a real RTL Word ``Table`` inserted
    inline at the walker's current position.

    Ports the behaviour of ``docx_engine._pp_general_book``: full-width fixed
    layout, ``Table Grid`` style, ``bidiVisual`` + ``jc=right`` (unless an
    explicit ``dir="ltr"``), column widths from ``<colgroup>`` then first-row
    cell widths then even split, and per-cell inline rendering through
    ``_walk_inline`` so nested formatting survives. Cell style cascades
    table → row → cell (cell wins).
    """
    from docx.shared import Pt

    attrs: dict[str, str] = dict(node.attrib) if hasattr(node, "attrib") else {}
    rtl = _table_rtl(node, attrs)
    table_style = attrs.get("style", "")

    rows = _collect_table_rows(node)
    n_rows = len(rows)
    n_cols = max((len(r) for r in rows), default=0)
    if n_rows == 0 or n_cols == 0:
        return

    doc = state["anchor"].part.document
    content_twips = _table_content_twips(state)
    fracs = _col_fractions(node, rows, n_cols)
    col_twips = [max(1, int(content_twips * f)) for f in fracs]

    # add_table appends at body end; we relocate it via _state_insert_table.
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.style = "Table Grid"
    tbl.autofit = False
    _set_table_rtl_and_width(tbl, content_twips, col_twips, rtl)

    for r_idx, row_cells in enumerate(rows):
        row_attrs = dict(row_cells[0].getparent().attrib) if row_cells else {}
        row_style = row_attrs.get("style", "")
        for c_idx in range(n_cols):
            cell = tbl.rows[r_idx].cells[c_idx]
            cell_node = row_cells[c_idx] if c_idx < len(row_cells) else None
            cell_attrs = dict(cell_node.attrib) if cell_node is not None else {}
            cell_tag = (
                (cell_node.tag or "").lower()
                if cell_node is not None and isinstance(cell_node.tag, str)
                else "td"
            )
            cell_style = cell_attrs.get("style", "")

            # Effective cascaded style (cell wins) for block/run/background.
            eff_style = "; ".join(s for s in (table_style, row_style, cell_style) if s)
            eff_attrs = {**attrs, **row_attrs, **cell_attrs, "style": eff_style}

            # Block format: RTL/right default unless the table opts out.
            cblk = _BlockFmt()
            cblk.rtl = rtl
            cblk.align = "right" if rtl else "left"
            cblk = _merge_block_fmt_from_attrs(cblk, cell_tag, eff_attrs)

            # Run format: Calibri/12pt cell default, <th> bold, then cascade.
            fmt = _Fmt()
            fmt.family = "Calibri"
            fmt.size_pt = 12.0
            if cell_tag == "th":
                fmt.bold = True
            fmt = _merge_fmt_from_attrs(fmt, cell_tag, eff_attrs)

            # Per-cell column width under fixed layout.
            tcPr = cell._tc.get_or_add_tcPr()
            for existing in tcPr.findall(qn("w:tcW")):
                tcPr.remove(existing)
            tcW = OxmlElement("w:tcW")
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(col_twips[c_idx]))
            tcPr.append(tcW)

            para = cell.paragraphs[0]
            for rr in para.runs:
                rr.text = ""
            _apply_block_fmt(para, cblk)

            if cell_node is not None:
                # first_used=False so the first block child (<p> — Word wraps
                # every cell's text in one) REUSES this paragraph instead of
                # leaving it empty. This was the 3-paragraphs-per-cell bug.
                sub: _WalkState = {
                    "anchor": para,
                    "current": para,
                    "first_used": False,
                    "parent_obj": para._parent,
                }
                _walk_inline(cell_node, para, fmt, cblk, sub, default_family, default_size)
                # Drop blank paragraphs the walker left behind (the speculative
                # paragraph allocated after each block child). Keep >= 1.
                paras = list(cell.paragraphs)
                removable = [q for q in paras if _paragraph_is_visually_empty(q)]
                if len(removable) == len(paras):
                    removable = removable[1:]
                for q in removable:
                    q._p.getparent().remove(q._p)

            # Hug EVERY remaining paragraph: zero spacing so rows don't render
            # taller than their content. An explicit line-height (set by
            # _apply_block_fmt from cascaded CSS) still wins.
            for q in cell.paragraphs:
                q.paragraph_format.space_before = Pt(0)
                q.paragraph_format.space_after = Pt(0)
                if q.paragraph_format.line_spacing is None:
                    q.paragraph_format.line_spacing = 1.0

            # Cell background shading from cascaded background / background-color.
            style = _parse_inline_style(eff_style)
            bg_raw = style.get("background-color") or style.get("background", "")
            bg_hex = _hparse_color(bg_raw) if bg_raw else None
            if bg_hex:
                for existing in tcPr.findall(qn("w:shd")):
                    tcPr.remove(existing)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), bg_hex.upper())
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                tcPr.append(shd)

    # Relocate the appended table to the current walker position (in order).
    # The cursor becomes the raw <w:tbl>, so the next block lands after it.
    _state_insert_table(state, tbl._tbl)


def _walk_block(
    node: Any,
    blk: _BlockFmt,
    state: _WalkState,
    default_family: str,
    default_size: float,
) -> None:
    tag = (node.tag or "").lower() if isinstance(node.tag, str) else ""
    attrs: dict[str, str] = dict(node.attrib) if hasattr(node, "attrib") else {}

    pb_before, pb_after = _is_page_break_node(tag, attrs)
    is_bare_pagebreak = "mce-pagebreak" in (attrs.get("class") or "").lower().split()
    if pb_before:
        _emit_page_break(state)
    if is_bare_pagebreak:
        # The mce-pagebreak div carries no narrative — the break is the content.
        return

    if tag == "table":
        _render_table(node, blk, state, default_family, default_size)
        if pb_after:
            _emit_page_break(state)
        return

    if tag in ("ul", "ol"):
        is_numbered = tag == "ol"
        for li in node:
            li_tag = (li.tag or "").lower() if isinstance(li.tag, str) else ""
            if li_tag != "li":
                continue
            li_attrs: dict[str, str] = dict(li.attrib) if hasattr(li, "attrib") else {}
            new_blk = _merge_block_fmt_from_attrs(blk, "li", li_attrs)
            paragraph = _state_new_paragraph(state)
            _apply_block_fmt(paragraph, new_blk)
            style_name = "List Number" if is_numbered else "List Bullet"
            try:
                paragraph.style = style_name
            except (KeyError, ValueError):
                idx = list(node).index(li) + 1
                prefix = f"{idx}. " if is_numbered else "• "
                run = paragraph.add_run(prefix)
                _apply_run_fmt(run, _Fmt(), default_family, default_size)
            li_fmt = _merge_fmt_from_attrs(_Fmt(), "li", li_attrs)
            sub: _WalkState = {
                "anchor": paragraph,
                "current": paragraph,
                "first_used": True,
                "parent_obj": state.get("parent_obj"),
            }
            _walk_inline(li, paragraph, li_fmt, new_blk, sub, default_family, default_size)
        if pb_after:
            _emit_page_break(state)
        return

    new_blk = _merge_block_fmt_from_attrs(blk, tag, attrs)
    paragraph = _state_new_paragraph(state)
    _apply_block_fmt(paragraph, new_blk)
    fmt0 = _merge_fmt_from_attrs(_Fmt(), tag, attrs)
    _walk_inline(node, paragraph, fmt0, new_blk, state, default_family, default_size)
    if pb_after:
        _emit_page_break(state)


def html_to_docx(
    html: str,
    paragraph: Any,
    **kwargs: Any,
) -> None:
    """Render TinyMCE-style HTML into a python-docx Paragraph (and siblings).

    The first emitted block reuses ``paragraph`` (its existing runs are
    cleared). Subsequent blocks become new ``<w:p>`` elements after the
    anchor. Inline formatting (bold/italic/underline/color/size/family/
    background) is inherited through nested tags via ``_Fmt``. All runs
    are stamped with ``w:rtl`` / ``w:cs`` / ``w:bidi`` for RTL Arabic.

    Supported HTML:
    - Block tags: ``<p>``, ``<div>``, ``<h1>``-``<h6>``, ``<blockquote>``
    - Inline styling: ``<b>/<strong>``, ``<i>/<em>``, ``<u>``, ``<s>/<del>``
    - ``<span style="...">`` with ``color``, ``font-size``, ``background-color``
    - Block alignment via ``style="text-align: ..."``
    - Lists: ``<ul>``, ``<ol>``, ``<li>`` (nested)
    - Inline images: ``<img src="..." style="width: Npx;">`` (base64 data URIs
      and absolute file paths)
    - ``<font face=... color=... size=...>``
    - Hyperlinks: ``<a href="...">`` → a real ``<w:hyperlink>`` (blue/underlined)
    - Block ``line-height`` → paragraph line spacing
    - Page breaks: ``<div class="mce-pagebreak">`` and
      ``page-break-before/after: always``
    - Tables: ``<table>`` (with ``<thead>/<tbody>/<tfoot>``, ``<colgroup>``)
      render as real RTL Word tables, full-width fixed layout, anchored inline.
    """
    default_family: str = kwargs.get("default_family", "Calibri")
    default_size: float = kwargs.get("default_size", 12.0)
    if not html or not html.strip():
        return

    try:
        from lxml import html as lhtml
    except ImportError:
        # lxml is a transitive dep of python-docx — should never be missing.
        run = paragraph.add_run(re.sub(r"<[^>]+>", "", html))
        _apply_run_fmt(run, _Fmt(), default_family, default_size)
        return

    root = lhtml.fragment_fromstring(html.strip(), create_parent="div")
    _collapse_html_whitespace(root)

    state: _WalkState = {
        "anchor": paragraph,
        "current": paragraph,
        "first_used": False,
        "parent_obj": paragraph._parent,
    }
    blk = _BlockFmt()  # RTL=True default

    if root.text and root.text.strip():
        p = _state_new_paragraph(state)
        _apply_block_fmt(p, blk)
        _emit_text_into_paragraph(p, root.text, _Fmt(), default_family, default_size)

    for child in root:
        if not isinstance(child.tag, str):
            # Comment / PI at body level: pagebreak comments become real
            # page breaks; anything else is skipped (tail still handled).
            if "pagebreak" in ((child.text or "").lower()):
                _emit_page_break(state)
            if child.tail and child.tail.strip():
                cur = state["current"]
                if not hasattr(cur, "add_run"):
                    p = _state_new_paragraph(state)
                    _apply_block_fmt(p, blk)
                else:
                    p = cur
                _emit_text_into_paragraph(p, child.tail, _Fmt(), default_family, default_size)
            continue
        ctag = child.tag.lower()
        if ctag in (_BLOCK_TAGS | {"ul", "ol", "table"}):
            _walk_block(child, blk, state, default_family, default_size)
        else:
            p = _state_new_paragraph(state)
            _apply_block_fmt(p, blk)
            attrs: dict[str, str] = dict(child.attrib) if hasattr(child, "attrib") else {}
            fmt0 = _merge_fmt_from_attrs(_Fmt(), ctag, attrs)
            sub: _WalkState = {
                "anchor": p,
                "current": p,
                "first_used": True,
                "parent_obj": state["parent_obj"],
            }
            _walk_inline(child, p, fmt0, blk, sub, default_family, default_size)
        if child.tail and child.tail.strip():
            cur = state["current"]
            if not hasattr(cur, "add_run"):
                p = _state_new_paragraph(state)
                _apply_block_fmt(p, blk)
            else:
                p = cur
            _emit_text_into_paragraph(p, child.tail, _Fmt(), default_family, default_size)
