"""DOCX template-filling engine — thin dispatcher over `docx_render.render`.

This replaces v3.5.4's 900-line `TemplateFiller` (with 16 hand-coded
`fill_*` methods each knowing specific cell coordinates) with a dispatcher
that looks up a per-form template + optional post-process hook and calls
the generic `render()` function.

## Public contract

    DocxEngine(templates_dir: Path | str)
        .fill(form_type: str, data: Mapping, output_path) -> Path
        .stamp_ref_number(docx_path, ref_number, style) -> bool   # @staticmethod

Adding a new form is now:
  1. Drop the tokenized DOCX into `templates_dir`.
  2. Register `(form_type, template_filename)` in `_FORM_REGISTRY` below.
  3. If the form needs post-processing (behind-text sig, dotted-line
     fallback, etc.), add a `post_process` callable.

No new code per form, no cell-coordinate maps, no fill_* methods. Tokens
in the template drive everything.

## Data convention (v3-compatible)

Callers pass v3-shaped dicts:
  * `sig1_path` (manager), `sig2_path` (employee), `submitter_sig_path`
  * Other keys named exactly as v3 used them.

The dispatcher renames v3 keys to v4 token names internally:
  * `sig1_path` → `manager_sig_path`  → token `{{ manager_sig }}`
  * `sig2_path` → `employee_sig_path` → token `{{ employee_sig }}`

So endpoints / services can keep their v3 data dicts unchanged.
"""

from __future__ import annotations

import contextlib
import logging
import re
from collections.abc import Callable, Mapping, Sequence
from datetime import datetime
from pathlib import Path
from typing import Any, ClassVar

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.exceptions import PackageNotFoundError
from docx.shared import Pt, RGBColor

from app.core._docx_helpers import (
    fill_image_behind_text_in_paragraph,
    float_inline_images_in_cell,
    replace_paragraph_text,
)
from app.core.constants import (
    ARABIC_WEEKDAYS,
    DEFAULT_MANAGER_NAME,
    DEFAULT_MANAGER_TITLE,
    PROJECT_LOCATION,
    STAMP_STYLE_HEADER,
    STAMP_STYLE_TOP_RIGHT,
    STAMP_STYLE_WATERMARK,
    TEMPLATE_FILES,
)
from app.core.docx_render import render
from app.core.signature_render import DEFAULT_SIG_BOLDNESS, DEFAULT_SIG_SIZE_MM

log = logging.getLogger(__name__)

_ARIAL = "Arial"
_CALIBRI = "Calibri"  # HugeRTE editor default body font
_TODAY_FMT = "%d/%m/%Y"

# Forms whose page-1 letterhead logo owns the TOP-LEFT corner, so the Aztec ref
# code is stamped TOP-RIGHT instead (verified collision-free + decodable by the
# 2026-06-23 placement audit — scripts/aztec_collision_audit.py). Everyone else
# gets top-left, the proven margin spot. These also carry a distinct first-page
# header, which is why the code is written there (see stamp_aztec_code).
_AZTEC_TOP_RIGHT: frozenset[str] = frozenset(
    {
        "General Book",
        "Warning Form",
        "Administrative Leave Form",
        "Leave Permit Form",
    }
)


def aztec_corner_for(template_id: str) -> str:
    """Page-1 corner for the Aztec ref code: ``top-right`` for the letterhead
    forms (logo owns the top-left), ``top-left`` for everyone else."""
    return "top-right" if template_id in _AZTEC_TOP_RIGHT else "top-left"


# Forms that cannot carry a scannable page-1 code. Empty today — every form
# fits one (the audit confirmed all 17). Kept as the single source of truth so a
# FUTURE form/service with no clear corner can be listed here and the Services
# tile will flag it automatically (operator decision 2026-06-23).
_NO_CODE_FORMS: frozenset[str] = frozenset()


def template_has_code(template_id: str) -> bool:
    """Whether a committed *template_id* document carries a scannable page-1 ref
    code. True for all forms today; False only for any listed in _NO_CODE_FORMS."""
    return template_id not in _NO_CODE_FORMS


# --- Data adapters --------------------------------------------------------
# Rename v3 keys to v4 token names. Done in one place rather than touching
# every caller — v3 data dicts flow through unchanged.


def _adapt_common(data: dict[str, Any]) -> dict[str, Any]:
    """Apply universal v3→v4 key renames + sensible defaults."""
    out = dict(data)
    # Signature key normalisation.
    if "sig1_path" in out:
        out["manager_sig_path"] = out.pop("sig1_path")
    if "sig2_path" in out:
        out["employee_sig_path"] = out.pop("sig2_path")
    # `doc_selections` is exposed to the frontend as a single `select` field
    # but the docx `check()` helper expects a Mapping keyed by option name.
    # Coerce a bare string into the single-key dict shape v3 used.
    sel = out.get("doc_selections")
    if isinstance(sel, str):
        out["doc_selections"] = {sel: True} if sel else {}
    elif isinstance(sel, list):
        out["doc_selections"] = {s: True for s in sel if isinstance(s, str)}
    # Defaults the template expects.
    today = datetime.now().strftime(_TODAY_FMT)
    out.setdefault("today", out.get("date") or today)
    out.setdefault("date", out.get("today") or today)
    out.setdefault("request_date", out.get("today") or today)
    out.setdefault(
        "any_other_checked",
        bool(str(out.get("please_specify", "") or "").strip()),
    )
    out.setdefault("location", PROJECT_LOCATION)
    out.setdefault("manager_name", "")
    out.setdefault("manager_title", DEFAULT_MANAGER_TITLE)
    out["_sig_size_mm"] = min(int(out.get("_sig_size_mm", 32)), 32)
    return out


def _adapt_resignation_letter(data: dict[str, Any]) -> dict[str, Any]:
    """Split today into day/month/year for the body cell's three X marks,
    and route purpose_plain → reason."""
    out = _adapt_common(data)
    today_str = out.get("today") or datetime.now().strftime(_TODAY_FMT)
    try:
        dt = datetime.strptime(today_str, _TODAY_FMT)
    except ValueError:
        dt = datetime.now()
    out["today"] = today_str
    out["day"] = dt.strftime("%d")
    out["month"] = dt.strftime("%m")
    out["year"] = dt.strftime("%Y")
    out["reason"] = (out.get("purpose_plain") or out.get("reason") or "").strip()
    return out


def _adapt_leave_undertaking(data: dict[str, Any]) -> dict[str, Any]:
    """Bottom block carries today's date only if a submitter is picked."""
    out = _adapt_common(data)
    out["submitter_date"] = out.get("today", "") if out.get("submitter_name") else ""
    return out


def _adapt_admin_leave(data: dict[str, Any]) -> dict[str, Any]:
    """Build the Administrative Leave Form's date-range token from the operator's
    ``start_date`` / ``end_date`` fields.

    The 301-005 template has a ``{{ leave_date_range }}`` token (the "تاريخ
    الإجازة" row) but the operator form only sends ``start_date`` / ``end_date``
    — so the dates vanished. v3 (``fill_admin_leave`` caller, gssg_manager.pyw)
    composed the range as ``f"{start}  -  {end}"`` (two spaces either side of the
    dash); we replicate that exactly. When only one bound is present, fall back to
    whichever is set.

    ``admin_leaves_this_month`` is DB-dependent (a per-employee count of this
    month's admin leaves) and is resolved upstream in
    ``document_service._build_template_data`` — see that module — so this pure
    adapter only defaults it to "" for the no-employee / test paths.
    """
    out = _adapt_common(data)
    if not (out.get("leave_date_range") or "").strip():
        start = (out.get("start_date") or "").strip()
        end = (out.get("end_date") or "").strip() or start
        if start and end:
            out["leave_date_range"] = f"{start}  -  {end}"
        else:
            out["leave_date_range"] = start or end
    out.setdefault("admin_leaves_this_month", "")
    return out


def _adapt_material_request(data: dict[str, Any]) -> dict[str, Any]:
    """Default `project_site` to the site code (v3 `fill_material_request`
    filled this cell with `PROJECT_LOCATION` when the operator left it blank).

    The template has a `{{ project_site }}` token but nothing in the form
    fields / common adapter supplies it, so it rendered blank. v3 used
    ``data.get('project_site', PROJECT_LOCATION)``.
    """
    out = _adapt_common(data)
    if not (out.get("project_site") or "").strip():
        out["project_site"] = PROJECT_LOCATION
    return out


def _adapt_resignation_declaration(data: dict[str, Any]) -> dict[str, Any]:
    """Use Arabic weekday + date for the header row."""
    out = _adapt_common(data)
    today_str = out.get("today") or datetime.now().strftime(_TODAY_FMT)
    try:
        dt = datetime.strptime(today_str, _TODAY_FMT)
    except ValueError:
        dt = datetime.now()
    out["today"] = today_str
    out["weekday_ar"] = ARABIC_WEEKDAYS[dt.weekday()]
    return out


def _adapt_employee_clearance(data: dict[str, Any]) -> dict[str, Any]:
    """Fan a frontend `clearance_table` payload into the two dicts the Jinja
    `clearance()` helper expects (`clearance_marks` + `clearance_remarks`).

    The v4 React form ships one composite field; v3 fixtures already pass the
    two flat dicts. Both shapes are accepted.
    """
    out = _adapt_common(data)
    table = out.pop("clearance_table", None)
    if isinstance(table, dict):
        marks = table.get("clearance_marks") or table.get("marks") or {}
        remarks = table.get("clearance_remarks") or table.get("remarks") or {}
        if marks:
            out.setdefault("clearance_marks", marks)
        if remarks:
            out.setdefault("clearance_remarks", remarks)
    # {{ employee_sig }} / {{ manager_sig }} repeat in every ~32mm-wide,
    # ~7.4mm-tall table row here — the global signature width (default 45mm)
    # overflows each cell and blows up the row. Cap it to sit inside the cell.
    _CLEARANCE_SIG_MM = 24
    global_mm = out.get("_sig_size_mm", _CLEARANCE_SIG_MM)
    out["_sig_size_mm"] = min(int(global_mm), _CLEARANCE_SIG_MM)
    return out


def _normalize_cc(value: Any) -> str:
    """Normalize a ``cc`` field payload into a single comma-joined string.

    Accepts both shapes:
      * ``list[str]`` — the new multi-picker frontend sends a list of selected
        recipient names.
      * ``str``       — the legacy textarea sent a single string with commas or
        newlines as separators.

    In either case the result is the cleaned, deduped, comma-joined form (e.g.
    ``"Ahmed Hassan, Saeed Rashed, Mohammed Ali"``) the ``{{ cc }}`` template
    token expects. Whitespace is trimmed, blanks are dropped, and duplicates are
    removed while preserving first-seen order.
    """
    if value is None:
        return ""
    if isinstance(value, list):
        parts: list[str] = [str(item).strip() for item in value]
    elif isinstance(value, str):
        # Legacy textarea: comma- or newline-separated.
        parts = [chunk.strip() for chunk in value.replace("\n", ",").split(",")]
    else:
        return ""
    seen: set[str] = set()
    cleaned: list[str] = []
    for part in parts:
        if part and part not in seen:
            seen.add(part)
            cleaned.append(part)
    return ", ".join(cleaned)


def _adapt_general_book(data: dict[str, Any]) -> dict[str, Any]:
    """Date format is `dd-mm-yyyy` for the General Book (not v3's default `dd/mm/yyyy`).

    ``recipient_id`` → ``recipient_name`` resolution happens upstream on the
    request session (``document_service`` / ``render_signed_pdf``, before
    ``engine.fill``); this adapter only defaults the token to "" so the template
    renders cleanly. It does NOT open its own DB session.

    ``cc`` (multi-picker selection or legacy textarea string) is normalized to a
    single comma-joined string for the ``{{ cc }}`` token — see ``_normalize_cc``.

    ``manager_name`` is whatever ``manager_override.apply`` produced (English
    name preferred, Arabic fallback — see ``core.manager_override``); when no
    manager is picked, falls back to ``DEFAULT_MANAGER_NAME`` so the template
    still has a value to render.
    """
    out = _adapt_common(data)
    # General Book overrides the date format and discards the default the common
    # adapter applied.
    out["date"] = data.get("date") or datetime.now().strftime("%d-%m-%Y")
    out.setdefault("subject", "")
    out.setdefault("body", "")
    # Preserve raw body HTML (when the service layer threaded it through) for
    # the post-process that rebuilds real Word tables from <table> elements.
    # Plain assignment so an explicit upstream empty string still wins.
    out["body_html"] = data.get("body_html", "")
    # Normalize cc to a comma-joined string first (keeps _normalize_cc's
    # contract intact for tests / other callers), then re-join with "\n" for
    # the template token. docxtpl's `resolve_listing` rewrites every "\n"
    # inside the rendered text into `</w:t><w:br/><w:t>` — so each recipient
    # lands on its own line in the document.
    #
    # Round 2 — Fix C: the bare-string path produced one `<w:r>` per line,
    # only the first carrying the template's `<w:rPr>` (with `<w:sz w:val="28"/>`
    # for 14pt). Subsequent runs fell back to the default font-size. We keep
    # the legacy "\n"-join (so existing string-equality tests still pass),
    # then a post-process stamps `<w:sz w:val="28"/>` on every CC run.
    cc_str = _normalize_cc(out.get("cc"))
    if not cc_str:
        out["cc"] = ""
    else:
        parts = [p.strip() for p in cc_str.split(",") if p.strip()]
        out["cc"] = "\n".join(parts)
    # Default manager_name when no picker / override resolved it — the template
    # never has to render a blank signature line.
    if not (out.get("manager_name") or "").strip():
        out["manager_name"] = DEFAULT_MANAGER_NAME
    # Submitter G-number for the footer (template token: {{ submitter_g }}).
    # document_service injects this from the authenticated caller; default to
    # "" so the Jinja {% if %} guard can hide the line cleanly when missing.
    out.setdefault("submitter_g", "")
    # recipient_id → recipient_name is resolved upstream on the REQUEST session
    # (document_service, before engine.fill) — see that module's canonical path.
    # We must NOT open a second SessionLocal() here: the generation path holds a
    # `BEGIN IMMEDIATE` write lock, and a second SQLite connection would risk a
    # "database is locked" / self-deadlock.
    out.setdefault("recipient_name", "")
    return out


# --- Post-process hooks ---------------------------------------------------


def _pp_resignation_letter(doc: Any, ctx: dict[str, Any]) -> None:
    """Overwrite the body cell's dotted-line paragraphs with `reason`.

    v3's fallback when the reason was longer than the inline X slot: find
    paragraphs that look like "...................................." in
    the body cell and replace the first one with the reason, clearing the
    rest. We do this unconditionally so reasons of any length render in
    the same place (just below "نظراً للأسباب التالية:").
    """
    reason = (ctx.get("reason") or "").strip()
    if not reason:
        return
    try:
        body_cell = doc.tables[1].rows[0].cells[0]
    except IndexError:
        return
    dotted = [
        p
        for p in body_cell.paragraphs
        if (s := p.text.strip()) and s.count(".") >= 20 and len(s) > 30
    ]
    if not dotted:
        return
    replace_paragraph_text(dotted[0], reason, size=11)
    for extra in dotted[1:]:
        replace_paragraph_text(extra, "", size=11)


def _find_body_paragraph(doc: Any, text: str) -> Any | None:
    """First body-level paragraph whose stripped text equals *text* (non-empty)."""
    text = text.strip()
    if not text:
        return None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    return None


def _finalize_manager_block(doc: Any, ctx: dict[str, Any]) -> None:
    """Left-align the manager name + title, OR blank the standalone designation
    when no manager is selected.

    The name/title ship as bidi body paragraphs with no explicit ``<w:jc>`` — a
    bidi paragraph's natural start edge is the RIGHT, so they hugged the right
    margin (Leave Permit) or split (Admin Leave, where the title had no bidi and
    rendered left while the name stayed right). Operators want a consistent
    LEFT-aligned block, so we drop ``<w:bidi/>`` and set an explicit left
    ``<w:jc>``.

    When no manager is picked, ``manager_name`` is blank but ``manager_title``
    still carries the ``DEFAULT_MANAGER_TITLE`` designation (set by
    ``_adapt_common``) — which then floated alone at the foot of the form. With
    no signer there is no designation to show, so we clear that paragraph.
    """
    from docx.oxml.ns import qn as _qn

    name = str(ctx.get("manager_name", "") or "").strip()
    title = str(ctx.get("manager_title", "") or "").strip()

    if not name:
        # No manager selected → don't leave the designation stranded on its own.
        if title and (title_para := _find_body_paragraph(doc, title)) is not None:
            replace_paragraph_text(title_para, "")
        return

    for value in (name, title):
        paragraph = _find_body_paragraph(doc, value)
        if paragraph is None:
            continue
        pPr = paragraph._p.get_or_add_pPr()
        for bidi in pPr.findall(_qn("w:bidi")):
            pPr.remove(bidi)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _place_manager_sig_above_name(doc: Any, ctx: dict[str, Any]) -> None:
    """Float the manager signature behind text on the signature line just above
    the manager name at the foot of the form (where the operator expects it),
    left-aligned to sit over the left-aligned name block.

    No-op when there is no signature or no manager name (the whole signer block
    is hidden in that case)."""
    sig_path = ctx.get("manager_sig_path")
    name = str(ctx.get("manager_name", "") or "").strip()
    if not sig_path or not name:
        return
    # Anchor to the blank paragraph directly above the name (the signature-line
    # gap); the float rests on its baseline and rises up into that gap.
    body = list(doc.paragraphs)
    idx = next((i for i, p in enumerate(body) if p.text.strip() == name), None)
    if idx is None:
        return
    anchor = body[idx - 1] if idx > 0 else body[idx]
    size_mm = float(ctx.get("_sig_size_mm", DEFAULT_SIG_SIZE_MM))
    boldness = int(ctx.get("_sig_boldness", DEFAULT_SIG_BOLDNESS))
    fill_image_behind_text_in_paragraph(
        anchor, sig_path, width_inches=size_mm / 25.4, dilate_radius_px=boldness
    )


_TATWEEL_WS = re.compile(r"[\sـ]+")  # whitespace + Arabic tatweel (U+0640)


def _norm_name(s: str) -> str:
    return _TATWEEL_WS.sub("", s or "")


_CC_LABEL_NORM = "نسخةإلى"  # "نسخة إلى" normalized — CC lines can quote the manager


def stamp_signature_above_name(
    docx_path: Path | str,
    sig_path: str,
    names: Sequence[str],
    *,
    size_mm: float = DEFAULT_SIG_SIZE_MM,
    boldness: int = DEFAULT_SIG_BOLDNESS,
) -> bool:
    """Float *sig_path* above the closing-name line of an ALREADY-RENDERED docx.

    Word-authored books have no Jinja tokens left, so the anchor is textual —
    normalized (whitespace + tatweel stripped; hand-made templates stretch
    names with tatweel). Search order:

    1. exact-equality match on a body paragraph (beats CC lines that merely
       QUOTE the manager's name after the closing block);
    2. containment match, skipping "نسخة إلى" lines;
    3. table-cell paragraphs (Word-paste-into-tables letters keep everything
       inside cells — ``doc.paragraphs`` never sees them);
    4. last non-empty paragraph (body, then tables).

    Body matches anchor on the paragraph ABOVE the name (the signature gap);
    table matches anchor on the matched paragraph itself (the float rises up).
    Returns False when the signature file is unusable or no anchor exists —
    callers must treat that as a FAILURE, not a soft skip (a "signed" paper
    without a visible signature is the defect this function exists to prevent).
    """
    if not sig_path or not Path(sig_path).is_file():
        return False
    doc = Document(str(docx_path))
    paras = list(doc.paragraphs)
    table_paras = [
        p for tbl in doc.tables for row in tbl.rows for c in row.cells for p in c.paragraphs
    ]
    wanted = [_norm_name(n) for n in names if n and _norm_name(n)]

    def _find(pool: list[Any], *, exact: bool) -> int | None:
        for i in range(len(pool) - 1, -1, -1):
            text = _norm_name(pool[i].text)
            if not text or text.startswith(_CC_LABEL_NORM):
                continue
            if exact:
                if any(w == text for w in wanted):
                    return i
            elif any(w in text for w in wanted):
                return i
        return None

    anchor = None
    idx = _find(paras, exact=True)
    if idx is None:
        idx = _find(paras, exact=False)
    if idx is not None:
        anchor = paras[idx - 1] if idx > 0 else paras[idx]
    else:
        t_idx = _find(table_paras, exact=True)
        if t_idx is None:
            t_idx = _find(table_paras, exact=False)
        if t_idx is not None:
            anchor = table_paras[t_idx]  # float rises up from the name's own line
    if anchor is None:
        last = next((p for p in reversed(paras) if p.text.strip()), None) or next(
            (p for p in reversed(table_paras) if p.text.strip()), None
        )
        anchor = last
    if anchor is None:
        return False
    placed = fill_image_behind_text_in_paragraph(
        anchor, sig_path, width_inches=size_mm / 25.4, dilate_radius_px=boldness
    )
    if placed:
        doc.save(str(docx_path))
    return bool(placed)


def _pp_leave_permit(doc: Any, ctx: dict[str, Any]) -> None:
    """Float the manager signature on the signature line above the name and
    left-align (or, with no manager, blank) the name/title block."""
    _place_manager_sig_above_name(doc, ctx)
    _finalize_manager_block(doc, ctx)


def _pp_admin_leave(doc: Any, ctx: dict[str, Any]) -> None:
    """Same signer-block layout as Leave Permit, plus clear the literal "x"
    placeholder in the branch-manager التوقيع value cell so it reads as a clean,
    empty manual-signature line."""
    try:
        branch_sig_cell = doc.tables[2].rows[3].cells[1]
        for para in branch_sig_cell.paragraphs:
            if para.text and "x" in para.text.lower():
                replace_paragraph_text(para, "")
    except IndexError:
        log.warning("Admin Leave: branch التوقيع cell (t2 r3 c1) missing")
    _place_manager_sig_above_name(doc, ctx)
    _finalize_manager_block(doc, ctx)


def _pp_material_request(doc: Any, ctx: dict[str, Any]) -> None:
    """Float the signature images behind text so they don't grow their row and
    push the form onto a second page (SC-0425 regression). The ``{{ *_sig }}``
    tokens place the images inline; we convert them in place to behind-text
    floats, bottom-aligned so each signature rests on its cell's line and
    overflows UP into the empty cell above (the unused Contact cell) instead of
    hanging down across the section divider.

    Cells (table 4): r2/c3 = employee "Signature", r5/c3 = manager
    "Signature". No-op for whichever signature wasn't embedded."""
    try:
        sig_cells = [doc.tables[4].rows[2].cells[3], doc.tables[4].rows[5].cells[3]]
    except IndexError:
        log.warning("Material Request: signature cell missing — float skipped")
        return
    for cell in sig_cells:
        float_inline_images_in_cell(cell, bottom_align=True)


def _pp_general_book_cc(doc: Any, ctx: dict[str, Any]) -> None:
    """Render each CC recipient as its own right-aligned, bulleted line, each
    prefixed with the "نسخة إلى:" label (not just the first recipient).

    The template renders ``نسخة إلى: {{ cc }}`` as one paragraph with the
    recipients separated by line breaks, so only the first line carried the
    label. Operators want every recipient on its own bullet with the label
    repeated:

        • نسخة إلى: HR Department
        • نسخة إلى: Operations Wing

    We rebuild the paragraph's runs from ``ctx["cc"]`` (the adapter's
    ``"\\n"``-joined recipient list), preserving the template's CC font size.

    Alignment: the template's CC paragraph carries ``<w:pStyle
    w:val="ListParagraph"/>`` (a left indent) plus an explicit ``<w:jc
    w:val="right"/>`` placed AFTER ``<w:rPr>`` — invalid OOXML order, so Word
    ignores it and the list indent pulls the line left. Verified against Word's
    PDF export: the lines that render hard-right (date, salutation, closing)
    carry ``<w:bidi/>`` and **no** ``<w:jc>`` — a bidi paragraph's natural start
    edge is the right, which right-aligns mixed Arabic/Latin content reliably
    (an explicit ``jc="right"`` does not). So we drop the list style, any indent,
    and every ``<w:jc>``, keeping only ``<w:bidi/>``.
    """
    from docx.oxml.ns import qn

    from app.core.arabic_rtl import stamp_paragraph, stamp_run

    recipients = [r.strip() for r in str(ctx.get("cc", "") or "").split("\n") if r.strip()]

    for paragraph in doc.paragraphs:
        if "نسخة إلى" not in (paragraph.text or ""):
            continue
        pPr = paragraph._p.get_or_add_pPr()
        # Drop the ListParagraph style so its left indent can't reassert.
        for pStyle in pPr.findall(qn("w:pStyle")):
            pPr.remove(pStyle)
        # Drop every <w:jc>; a bidi paragraph's natural start edge is the right,
        # which right-aligns mixed Arabic/Latin content reliably (an explicit
        # jc="right" does not — see docstring).
        for jc in pPr.findall(qn("w:jc")):
            pPr.remove(jc)
        # Remove any left/right indent (drops <w:ind>).
        for ind in pPr.findall(qn("w:ind")):
            pPr.remove(ind)
        # Keep the paragraph RTL (adds <w:bidi/> only if missing).
        stamp_paragraph(paragraph)

        if recipients:
            # Preserve the template CC font size/family, then rebuild as one
            # bulleted, labelled line per recipient (label repeated on each).
            first = paragraph.runs[0] if paragraph.runs else None
            size = first.font.size if first is not None else None
            family = (first.font.name if first is not None else None) or _CALIBRI
            for r in list(paragraph.runs):
                r._element.getparent().remove(r._element)
            for i, rec in enumerate(recipients):
                run = paragraph.add_run(f"• نسخة إلى: {rec}")
                if size is not None:
                    run.font.size = size
                run.font.name = family
                stamp_run(run, family)
                if i < len(recipients) - 1:
                    run.add_break()
        break


def _pp_general_book(doc: Any, ctx: dict[str, Any]) -> None:
    """Render the WHOLE General Book body via the RTL-aware HTML renderer.

    The body field is a single ``{{ body }}`` token in the General Book
    template. The service layer renders a content-independent sentinel
    (``GENERAL_BOOK_BODY_SENTINEL``) for that token and threads the raw HugeRTE
    HTML through ``ctx["body_html"]``. This post-process locates the sentinel
    paragraph (the anchor), clears it, and hands the full body — narrative
    (headings, bold/italic/color/size, alignment, lists, links, images) AND any
    real Word tables, in document order — to ``arabic_rtl.html_to_docx``. That
    is the single lossless renderer (PART A); nothing about the body is
    flattened here.

    Calibri/12pt is the HugeRTE editor's default body font, so plain text
    renders at that; inner spans/headings override per their inline styling.

    Signature images are not auto-embedded — operators hand-sign or paste
    images manually.
    """
    from app.core.arabic_rtl import html_to_docx

    # CC right-alignment fix runs first — independent of the body.
    _pp_general_book_cc(doc, ctx)
    _format_general_book_ref_line(doc)

    anchor = _find_general_book_body_anchor(doc)
    if anchor is None:
        log.warning("General Book: no body anchor paragraph found — body not rendered")
        return

    body_html = ctx.get("body_html", "") or ""
    if not body_html.strip():
        # No body content — just clear the sentinel so it never shows.
        for run in list(anchor.runs):
            run.text = ""
        return

    html_to_docx(body_html, anchor, default_family=_CALIBRI, default_size=12.0)


def _format_general_book_ref_line(doc: Any) -> None:
    """Keep the Arabic label RTL and the rendered reference in LTR order."""
    from app.core.arabic_rtl import stamp_paragraph, stamp_run

    for paragraph in doc.paragraphs:
        match = re.match(r"^\s*الرقم\s*[:：]\s*(.+?)\s*$", paragraph.text or "")  # noqa: RUF001
        if not match:
            continue

        for run in list(paragraph.runs):
            run._element.getparent().remove(run._element)
        stamp_paragraph(paragraph)
        label = paragraph.add_run("الرقم: ")
        value = paragraph.add_run(match.group(1))
        for run in (label, value):
            run.font.name = _CALIBRI
            run.font.size = Pt(16)
            run.font.italic = True
            stamp_run(run, _CALIBRI)
        value.font.rtl = False
        break


def _find_general_book_body_anchor(doc: Any) -> Any | None:
    """Return the paragraph carrying the body sentinel, or ``None``.

    Matches by substring so the anchor is found regardless of the surrounding
    body content (the sentinel is content-independent — see
    ``document_service.GENERAL_BOOK_BODY_SENTINEL``).
    """
    from app.services.document_service import GENERAL_BOOK_BODY_SENTINEL

    for p in doc.paragraphs:
        if GENERAL_BOOK_BODY_SENTINEL in (p.text or ""):
            return p
    return None


def _postprocess_general_book_footer(docx_path: str | Path) -> None:
    """Sync ``word/footer2.xml`` (default / pages 2+) with ``word/footer3.xml``
    (page 1) so the submitter G-number + letterhead repeat on every page.

    The General Book template has ``<w:titlePg/>`` set, so page 1 uses
    ``footer3.xml`` (which carries ``{{ submitter_g }}`` + letterhead) and
    pages 2+ use ``footer2.xml`` (only a page number + logo). Operators want
    the same footer on every page. We replace footer2's whole XML body with
    footer3's content — both share the Office XML ``<w:ftr>`` schema, and
    ``document.xml.rels`` still points at ``footer2.xml`` (the file name is
    unchanged; only its contents are now footer3's).

    The zipfile module can't replace entries in place, so we copy through a
    sibling temp file and atomically replace. No-op (logged) if either footer
    is missing — the template might not always carry both.
    """
    import shutil
    import zipfile

    p = Path(docx_path)
    tmp = p.with_suffix(p.suffix + ".__footer_patch__")
    try:
        with zipfile.ZipFile(p, "r") as zin:
            names = zin.namelist()
            if "word/footer3.xml" not in names or "word/footer2.xml" not in names:
                log.warning(
                    "General Book footer post-process: footer2/footer3 missing in %s",
                    p,
                )
                return
            footer3 = zin.read("word/footer3.xml")
            with zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
                for name in names:
                    data = zin.read(name)
                    if name == "word/footer2.xml":
                        data = footer3
                    zout.writestr(name, data)
        shutil.move(str(tmp), str(p))
    except (OSError, zipfile.BadZipFile) as exc:
        log.warning("General Book footer post-process failed for %s: %s", p, exc)
        if tmp.exists():
            with contextlib.suppress(OSError):
                tmp.unlink()


# --- Form registry --------------------------------------------------------


_FORM_REGISTRY: dict[str, dict[str, Any]] = {
    "Acknowledgment Form": {"adapter": _adapt_common, "post_process": None},
    "Salary Transfer Request": {"adapter": _adapt_common, "post_process": None},
    "Salary Deduction Form": {"adapter": _adapt_common, "post_process": None},
    "Violation Form": {"adapter": _adapt_common, "post_process": None},
    "Employee Clearance Form": {"adapter": _adapt_employee_clearance, "post_process": None},
    "HR Request Form": {"adapter": _adapt_common, "post_process": None},
    "Resignation Declaration": {
        "adapter": _adapt_resignation_declaration,
        "post_process": None,
    },
    "Resignation Letter": {
        "adapter": _adapt_resignation_letter,
        "post_process": _pp_resignation_letter,
    },
    "Leave Undertaking": {"adapter": _adapt_leave_undertaking, "post_process": None},
    "Material Request Form": {
        "adapter": _adapt_material_request,
        "post_process": _pp_material_request,
    },
    "Leave Application Form": {"adapter": _adapt_common, "post_process": None},
    "Passport Release Form": {"adapter": _adapt_common, "post_process": None},
    "Duty Resumption Form": {"adapter": _adapt_common, "post_process": None},
    "General Book": {"adapter": _adapt_general_book, "post_process": _pp_general_book},
    "Leave Permit Form": {"adapter": _adapt_common, "post_process": _pp_leave_permit},
    "Administrative Leave Form": {
        "adapter": _adapt_admin_leave,
        "post_process": _pp_admin_leave,
    },
    "Warning Form": {"adapter": _adapt_common, "post_process": None},
    # Multi-employee passport list — rows come from fields["items"], rendered by
    # the item() Jinja global; no per-employee binding (admin category).
    "Passport Release List": {"adapter": _adapt_common, "post_process": None},
    # Report: no-classification, no-ref General Book paper; reuses same adapter
    # and post-process as General Book (body + footer pipeline).
    "Report": {"adapter": _adapt_general_book, "post_process": _pp_general_book},
}


# --- Engine ---------------------------------------------------------------


def _normalize_form_type(form_type: str) -> str:
    """Strip the ``" - Arabic"`` half off a UI label."""
    return form_type.split(" - ")[0] if " - " in form_type else form_type


class DocxEngine:
    """Thin dispatcher: form_type → template + adapter + optional post_process."""

    _REGISTRY: ClassVar[dict[str, dict[str, Any]]] = _FORM_REGISTRY

    def __init__(self, templates_dir: Path | str) -> None:
        self.templates_dir = Path(templates_dir)

    def fill(
        self,
        form_type: str,
        data: Mapping[str, Any],
        output_path: Path | str,
    ) -> Path:
        """Render `form_type` with `data` and save to `output_path`.

        Raises ``KeyError`` if `form_type` is unknown and ``FileNotFoundError``
        if the template DOCX is missing from `templates_dir`.
        """
        short = _normalize_form_type(form_type)
        if short not in self._REGISTRY:
            raise KeyError(f"Unknown form_type: {form_type!r}")
        if short not in TEMPLATE_FILES:
            raise KeyError(f"No template registered for {form_type!r}")

        spec = self._REGISTRY[short]
        template = self.templates_dir / TEMPLATE_FILES[short]
        if not template.exists():
            raise FileNotFoundError(f"Template for {form_type!r} not found: {template}")

        adapter: Callable[[dict[str, Any]], dict[str, Any]] = spec["adapter"]
        post_process = spec.get("post_process")
        prepared = adapter(dict(data))
        return render(template, prepared, Path(output_path), post_process=post_process)

    def fill_general_book_path(
        self,
        template_path: Path,
        data: Mapping[str, Any],
        output_path: Path | str,
        *,
        sandboxed: bool = False,
        strict: bool = False,
    ) -> Path:
        """Render an arbitrary docx through the General Book adapter +
        post-process — for library boilerplate templates (untrusted; render
        sandboxed). Raises FileNotFoundError if the file is gone."""
        if not Path(template_path).exists():
            raise FileNotFoundError(template_path)
        spec = self._REGISTRY["General Book"]
        adapter: Callable[[dict[str, Any]], dict[str, Any]] = spec["adapter"]
        prepared = adapter(dict(data))
        return render(
            Path(template_path),
            prepared,
            Path(output_path),
            post_process=spec.get("post_process"),
            strict=strict,
            sandboxed=sandboxed,
        )

    @staticmethod
    def stamp_ref_number(
        docx_path: Path | str,
        ref_number: str,
        style: str = STAMP_STYLE_HEADER,
    ) -> bool:
        """Stamp a reference number into the document header.

        Style strings come from `core.constants.STAMP_STYLES`. Returns False
        on any docx error — v3 silently swallowed errors here and we match
        that to avoid breaking downstream PDF export.
        """
        try:
            doc = Document(str(docx_path))
            section = doc.sections[0]
            header = section.header
            header.is_linked_to_previous = False
            p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            for run in p.runs:
                run.text = ""

            if style == STAMP_STYLE_TOP_RIGHT or style.startswith("Bold"):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                run = p.add_run(f"Ref: {ref_number}")
                run.font.size = Pt(11)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 51, 102)
            elif style == STAMP_STYLE_WATERMARK or style.startswith("Watermark"):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"  {ref_number}  ")
                run.font.size = Pt(28)
                run.font.bold = True
                run.font.color.rgb = RGBColor(200, 200, 200)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(f"Ref: {ref_number}")
                run.font.size = Pt(9)
                run.font.bold = True
                run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.name = _ARIAL

            doc.save(str(docx_path))
            return True
        except (OSError, ValueError, KeyError, PackageNotFoundError) as e:
            log.warning("stamp_ref_number failed for %s: %s", docx_path, e)
            return False

    @staticmethod
    def stamp_aztec_code(
        docx_path: Path | str,
        ref_number: str,
        *,
        corner: str = "top-left",
        size_cm: float = 1.1,
        margin_cm: float = 0.3,
    ) -> bool:
        """Stamp a small floating Aztec of *ref_number* into the page-1 header.

        Writes to the section's *first-page* header when it has a distinct
        titlePg header (the letterhead forms), else the primary header — so the
        code always lands on page 1. Page-anchored to *corner* (in front of
        content) so it never reflows the letterhead; accompanies — does not
        replace — the human-readable ``Ref:`` text stamp. Returns False on any
        error (logs a warning) so a code failure never breaks PDF export,
        matching ``stamp_ref_number``'s contract.
        """
        try:
            from docx.shared import Cm

            from app.core._docx_helpers import insert_floating_image_in_header
            from app.core.qr import make_aztec_png

            png = make_aztec_png(ref_number)
            doc = Document(str(docx_path))
            section = doc.sections[0]
            # The header that renders on page 1: first-page header for titlePg
            # forms (their logo lives there), else the primary header.
            if section.different_first_page_header_footer:
                header = section.first_page_header
            else:
                header = section.header
            header.is_linked_to_previous = False

            pw = int(section.page_width or 0)
            ph = int(section.page_height or 0)
            size = int(size_cm * 360000)
            margin = int(margin_cm * 360000)
            if corner == "top-right":
                x, y = pw - size - margin, margin
            elif corner == "bottom-right":
                x, y = pw - size - margin, ph - size - margin
            elif corner == "bottom-left":
                x, y = margin, ph - size - margin
            else:  # top-left
                x, y = margin, margin

            ok = insert_floating_image_in_header(header, png, x_emu=x, y_emu=y, size_emu=size)
            if ok:
                # Only the top-left code shares the header corner with the "Ref:"
                # text stamp; indent that paragraph clear of the (page-anchored)
                # code. Top-right forms keep the ref text untouched at top-left.
                if corner == "top-left" and header.paragraphs:
                    header.paragraphs[0].paragraph_format.left_indent = Cm(
                        margin_cm + size_cm + 0.25
                    )
                doc.save(str(docx_path))
            return ok
        except Exception as e:
            log.warning("stamp_aztec_code failed for %s: %s", docx_path, e)
            return False


# Suppress unused-import warning — replace_paragraph_text is intentionally
# re-exported for callers that still build documents directly (e.g., Phase 04
# services). Will be cleaner once those callers also move to render().
_ = replace_paragraph_text
