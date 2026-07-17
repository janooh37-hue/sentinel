"""book_text -- text extraction and Arabic normalisation for Books FTS.

``normalize_ar`` is the canonical normaliser: both the indexing path (Task 15)
and the search-query path (Task 16) must run text through this function so the
FTS index and the query share identical forms.

Usage::

    from app.core.book_text import build_search_text, docx_to_text, html_to_text, normalize_ar
"""

from __future__ import annotations

import re
from html.parser import HTMLParser
from pathlib import Path

# ---------------------------------------------------------------------------
# Arabic normalisation
# ---------------------------------------------------------------------------

# Tatweel (kashida) -- U+0640
_TATWEEL = "ـ"

# Alef variants: U+0623 U+0625 U+0622 U+0671 -> U+0627 (plain alef)
_ALEF_VARIANTS_RE = re.compile("[أإآٱ]")

# U+0649 (alef maqsura) -> U+064A (ya)
# U+0629 (taa marbuta)  -> U+0647 (ha)
_CHAR_MAP = str.maketrans({"ى": "ي", "ة": "\u0647"})

# Arabic diacritics (harakat) U+064B-U+065F -- stripped for cleaner corpus.
# The FTS tokenizer's remove_diacritics=2 also strips them, but normalising
# here means the stored search_text column is already diacritic-free.
_DIACRITIC_RE = re.compile("[ً-ٟ]")

_WHITESPACE_RE = re.compile(r"\s+")


def normalize_ar(text: str) -> str:
    """Normalise Arabic text for FTS indexing/querying.

    Transformations (in order):
    1. Strip tatweel (U+0640).
    2. Unify alef variants (U+0623/0625/0622/0671) to plain alef (U+0627).
    3. Alef maqsura (U+0649) -> ya (U+064A); taa marbuta (U+0629) -> ha (U+0647).
    4. Strip Arabic diacritics (harakat U+064B-U+065F).
    5. Collapse whitespace runs to a single space and strip.
    """
    text = text.replace(_TATWEEL, "")
    text = _ALEF_VARIANTS_RE.sub("\u0627", text)
    text = text.translate(_CHAR_MAP)
    text = _DIACRITIC_RE.sub("", text)
    text = _WHITESPACE_RE.sub(" ", text)
    return text.strip()


# ---------------------------------------------------------------------------
# DOCX text extraction
# ---------------------------------------------------------------------------


def docx_to_text(path: Path) -> str:
    """Extract plain text from a .docx file (paragraphs + table cells).

    Joins all parts with newlines so whitespace is predictable.
    """
    from docx import Document  # python-docx; always available (used throughout codebase)

    doc = Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        t = para.text
        if t:
            parts.append(t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                t = cell.text
                if t:
                    parts.append(t)

    return "\n".join(parts)


# ---------------------------------------------------------------------------
# HTML -> plain text
# ---------------------------------------------------------------------------


class _TextExtractor(HTMLParser):
    """Minimal HTMLParser subclass that collects text nodes."""

    def __init__(self) -> None:
        super().__init__()
        self._chunks: list[str] = []
        self._block_tags = frozenset(
            {"p", "div", "li", "tr", "h1", "h2", "h3", "h4", "h5", "h6", "blockquote", "br"}
        )

    def handle_data(self, data: str) -> None:
        self._chunks.append(data)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag.lower() in self._block_tags:
            self._chunks.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag.lower() in self._block_tags:
            self._chunks.append("\n")

    def result(self) -> str:
        text = "".join(self._chunks)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()


def html_to_text(html: str) -> str:
    """Strip HTML tags to plain text.

    Uses stdlib ``html.parser`` -- no new dependency (lxml is installed but
    stdlib is simpler and the output is identical for the tag set HugeRTE emits).
    Entity decoding is handled automatically by HTMLParser.
    """
    if not html or "<" not in html:
        return html or ""
    parser = _TextExtractor()
    parser.feed(html)
    return parser.result()


# ---------------------------------------------------------------------------
# Build search text
# ---------------------------------------------------------------------------


def build_search_text(*, subject: str | None, ref: str, body: str) -> str:
    """Concatenate subject + ref + body and normalise for FTS indexing.

    Falsy parts are skipped so empty strings do not bloat the corpus.
    """
    parts = [p for p in (subject, ref, body) if p]
    return normalize_ar("  ".join(parts))

