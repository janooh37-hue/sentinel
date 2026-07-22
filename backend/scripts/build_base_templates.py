"""Build the single canonical General Book table template.

The normal no-table document already exists as the "no template" path.  This
builder starts from that real document so the table version keeps the exact
letterhead, logos, page setup, closing, author block, and footer.
"""

from __future__ import annotations

import argparse
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.core.book_template_retokenize import retokenize_general_book, validate_book_template
from app.core.docx_engine import DocxEngine

_BACKEND = Path(__file__).resolve().parents[1]
_NAME = "base_table.docx"


def _build_table(path: Path) -> None:
    DocxEngine(_BACKEND / "templates").fill(
        "General Book",
        {
            "ref": "1/1/1",
            "date": "01-01-2026",
            "recipient_name": "",
            "subject": "",
            "body": "",
            "cc": "",
            "submitter_g": "G-0000",
        },
        path,
    )

    doc = Document(str(path))
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    headers = ("الرقم الوظيفي", "المسمى الوظيفي", "الاسم")
    for cell, label in zip(table.rows[0].cells, headers, strict=True):
        cell.text = label
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.rtl = True
    for cell in table.rows[1].cells:
        cell.text = "—"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.rtl = True

    # Put the grid before the standard closing, not after the footer text.
    closing = next((p for p in doc.paragraphs if "للتفضل بالعلم" in p.text), None)
    if closing is not None:
        closing._p.addprevious(table._tbl)
    doc.save(str(path))

    retokenize_general_book(path, submitter_g="G-0000")
    validate_book_template(path)


def build_templates(output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    _build_table(output_dir / _NAME)
    (output_dir / "base_text.docx").unlink(missing_ok=True)


def check_templates(output_dir: Path) -> bool:
    path = output_dir / _NAME
    if not path.is_file():
        return False
    try:
        validate_book_template(path)
    except ValueError:
        return False
    return True


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("output_dir", nargs="?")
    parser.add_argument("--check", action="store_true")
    args = parser.parse_args()

    if args.output_dir:
        out = Path(args.output_dir)
    else:
        from app.config import get_settings

        out = get_settings().data_dir / "book_templates"

    if args.check:
        if check_templates(out):
            print(f"OK: {_NAME} is present and valid in {out}")
            return 0
        print(f"FAIL: {_NAME} is missing or invalid in {out}", file=sys.stderr)
        return 1

    build_templates(out)
    print(f"Built {_NAME} in {out}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
