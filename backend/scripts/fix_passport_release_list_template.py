"""One-time, idempotent: repair the Passport Release List template geometry.

Two defects, both from the original Excel paste (verified by XML dump):

1. Only the first 3 data rows (slots 0-2) were ever formatted. Rows for slots
   3-14 carry `tcBorders` all `nil`, no `<w:jc w:val="center"/>`, and Times New
   Roman 10pt instead of Arial 11pt bold — so the 4th employee onward rendered
   as naked, left-shifted, borderless text hanging under the table. Fixed by
   cloning the styled slot-0 `<w:tr>` 14 times with the item index rewritten.

2. The paper was ALWAYS 2 pages, even with zero employees: the 28 rows sum to
   ~10946 twips (7.6in) against 6.0in usable (letter landscape, 1.25in top and
   bottom margins), so the signature block could never share page 1. Fixed by
   dropping the top/bottom margins to 0.5in, normalising every data row to a
   uniform 300-twip height, and shrinking the mandatory trailing body paragraph
   (a full-size empty paragraph after the table is enough on its own to push a
   blank second page).

The remaining half of defect 2 — the 15 fixed slots leaving phantom blank rows
when fewer employees are listed — is handled at render time by
`_pp_passport_release_list` in `docx_engine.py`, which deletes the unused rows.

The modified .docx is committed intentionally (template-churn rule: only THIS
file changes).

Usage:  venv\\Scripts\\python.exe backend\\scripts\\fix_passport_release_list_template.py
"""

from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches

TEMPLATE = (
    Path(__file__).resolve().parents[1]
    / "templates"
    / "GSSG-HR_300-004b_Passport_Release_List.docx"
)

FIRST_ITEM_ROW = 7  # table row index of item slot 0
ITEM_SLOTS = 15  # template capacity; matches MAX_ROWS in EmployeesTableField.tsx
SIGNATURE_ROW_0 = FIRST_ITEM_ROW + ITEM_SLOTS  # "تم إستلام..." lead line + receipt boxes
ROW_HEIGHT_TWIPS = "300"


def _restyle_item_rows(doc: Document) -> bool:
    """Clone the styled slot-0 row over slots 1-14, rewriting the item index."""
    table = doc.tables[0]
    template_tr = table.rows[FIRST_ITEM_ROW]._tr
    changed = False
    for slot in range(1, ITEM_SLOTS):
        row = table.rows[FIRST_ITEM_ROW + slot]
        if row.cells[0].text.strip() != f"{{{{ item({slot}, 'employee_id') }}}}":
            raise SystemExit(f"row {FIRST_ITEM_ROW + slot} is not item slot {slot} — aborting")
        borders = row.cells[0]._tc.tcPr.find(qn("w:tcBorders"))
        already_styled = (
            borders is not None and borders.find(qn("w:left")).get(qn("w:val")) != "nil"
        )
        if already_styled:
            continue
        new_tr = copy.deepcopy(template_tr)
        for node in new_tr.iter(qn("w:t")):
            if node.text and "item(0," in node.text:
                node.text = node.text.replace("item(0,", f"item({slot},")
        row._tr.addprevious(new_tr)
        row._tr.getparent().remove(row._tr)
        changed = True
    return changed


def _normalize_row_heights(doc: Document) -> bool:
    """Give every data row the same compact height (the paste left 458/394/414
    on the styled rows and 300 on the rest)."""
    table = doc.tables[0]
    changed = False
    for slot in range(ITEM_SLOTS):
        tr = table.rows[FIRST_ITEM_ROW + slot]._tr
        trPr = tr.get_or_add_trPr()
        height = trPr.find(qn("w:trHeight"))
        if height is None:
            height = trPr.makeelement(qn("w:trHeight"), {})
            trPr.append(height)
        if height.get(qn("w:val")) != ROW_HEIGHT_TWIPS:
            height.set(qn("w:val"), ROW_HEIGHT_TWIPS)
            changed = True
    return changed


def _tighten_page(doc: Document) -> bool:
    """0.5in top/bottom margins + a hairline trailing paragraph, so the
    signature block can share page 1."""
    changed = False
    section = doc.sections[0]
    if section.top_margin != Inches(0.5):
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        changed = True

    last = doc.paragraphs[-1]
    if last.text.strip():
        raise SystemExit("last body paragraph is not empty — aborting")
    pPr = last._p.get_or_add_pPr()
    rPr = pPr.find(qn("w:rPr"))
    if rPr is None:
        rPr = pPr.makeelement(qn("w:rPr"), {})
        pPr.insert(0, rPr)
    for tag in ("w:sz", "w:szCs"):
        el = rPr.find(qn(tag))
        if el is None:
            el = rPr.makeelement(qn(tag), {})
            rPr.append(el)
        if el.get(qn("w:val")) != "2":
            el.set(qn("w:val"), "2")
            changed = True
    return changed


def _keep_signature_block_together(doc: Document) -> bool:
    """A full 15-employee list is taller than one landscape page can hold, and
    Word was tearing the receipt block between its الإسم and ID lines. Chain the
    block's rows with keepNext so an unavoidable overflow moves it intact."""
    table = doc.tables[0]
    changed = False
    for row in table.rows[SIGNATURE_ROW_0:-1]:
        for cell in row.cells:
            for para in cell.paragraphs:
                if para.paragraph_format.keep_with_next is not True:
                    para.paragraph_format.keep_with_next = True
                    changed = True
    return changed


def main() -> None:
    doc = Document(str(TEMPLATE))
    changed = any(
        [
            _restyle_item_rows(doc),
            _normalize_row_heights(doc),
            _tighten_page(doc),
            _keep_signature_block_together(doc),
        ]
    )
    if not changed:
        print(f"{TEMPLATE.name}: already repaired — nothing to do.")
        return
    doc.save(str(TEMPLATE))
    print(f"{TEMPLATE.name}: repaired {ITEM_SLOTS} item rows + page geometry.")


if __name__ == "__main__":
    main()
