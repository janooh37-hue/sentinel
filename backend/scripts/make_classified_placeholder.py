"""Build the placeholder classified-book template docx.

PLACEHOLDER — this generates a minimal Word document with the correct Jinja
tokens so the service can render it immediately. The user MUST replace the
resulting file (backend/templates/GSSG-GS_301-001_Classified_Standard.docx)
with the real government letterhead layout once it is available.

Usage:
    python backend/scripts/make_classified_placeholder.py

The script is idempotent: re-running overwrites the file.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

OUTPUT = (
    Path(__file__).resolve().parents[1] / "templates" / "GSSG-GS_301-001_Classified_Standard.docx"
)


def _rtl_para(doc: Document, text: str, bold: bool = False) -> None:
    """Add a right-aligned RTL paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # RTL paragraph property
    pPr = p._p.get_or_add_pPr()
    bidi = pPr.find(qn("w:bidi"))
    if bidi is None:
        from docx.oxml import OxmlElement

        bidi = OxmlElement("w:bidi")
        pPr.insert(0, bidi)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(14)


def _guide_para(doc: Document, text: str) -> None:
    """Add a grey guide paragraph marking this as a placeholder."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run.font.size = Pt(10)
    run.italic = True


def main() -> None:
    doc = Document()

    # Guard paragraph — clearly marks this as a placeholder for the user
    _guide_para(
        doc,
        "*** PLACEHOLDER — replace with the real government letterhead layout ***",
    )

    _rtl_para(doc, "الرقم: {{ ref }}", bold=True)
    _rtl_para(doc, "التاريخ: {{ date }}")
    _rtl_para(doc, "السيد / {{ recipient_name }}")
    _rtl_para(doc, "الموضوع: {{ subject }}", bold=True)

    _guide_para(doc, "(body of the classified book goes here)")
    doc.add_paragraph()  # spacer

    _rtl_para(doc, "{{ cc }}")

    # Manager block — uses the SAME token names as the General Book template
    # (manager_name, manager_title) so the service layer can reuse identical
    # data-dict assembly.
    doc.add_paragraph()
    _rtl_para(doc, "{{ manager_name }}", bold=True)
    _rtl_para(doc, "{{ manager_title }}")

    # Author G-number footer — same token as the General Book footer
    _rtl_para(doc, "{{ submitter_g }}")

    _guide_para(
        doc,
        "*** END PLACEHOLDER — tokens: ref, date, recipient_name, subject, "
        "cc, manager_name, manager_title, submitter_g ***",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Written: {OUTPUT}")


if __name__ == "__main__":
    main()
