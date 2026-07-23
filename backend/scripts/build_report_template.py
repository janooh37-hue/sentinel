# backend/scripts/build_report_template.py
"""One-shot: derive the Report template from the General Book template.

Run once, then commit backend/templates/GSSG-GS_300-004_Report.docx:
    venv\\Scripts\\python.exe backend/scripts/build_report_template.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = Path("backend/templates/GSSG-GS_300-003_General_Book.docx")
DST = Path("backend/templates/GSSG-GS_300-004_Report.docx")


def _set_text(paragraph, text: str) -> None:
    """Replace the paragraph's text, keeping the first run's formatting."""
    runs = paragraph.runs
    if not runs:
        paragraph.add_run(text)
        return
    runs[0].text = text
    for r in runs[1:]:
        r._element.getparent().remove(r._element)


def _delete(paragraph) -> None:
    paragraph._element.getparent().remove(paragraph._element)


def main() -> None:
    doc = Document(str(SRC))

    def find(token: str):
        return next(p for p in doc.paragraphs if token in p.text)

    # 1) Remove the ref block entirely (guarantees no الرقم prints, whatever
    #    the caller passes).
    _delete(find("{%p if ref %}"))
    _delete(find("الرقم: {{ ref }}"))
    _delete(find("{%p endif %}"))  # the FIRST endif == the ref block's

    # 2) Closing courtesy + labelled, reordered author block. Re-find AFTER
    #    the deletions so we operate on the live tree.
    lut = find("للتفضل بالعلم وإجراءاتكم")
    _set_text(lut, "للتفضل بالعلم وإجراءاتكم لطفاً،،،")

    sig_para = find("{{ manager_sig }}")  # was first in the block (CENTER)
    name_para = find("{{ manager_name }}")  # bold 14pt
    title_para = find("{{ manager_title }}")  # bold 14pt

    # The blank CENTER paragraph directly above the signature holds the closing.
    closing = sig_para._element.getprevious()
    from docx.text.paragraph import Paragraph  # local import; avoids top clutter

    closing_para = Paragraph(closing, sig_para._parent) if closing is not None else None
    if closing_para is not None and not closing_para.text.strip():
        _set_text(closing_para, "وتفضلوا بقبول فائق الاحترام والتقدير،،،")
        closing_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Reorder by content swap (no element moves): top→bottom becomes name,
    # title, signature.
    for p, txt in (
        (sig_para, "الاسم: {{ manager_name }}"),
        (name_para, "المسمى الوظيفي: {{ manager_title }}"),
        (title_para, "التوقيع: {{ manager_sig }}"),
    ):
        _set_text(p, txt)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    DST.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(DST))
    print(f"wrote {DST}")


if __name__ == "__main__":
    main()
