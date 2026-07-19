"""One-time, idempotent: insert the guarded Arabic ref line above the date
line in the canonical General Book template. The modified .docx is committed
intentionally (template-churn rule: only THIS file changes).

Re-running on an already-edited template runs a REPAIR pass instead: it
ensures the {{ ref }} run carries <w:rtl/> — the EXACT encoding of the
hand-typed legacy books (verified by XML dump: their digit runs are
RTL-marked), which makes Word order the segments right-to-left so the
bumping serial reads LAST on the line. Both no-mark and a forced
<w:rtl w:val="0"/> made Word lay the value as one LTR unit with the serial
landing right next to الرقم:.
"""

import copy
import sys
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

TEMPLATE = Path(__file__).resolve().parents[1] / "templates" / "GSSG-GS_300-003_General_Book.docx"


def _repair_ref_run_direction(doc: Document) -> bool:
    """Ensure the {{ ref }} run carries <w:rtl/> (legacy-book encoding).
    Returns True if anything changed."""
    changed = False
    for para in doc.paragraphs:
        if "{{ ref }}" not in para.text:
            continue
        for run in para.runs:
            if "{{ ref }}" in run.text and run.font.rtl is not True:
                rpr = run._element.get_or_add_rPr()
                for el in rpr.findall(qn("w:rtl")):
                    rpr.remove(el)
                run.font.rtl = True
                changed = True
    return changed


def main() -> None:
    doc = Document(str(TEMPLATE))
    if any("{{ ref }}" in p.text for p in doc.paragraphs):
        if _repair_ref_run_direction(doc):
            doc.save(str(TEMPLATE))
            print("ref line present; ref run now RTL-marked (repair)")
        else:
            print("already has ref line; nothing to do")
        return
    date_p = next(p for p in doc.paragraphs if "التاريخ" in p.text and "{{ date }}" in p.text)

    def clone_empty_before() -> Paragraph:
        """Deep-copy the date paragraph (keeps RTL pPr/alignment), strip its
        runs, insert before date_p. Successive calls stack in call order."""
        new_p = copy.deepcopy(date_p._p)
        date_p._p.addprevious(new_p)
        para = Paragraph(new_p, date_p._parent)
        for r in list(para.runs):
            r._element.getparent().remove(r._element)
        return para

    src_run = date_p.runs[0]

    def style_like_date(run) -> None:
        run.font.name = src_run.font.name
        run.font.size = src_run.font.size
        run.font.bold = src_run.font.bold

    p_if = clone_empty_before()
    p_if.add_run("{%p if ref %}")

    p_ref = clone_empty_before()
    label = p_ref.add_run("الرقم: ")
    style_like_date(label)
    ref_run = p_ref.add_run("{{ ref }}")
    style_like_date(ref_run)
    ref_run.font.rtl = True  # legacy-book encoding — see module docstring

    p_endif = clone_empty_before()
    p_endif.add_run("{%p endif %}")

    doc.save(str(TEMPLATE))
    print("ref line added")


if __name__ == "__main__":
    sys.exit(main())
