"""Build backend/templates/GSSG-NAT_300-005_Inmate_Conduct_Violations.docx.

Takes the hand-made source docx and replaces its malformed placeholder tokens
with valid Jinja ones:

  * The two sample inmate rows collapse into ONE row wrapped in a
    ``{%tr for i in inmates %}`` / ``{%tr endfor %}`` pair, so the paper grows a
    row per inmate and prints no blank filler rows.
  * The three hard-coded supervisor-action bullets collapse into ONE bullet
    wrapped in ``{%p for a in actions %}`` / ``{%p endfor %}``.
  * ``{{Personas name}}`` / ``{( personal ID for persionar ))`` /
    ``{{ EMPLOYEE NAME }}`` / ``{{ G-NUMBER}}`` / ``{{MANAGER-SIGN}}`` are not
    legal Jinja (spaces, hyphens, wrong braces) — each is rewritten.
  * The inherited first-page footer's ``{{ submitter_id }}`` — a token no
    framework code ever fills — is rewritten to the canonical
    ``{{ submitter_g }}`` (see ``fix_footer_token``).

Tokens are split across ``w:r`` runs in the source, so a ``w:t`` string replace
misses them: every cell is rewritten run-and-all via ``set_cell_text``.

Run:  venv\\Scripts\\python.exe backend/scripts/build_inmate_violations_template.py "<source.docx>"
"""

from __future__ import annotations

import copy
import shutil
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

DEST = (
    Path(__file__).resolve().parents[1]
    / "templates"
    / "GSSG-NAT_300-005_Inmate_Conduct_Violations.docx"
)

#: Column order of the inmate row, right-to-left as Word lays it out.
ROW_TOKENS = (
    "{{ loop.index }}",
    "{{ i.name }}",
    "{{ i.nationality }}",
    "{{ i.wing }}",
    "{{ i.uid }}",
    "{{ i.holding_no }}",
)


def set_cell_text(tc: Any, text: str) -> None:
    """Replace a cell's whole text with `text`, keeping the first paragraph's
    formatting (font, size, RTL) and dropping every other run/paragraph."""
    paras = tc.findall(qn("w:p"))
    for extra in paras[1:]:
        tc.remove(extra)
    p = paras[0]
    for r in p.findall(qn("w:r")):
        p.remove(r)
    Paragraph(p, None).add_run(text)  # type: ignore[arg-type]


def set_para_text(para: Any, text: str) -> None:
    """Same, for a body paragraph — keep run 0's formatting, drop the rest."""
    for r in para.runs[1:]:
        r._r.getparent().remove(r._r)
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)


def fix_footer_token(doc: Any) -> None:
    """The source docx's first-page footer (``word/footer3.xml``) carries a
    stray, non-framework token — ``{{ submitter_id }}``, split across w:r
    runs same as the body tokens above — that nothing in the app ever fills.
    The canonical footer token every other form uses is ``{{ submitter_g }}``
    (see ``docx_engine.py``'s ``out.setdefault("submitter_g", "")``).

    The token's paragraph shares its table cell with a sibling
    "www.gss-group.net" paragraph — a cell-level rewrite (``set_cell_text``)
    would delete that sibling. Target the exact ``<w:p>`` instead, found by
    walking every paragraph in the footer part (tokens can be nested
    arbitrarily deep in tables), so only the token paragraph is touched.
    """
    footer_root = doc.sections[0].first_page_footer.part.element
    hits = 0
    for p in footer_root.iter(qn("w:p")):
        text = "".join(t.text or "" for t in p.iter(qn("w:t")))
        if "submitter_id" in text:
            set_para_text(Paragraph(p, None), "{{ submitter_g }}")  # type: ignore[arg-type]
            hits += 1
    if hits != 1:
        raise RuntimeError(
            f"expected exactly one paragraph with the stray submitter_id footer token, found {hits}"
        )


def build(src: Path) -> None:
    shutil.copy(src, DEST)
    doc = Document(str(DEST))
    tbl = doc.tables[1]
    trs = tbl._tbl.findall(qn("w:tr"))

    # --- 1. date / day / time: replace the three "Auto_edit" placeholders ---
    hdr = trs[0].findall(qn("w:tc"))
    set_cell_text(hdr[1], "{{ today }}")
    set_cell_text(hdr[3], "{{ weekday_ar }}")
    set_cell_text(hdr[5], "{{ now_time }}")

    # --- 2. inmate rows: token row becomes the loop body, sample row is cut ---
    data_row, sample_row = trs[2], trs[3]
    for tc, token in zip(data_row.findall(qn("w:tc")), ROW_TOKENS, strict=True):
        set_cell_text(tc, token)

    for_row = copy.deepcopy(sample_row)
    end_row = copy.deepcopy(sample_row)
    for marker_row, tag in ((for_row, "{%tr for i in inmates %}"), (end_row, "{%tr endfor %}")):
        cells = marker_row.findall(qn("w:tc"))
        set_cell_text(cells[0], tag)
        for c in cells[1:]:
            set_cell_text(c, "")
    data_row.addprevious(for_row)
    data_row.addnext(end_row)
    tbl._tbl.remove(sample_row)

    # --- 3. violation details ---
    detail_cell = tbl._tbl.findall(qn("w:tr"))[7].findall(qn("w:tc"))[0]
    set_cell_text(detail_cell, "{{ violation_details }}")

    # --- 4. supervisor actions: 3 fixed bullets -> one looped bullet ---
    acts_cell = tbl._tbl.findall(qn("w:tr"))[9].findall(qn("w:tc"))[0]
    bullets = acts_cell.findall(qn("w:p"))
    keep = bullets[0]  # keeps the ListParagraph + numPr
    for extra in bullets[1:]:
        acts_cell.remove(extra)
    set_para_text(Paragraph(keep, None), "{{ a }}")  # type: ignore[arg-type]
    for_p, end_p = copy.deepcopy(keep), copy.deepcopy(keep)
    set_para_text(Paragraph(for_p, None), "{%p for a in actions %}")  # type: ignore[arg-type]
    set_para_text(Paragraph(end_p, None), "{%p endfor %}")  # type: ignore[arg-type]
    keep.addprevious(for_p)
    keep.addnext(end_p)

    # --- 5. reporter row + manager signature block ---
    last = tbl._tbl.findall(qn("w:tr"))[-1].findall(qn("w:tc"))
    set_cell_text(last[1], "{{ reporter_name }}")
    set_cell_text(last[3], "{{ reporter_g }}")
    for para in doc.paragraphs:
        if "الإســـــــــم" in para.text:
            set_para_text(para, "الإســـــــــم: {{ manager_name }}")
        elif "MANAGER-SIGN" in para.text:
            set_para_text(para, "التوقيع: {{ manager_sig }}")

    # --- 6. footer: stray {{ submitter_id }} -> canonical {{ submitter_g }} ---
    fix_footer_token(doc)

    doc.save(str(DEST))
    print(f"wrote {DEST}")


if __name__ == "__main__":
    build(Path(sys.argv[1]))
